"""
template_import.py: DOCX template importer (C7, style axis; execution-plan chunk
C7's DOCX-first build order).

"render import-template corporate.docx" derives a template-profile.yaml FROM a
branded Word template so the FIRST render through render-doc.sh reproduces the
template's own look, no hand-written profile, no surgical post-edits. Consumes
docstyle/ooxml_theme.py (the shared DrawingML theme parser) plus python-docx's
own documented Styles/Font/Section APIs (adopted directly: these are real,
maintained APIs, unlike the theme part; see docs/prior-art-template-analysis.md
section 1).

Cascade scope (v1, deliberate): style-property resolution implements ONE level of
basedOn fallback (a style's own explicit value, else its base_style's explicit
value, else the theme). Full docx4j-style PropertyResolver cascading (following
basedOn arbitrarily deep, merging numbering/paragraph-format inheritance too) is
deliberately out of scope for v1, see docs/prior-art-template-analysis.md
section 2 (the docx4j pattern is [imitate], not adopted wholesale).

Only genuinely derivable keys are ever written uncommented into the generated
profile: honesty over guessing. Everything the importer could not derive from
THIS template is left as a commented line at the built-in default, with the
reason stated, rather than silently inventing a value.

The idempotency gate (--check PROBE.md) is the scoped style-diff the C7 prior-art
pass confirmed has no adoptable OSS implementation (docs/prior-art-template-
analysis.md section 5): it compares the DERIVED properties (not a general
effective-style differ) between the template and a probe rendered through the
real pipeline (render-doc.sh) with the derived profile applied.
"""

from __future__ import annotations

import argparse
import hashlib
import os
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

REPO_ROOT = Path(__file__).resolve().parent.parent

# Ensure this module's own directory is importable as a flat module (ooxml_theme
# is a sibling, not a subpackage) regardless of HOW template_import.py is
# invoked (direct script run, `from docstyle import template_import`, or a
# subprocess `python docstyle/template_import.py ...`). Mirrors the pattern
# roundtrip/provenance.py already uses for its own sibling import.
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))
from ooxml_theme import Theme, ThemeError, read_theme  # noqa: E402

# THEME keys docstyle/style_postprocess.py actually consumes, in the order the
# derived profile lists them.
PROFILE_KEY_ORDER = (
    "font", "accent", "body", "body_muted", "table_body", "zebra",
    "page_width_cm", "page_height_cm", "margin_cm",
)

# python-docx Font.color.theme_color members (docx.enum.dml.MSO_THEME_COLOR_INDEX)
# mapped onto the raw clrScheme role keys ooxml_theme.Theme.colors uses. This is
# the STANDARD/default clrMap (tx1=dk1, bg1=lt1, tx2=dk2, bg2=lt2); a document
# whose settings.xml overrides clrMap to something non-default is out of scope
# v1 (same cascade-depth tradeoff as the basedOn fallback above).
THEME_COLOR_ROLE_MAP = {
    "DARK_1": "dk1", "TEXT_1": "dk1",
    "LIGHT_1": "lt1", "BACKGROUND_1": "lt1",
    "DARK_2": "dk2", "TEXT_2": "dk2",
    "LIGHT_2": "lt2", "BACKGROUND_2": "lt2",
    "ACCENT_1": "accent1", "ACCENT_2": "accent2", "ACCENT_3": "accent3",
    "ACCENT_4": "accent4", "ACCENT_5": "accent5", "ACCENT_6": "accent6",
    "HYPERLINK": "hlink", "FOLLOWED_HYPERLINK": "folHlink",
}


# --------------------------------------------------------------------------
# Style / geometry extraction
# --------------------------------------------------------------------------

def _get_style(doc, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def _theme_color_hex(theme_color, theme: Theme) -> str | None:
    role = THEME_COLOR_ROLE_MAP.get(getattr(theme_color, "name", None))
    if role is None:
        return None
    return theme.colors.get(role)


def _resolve_color(font, theme: Theme) -> str | None:
    """Resolve a python-docx Font's effective colour to a 6-hex string: a
    theme-colour reference resolved against the parsed Theme if one is set, else
    an explicit RGB value, else None.

    Theme reference is checked FIRST, not RGB: per python-docx's own ColorFormat.
    rgb docstring, Word writes a "last known good" RGB value alongside a
    themeColor attribute when a theme colour is assigned, and "the theme color
    takes precedence at rendering time." Checking .rgb first (the more obvious
    read) silently returns that stale baked-in value instead of resolving the
    live theme reference (caught by hand: python-docx's own built-in default
    Heading 1 style carries both, and RGB-first returned Office's stock default
    accent, 365F91, instead of this template's actual accent1)."""
    color = font.color
    if color is None or color.type is None:
        return None
    if color.theme_color is not None:
        return _theme_color_hex(color.theme_color, theme)
    if color.rgb is not None:
        return str(color.rgb).upper()
    return None


def style_font_info(style, theme: Theme) -> dict[str, Any]:
    """{name, size_pt, bold, color} for a docx paragraph Style, with ONE level of
    basedOn fallback for any property the style itself leaves unset (see module
    docstring: full cascade is out of scope v1)."""
    font = style.font
    name = font.name
    size = font.size
    bold = font.bold
    color = _resolve_color(font, theme)

    base = style.base_style
    if base is not None:
        base_font = base.font
        if name is None:
            name = base_font.name
        if size is None:
            size = base_font.size
        if bold is None:
            bold = base_font.bold
        if color is None:
            color = _resolve_color(base_font, theme)

    return {
        "name": name,
        "size_pt": size.pt if size is not None else None,
        "bold": bold,
        "color": color,
    }


def extract_docx_styles(doc, theme: Theme) -> dict[str, dict[str, Any]]:
    """style_font_info() for every style relevant to C7 (Heading 1-3, Normal,
    Body Text) that is actually present in the document."""
    out: dict[str, dict[str, Any]] = {}
    for name in ("Normal", "Body Text", "Heading 1", "Heading 2", "Heading 3"):
        style = _get_style(doc, name)
        if style is not None:
            out[name] = style_font_info(style, theme)
    return out


def derive_style_font_overrides(doc, theme: Theme, global_font: str | None) -> dict[str, str]:
    """Walk EVERY paragraph w:style definition's w:rPr/w:rFonts (not just Normal),
    with the same one-level basedOn fallback style_font_info() already applies, and
    return only the styles whose resolved font is a GENUINE override: explicit,
    and different from global_font (the single value the top-level 'font' key
    would otherwise apply to every paragraph, see issue #97).

    global_font=None (the top-level key itself was not derivable) means there is
    no baseline to diff against, so any style with an explicit resolvable font is
    reported as an override. Keeps the profile minimal: a template where every
    style resolves to the same font yields an empty dict, not a redundant
    per-style entry that just repeats the global default."""
    from docx.enum.style import WD_STYLE_TYPE

    overrides: dict[str, str] = {}
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        try:
            name = style.name
        except Exception:
            continue
        if not name:
            continue
        info = style_font_info(style, theme)
        font_name = info["name"]
        if not font_name:
            continue
        if global_font is not None and font_name == global_font:
            continue
        overrides[name] = font_name
    return overrides


def _margins_cm(section) -> dict[str, float | None]:
    sides = {"top": section.top_margin, "bottom": section.bottom_margin,
             "left": section.left_margin, "right": section.right_margin}
    return {k: (round(v.cm, 2) if v is not None else None) for k, v in sides.items()}


def _uniform_margin_cm(section) -> float | None:
    sides = _margins_cm(section)
    values = list(sides.values())
    if any(v is None for v in values):
        return None
    if len(set(values)) == 1:
        return values[0]
    return None


def section_geometry_cm(section) -> dict[str, Any]:
    pw = section.page_width
    ph = section.page_height
    return {
        "page_width_cm": round(pw.cm, 2) if pw is not None else None,
        "page_height_cm": round(ph.cm, 2) if ph is not None else None,
        "margin_cm": _uniform_margin_cm(section),
        "_margin_sides_cm": _margins_cm(section),  # diagnostic only, not a profile key
    }


# --------------------------------------------------------------------------
# Derivation
# --------------------------------------------------------------------------

@dataclass
class DerivedProfile:
    derived: dict[str, Any] = field(default_factory=dict)
    not_derived: dict[str, str] = field(default_factory=dict)
    style_fonts: dict[str, str] = field(default_factory=dict)
    body_style: str | None = None


_NOT_DERIVABLE_ALWAYS = {
    "body_muted": "no corresponding OOXML style/theme concept is derivable from a DOCX "
                  "template in v1 scope (there is no 'muted body' role in DrawingML or "
                  "the Normal/Body Text styles)",
    "table_body": "no corresponding OOXML style/theme concept is derivable from a DOCX "
                  "template in v1 scope (pandoc's own Table style, not the source "
                  "template, drives rendered table cell text)",
    "zebra": "alternate-row table shading is a rendering choice, not a property a DOCX "
             "template's styles or theme expose",
}


def derive_profile(doc, theme: Theme) -> DerivedProfile:
    """Derive the THEME keys docstyle/style_postprocess.py consumes from a parsed
    template Document + Theme. Only keys genuinely derivable are set; everything
    else carries its non-derivability reason for the commented output."""
    styles = extract_docx_styles(doc, theme)
    normal = styles.get("Normal", {})
    body_text = styles.get("Body Text", {})
    h1 = styles.get("Heading 1", {})

    derived: dict[str, Any] = {}
    not_derived: dict[str, str] = {}

    font = normal.get("name") or theme.fonts.get("minor")
    if font:
        derived["font"] = font
    else:
        not_derived["font"] = ("the Normal style has no explicit font and the template's "
                               "theme has no minor-latin typeface")

    style_fonts = derive_style_font_overrides(doc, theme, derived.get("font"))

    # Default body-paragraph style detection (issue #101). pandoc's docx writer
    # styles an ordinary body paragraph "Normal" regardless of whether the
    # template defines its own dedicated body style (commonly named "Body" or
    # Word's built-in "Body Text") with its own font -- so that style sits
    # unused and the paragraph falls back to Normal's (often undefined) font.
    # Prefer "Body" (seen in real institutional templates) over "Body Text"
    # (Word's own built-in name) if a template somehow defines both distinctly.
    # Only counts as a real body style if it's already confirmed in style_fonts
    # (i.e. derive_style_font_overrides found it has a genuine font override,
    # distinct from Normal) -- a style with the right name but no distinguishing
    # font isn't worth remapping paragraphs into, nothing would change.
    body_style = None
    for candidate in ("Body", "Body Text"):
        if candidate in style_fonts:
            body_style = candidate
            break

    accent = h1.get("color") or theme.colors.get("accent1")
    if accent:
        derived["accent"] = accent
    else:
        not_derived["accent"] = ("the Heading 1 style has no explicit or theme-resolved "
                                 "colour and the template's theme has no accent1")

    body = normal.get("color") or body_text.get("color") or theme.colors.get("dk1")
    if body:
        derived["body"] = body
    else:
        not_derived["body"] = ("the Normal/Body Text styles have no explicit or "
                               "theme-resolved colour and the template's theme has no dk1")

    not_derived.update(_NOT_DERIVABLE_ALWAYS)

    geom = section_geometry_cm(doc.sections[0])
    if geom["page_width_cm"] is not None:
        derived["page_width_cm"] = geom["page_width_cm"]
    else:
        not_derived["page_width_cm"] = "the template's first section has no page width"
    if geom["page_height_cm"] is not None:
        derived["page_height_cm"] = geom["page_height_cm"]
    else:
        not_derived["page_height_cm"] = "the template's first section has no page height"
    if geom["margin_cm"] is not None:
        derived["margin_cm"] = geom["margin_cm"]
    else:
        sides = geom["_margin_sides_cm"]
        sides_str = ", ".join(f"{k}={v}" for k, v in sides.items())
        not_derived["margin_cm"] = (
            "the template's margins are not uniform on all four sides "
            f"({sides_str} cm) and the profile only supports one margin_cm value"
        )

    return DerivedProfile(derived=derived, not_derived=not_derived, style_fonts=style_fonts,
                          body_style=body_style)


def extract_rendered_style_properties(docx_path: Path) -> dict[str, Any]:
    """Re-extract the SAME property set from an already-rendered probe DOCX, for
    the --check idempotency gate. style_postprocess.py bakes accent/body/font
    directly onto RUNS (not style definitions), so this looks at the first
    heading-1 / body run's actual formatting rather than style definitions;
    page geometry is section-level either way."""
    from docx import Document

    doc = Document(str(docx_path))

    # Front-matter-driven styles (Title/Subtitle/Date) get accent/muted colouring
    # of their own in style_postprocess.py, not the body colour: excluded here
    # so a probe with a `title:` in its frontmatter doesn't get its Title
    # paragraph mistaken for a body paragraph (confirmed by hand: without this
    # exclusion, 'body' read back as the accent colour, not the body colour).
    NON_BODY_STYLES = {"Title", "Subtitle", "Date", "Author"}
    h1_run = None
    body_run = None
    for p in doc.paragraphs:
        style_name = p.style.name if p.style and p.style.name else ""
        if not p.runs:
            continue
        if h1_run is None and style_name == "Heading 1":
            h1_run = p.runs[0]
        if (body_run is None and p.text.strip()
                and not style_name.startswith("Heading") and style_name not in NON_BODY_STYLES):
            body_run = p.runs[0]
        if h1_run is not None and body_run is not None:
            break

    def _run_color(run) -> str | None:
        if run is None:
            return None
        rgb = run.font.color.rgb if run.font.color is not None else None
        return str(rgb).upper() if rgb is not None else None

    def _run_font(run) -> str | None:
        return run.font.name if run is not None else None

    geom = section_geometry_cm(doc.sections[0])
    return {
        "font": _run_font(h1_run) or _run_font(body_run),
        "accent": _run_color(h1_run),
        "body": _run_color(body_run),
        "page_width_cm": geom["page_width_cm"],
        "page_height_cm": geom["page_height_cm"],
        "margin_cm": geom["margin_cm"],
    }


# --------------------------------------------------------------------------
# Provenance + yaml rendering
# --------------------------------------------------------------------------

def _tool_version() -> str:
    sys.path.insert(0, str(REPO_ROOT / "roundtrip"))
    import provenance  # roundtrip/provenance.py

    return provenance.tool_version()


def build_import_provenance(template_path: Path, date_str: str | None = None) -> dict[str, str]:
    sha256 = hashlib.sha256(template_path.read_bytes()).hexdigest()
    return {
        "source_template": template_path.name,
        "source_sha256": sha256,
        "import_date": date_str or date.today().isoformat(),
        "tool_version": _tool_version(),
    }


def _yaml_value(key: str, value: Any) -> str:
    if key == "font":
        return f"{key}: {value}"
    if isinstance(value, str):
        return f'{key}: "{value}"'
    return f"{key}: {value}"


def render_profile_yaml(provenance: dict[str, str], dp: DerivedProfile) -> str:
    lines = [
        "# template-profile.yaml: derived by 'render import-template' (C7) from a "
        "corporate DOCX template.",
        f"# source template: {provenance['source_template']}",
        f"# source sha256:   {provenance['source_sha256']}",
        f"# import date:     {provenance['import_date']}",
        f"# tool_version:    {provenance['tool_version']}",
        "#",
        "# Consumed by docstyle/style_postprocess.py (--template-profile) and",
        "# docstyle/heading_numbering.py (--profile). Only keys this importer could",
        "# genuinely derive from the template are set below (honesty over guessing);",
        "# every other key stays commented, at the built-in default, with the reason",
        "# it could not be derived. Colours are 6-hex, no leading '#'.",
        "",
    ]
    for key in PROFILE_KEY_ORDER:
        if key in dp.derived:
            lines.append(_yaml_value(key, dp.derived[key]))
        else:
            reason = dp.not_derived.get(key, "not derivable from this template")
            lines.append(f"# {key}: not derivable from this template ({reason}); kept at built-in default")
    lines.append("")
    if dp.style_fonts:
        lines += [
            "# Per-style font overrides: this template defines a distinct font for these",
            "# named paragraph styles (derived by walking each w:style's w:rPr/w:rFonts,",
            "# one level of basedOn fallback, same as 'font' above). Falls back to 'font'",
            "# for any paragraph style not listed here.",
            "styles:",
        ]
        for name in sorted(dp.style_fonts):
            lines.append(f'  "{name}":')
            lines.append(f"    font: {dp.style_fonts[name]}")
    else:
        lines += [
            "# No per-style font overrides found: every paragraph style in this template",
            "# resolves to the same font as 'font' above.",
        ]
    if dp.body_style:
        lines += [
            "",
            "# Default body-paragraph style (issue #101): pandoc styles an ordinary body",
            '# paragraph "Normal" regardless of whether the template defines its own',
            f'# dedicated body style -- this template does ("{dp.body_style}", see styles:',
            "# above). style_postprocess.py remaps Normal-styled body paragraphs to this",
            "# style before font-styling runs, so its own font/size is respected via the",
            "# existing per-style-font-override + custom-style-preservation logic",
            "# (issues #97/#98) instead of falling back to Normal's (often undefined) font.",
            f'body_style: "{dp.body_style}"',
        ]
    else:
        lines += [
            "",
            "# body_style: not derivable -- no paragraph style named \"Body\" or \"Body Text\"",
            "# with its own distinct font was found in this template; ordinary body",
            "# paragraphs render via Normal as before.",
        ]
    lines += [
        "",
        "# Optional marking / cover behaviour: authoring hints, NOT derived from the",
        "# template (a DOCX template carries no marking-replacement or cover-label",
        "# data of this shape), copy the shape from template-profile-example.yaml",
        "# and fill in if this skin needs it:",
        "#",
        "# classification:",
        "#   header_footer_replacements:",
        '#     - find: ["INTERNAL USE ONLY"]',
        '#       replace: "INTERNAL USE ONLY (POLICY-REF)"',
        "#   brief_replacements: []",
        '#   strip_cover_banner_prefixes: ["Classification:", "Distribution:"]',
        "#   strip_cover_banner_contains: []",
        "#",
        "# cover:",
        '#   part_heading_prefix: "Part"',
        '#   version_label: "Version {version} - {date}"',
        "",
    ]
    return "\n".join(lines) + "\n"


# --------------------------------------------------------------------------
# --check idempotency gate
# --------------------------------------------------------------------------

@dataclass
class ComparisonRow:
    key: str
    expected: Any
    actual: Any
    status: str  # PASS | DRIFT | SKIP


def _values_match(key: str, expected: Any, actual: Any) -> bool:
    if key in ("page_width_cm", "page_height_cm", "margin_cm"):
        if expected is None or actual is None:
            return expected == actual
        return abs(float(expected) - float(actual)) <= 0.05
    if isinstance(expected, str) and isinstance(actual, str):
        return expected.strip().upper() == actual.strip().upper()
    return expected == actual


def compare_properties(expected: dict[str, Any], actual: dict[str, Any]) -> list[ComparisonRow]:
    """Per-property pass/drift/skip comparison. A property whose EXPECTED
    (template-derived) value is None (not derivable from the template) is
    SKIPPED rather than flagged, since there is nothing to check drift against."""
    rows = []
    for key in PROFILE_KEY_ORDER:
        if key not in expected and key not in actual:
            continue
        exp = expected.get(key)
        act = actual.get(key)
        if exp is None:
            rows.append(ComparisonRow(key, exp, act, "SKIP"))
            continue
        status = "PASS" if _values_match(key, exp, act) else "DRIFT"
        rows.append(ComparisonRow(key, exp, act, status))
    return rows


def format_comparison_table(rows: list[ComparisonRow]) -> str:
    header = f"{'property':<15} {'expected':<22} {'actual':<22} status"
    sep = "-" * len(header)
    lines = [header, sep]
    for r in rows:
        exp = "(not derivable)" if r.expected is None else str(r.expected)
        act = "-" if r.actual is None else str(r.actual)
        lines.append(f"{r.key:<15} {exp:<22} {act:<22} {r.status}")
    return "\n".join(lines)


def _find_bash() -> str | None:
    if os.name == "nt":
        program_files = os.environ.get("ProgramFiles", r"C:\Program Files")
        for cand in (
            Path(program_files) / "Git" / "bin" / "bash.exe",
            Path(program_files) / "Git" / "usr" / "bin" / "bash.exe",
        ):
            if cand.exists():
                return str(cand)
    for name in ("bash", "bash.exe"):
        found = shutil.which(name)
        if found:
            return found
    return None


def run_idempotency_check(probe_md: Path, template: Path, profile_path: Path,
                          expected: dict[str, Any]) -> int:
    """Render probe_md through container/render-doc.sh with TEMPLATE_DOCX=template
    and TEMPLATE_PROFILE=profile_path (env-set, subprocess bash, same pattern
    render.py's run_docx uses), then compare the re-extracted rendered properties
    against `expected`. Prints the pass/drift table; returns 1 on any drift, 0
    clean, 3 (skip cleanly) when bash or pandoc is unavailable."""
    bash = _find_bash()
    if not bash:
        print("SKIP: --check requires bash on PATH (render-doc.sh needs it; on "
              "Windows install git-bash/MSYS2).")
        return 3
    if not shutil.which("pandoc"):
        print("SKIP: --check requires pandoc on PATH.")
        return 3

    render_doc_sh = REPO_ROOT / "container" / "render-doc.sh"
    if not render_doc_sh.exists():
        print(f"SKIP: {render_doc_sh} not found.")
        return 3

    with tempfile.TemporaryDirectory(prefix="render-import-template-check-") as tmp:
        env = dict(os.environ)
        env["TEMPLATE_DOCX"] = str(template.resolve())
        env["TEMPLATE_PROFILE"] = str(profile_path.resolve())
        env["OUTPUT_DIR"] = tmp
        # Pin --profile reference: 'compact' (render-doc.sh's own default profile
        # NAME is 'reference', but style_postprocess.py's own default profile is
        # 'compact') colours body text from body_muted, not the derived 'body'
        # key at all: 'reference' is the only profile where THEME['body']
        # actually reaches rendered text, so it is the only profile this gate can
        # meaningfully check 'body' drift against. 'reference' also rebuilds a
        # cover page (drops the first Heading 1 UNLESS it starts with the
        # profile's part_heading_prefix, default "Part"): the probe must use a
        # Part-prefixed top heading (e.g. "# Part I: ...") or its Heading 1 gets
        # removed before extraction and 'accent' reads back as not-present.
        result = subprocess.run(
            [bash, str(render_doc_sh), str(probe_md), "--name", "c7-check", "--profile", "reference"],
            cwd=str(REPO_ROOT), env=env, capture_output=True, text=True, timeout=180,
        )
        if result.returncode != 0:
            print("ERROR: render-doc.sh failed during the --check probe render:")
            print(result.stdout[-3000:])
            print(result.stderr[-3000:])
            return 1

        produced = sorted(Path(tmp).glob("c7-check_*.docx"))
        if not produced:
            print("ERROR: --check probe render produced no .docx output.")
            print(result.stdout[-3000:])
            return 1
        actual = extract_rendered_style_properties(produced[-1])

    rows = compare_properties(expected, actual)
    print(format_comparison_table(rows))
    return 1 if any(r.status == "DRIFT" for r in rows) else 0


# --------------------------------------------------------------------------
# Guidance-doc scan (issue #100)
# --------------------------------------------------------------------------
#
# A branded template often ships alongside a SEPARATE document -- a policy or
# methodology paper explaining what each section is for, what's out of scope,
# how the document fits the surrounding process. import-template previously had
# no awareness that this second artifact usually exists; mining it for
# authoring doctrine was a fully manual, easy-to-forget step done after the
# fact. This is deliberately a MECHANICAL structural scan (heading/paragraph
# counts + a heading-text preview), not extraction: judging what the doctrine
# actually says and seeding editorial-doctrine.yaml (issue #84's concept, not
# yet built) from it is a judgment-heavy summarization task left to the
# operator, not automated here.

_HEADING_STYLE_PREFIXES = ("Heading", "Title", "Subtitle")
_MD_HEADING_RE = re.compile(r"^#{1,6}\s+(.+)$")


class GuidanceScanError(Exception):
    """Raised when --guidance-doc points at a file type this scan can't read."""


@dataclass
class GuidanceScan:
    path: Path
    heading_count: int
    paragraph_count: int
    headings: list[str]          # preview, capped at heading_preview_cap
    headings_truncated: bool


def scan_guidance_doc(path: Path, heading_preview_cap: int = 12) -> GuidanceScan:
    """Mechanical structural scan of an accompanying guidance/usage document.

    .docx: a paragraph is a heading if its style name starts with
    Heading/Title/Subtitle (same convention KNOWN_NON_CUSTOM_STYLE_NAMES-
    adjacent code elsewhere in this repo already uses), else it counts as a
    body paragraph if non-empty. .md/.markdown/.txt: an ATX heading line
    (# through ######) starting a blank-line-separated block counts as a
    heading; any other non-empty block counts as a paragraph.
    """
    suffix = path.suffix.lower()
    headings: list[str] = []
    paragraph_count = 0

    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name if p.style else "") or ""
            if style_name.startswith(_HEADING_STYLE_PREFIXES):
                headings.append(text)
            else:
                paragraph_count += 1
    elif suffix in (".md", ".markdown", ".txt"):
        text = path.read_text(encoding="utf-8", errors="replace")
        for block in re.split(r"\n\s*\n", text.strip()):
            block = block.strip()
            if not block:
                continue
            m = _MD_HEADING_RE.match(block.splitlines()[0])
            if m:
                headings.append(m.group(1).strip())
            else:
                paragraph_count += 1
    else:
        raise GuidanceScanError(
            f"unsupported --guidance-doc type: {suffix or '(no extension)'} "
            "(expected .docx, .md, .markdown, or .txt)"
        )

    return GuidanceScan(
        path=path,
        heading_count=len(headings),
        paragraph_count=paragraph_count,
        headings=headings[:heading_preview_cap],
        headings_truncated=len(headings) > heading_preview_cap,
    )


def format_guidance_scan_report(scan: GuidanceScan) -> str:
    lines = [
        f"Guidance-doc scan: {scan.path}",
        f"  {scan.heading_count} section heading(s), {scan.paragraph_count} paragraph(s) "
        "that look like authoring guidance",
    ]
    if scan.headings:
        preview = "; ".join(scan.headings)
        if scan.headings_truncated:
            preview += "; ..."
        lines.append(f"  Headings: {preview}")
    lines.append(
        "  Mechanical structural scan only, nothing extracted or summarized -- consider "
        "seeding editorial-doctrine.yaml from this by hand (issue #84 covers that schema)."
    )
    return "\n".join(lines)


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        prog="render import-template",
        description="Derive a template-profile.yaml (style axis) from a branded corporate "
                    "DOCX template, so the first render through render-doc.sh reproduces "
                    "the template's own look.",
    )
    ap.add_argument("template", type=Path, help="corporate .docx template to import")
    ap.add_argument("--out-dir", type=Path, default=Path("skin"),
                    help="output directory for the derived skin (default: ./skin)")
    ap.add_argument("--copy-reference", action="store_true",
                    help="also copy the template into --out-dir as reference.docx")
    ap.add_argument("--date", default=None,
                    help="import date stamped into the provenance header (YYYY-MM-DD); "
                         "default: today")
    ap.add_argument("--guidance-doc", type=Path, default=None, metavar="DOC",
                    help="an accompanying style/usage guide for this template (.docx/.md/.txt), "
                         "if one exists -- gets a mechanical structural scan (heading/paragraph "
                         "count + heading preview) surfaced back to you as a pointer toward "
                         "seeding editorial-doctrine.yaml (issue #84); not auto-extracted")
    ap.add_argument("--check", type=Path, default=None, metavar="PROBE.md",
                    help="idempotency gate: render PROBE.md through render-doc.sh (--profile "
                         "reference, the only profile that actually uses the derived 'body' "
                         "key) with this template + the derived profile, then compare. PROBE.md "
                         "needs a Part-prefixed top heading, e.g. '# Part I: ...' (the default "
                         "cover.part_heading_prefix), or its Heading 1 is dropped by the "
                         "reference-profile cover rebuild before this gate can see it")
    args = ap.parse_args(argv)

    template = args.template
    if not template.exists():
        print(f"ERROR: template not found: {template}", file=sys.stderr)
        return 2
    if template.suffix.lower() != ".docx":
        print(f"ERROR: {template} is not a .docx (PPTX/XLSX import is a later C7 build order "
              "step, DOCX first)", file=sys.stderr)
        return 2

    try:
        theme = read_theme(template)
    except ThemeError as exc:
        print(f"WARNING: {exc}; continuing with style-only derivation (no theme fallback).",
              file=sys.stderr)
        theme = Theme(colors={}, fonts={})

    from docx import Document

    doc = Document(str(template))
    dp = derive_profile(doc, theme)

    args.out_dir.mkdir(parents=True, exist_ok=True)
    provenance = build_import_provenance(template, args.date)
    profile_path = args.out_dir / "template-profile.yaml"
    profile_path.write_text(render_profile_yaml(provenance, dp), encoding="utf-8")
    print(f"Derived: {profile_path}")
    print(f"  {len(dp.derived)} key(s) derived: {', '.join(sorted(dp.derived))}")
    print(f"  {len(dp.not_derived)} key(s) not derivable, kept at built-in default: "
          f"{', '.join(sorted(dp.not_derived))}")
    if dp.style_fonts:
        print(f"  {len(dp.style_fonts)} per-style font override(s) derived: "
              f"{', '.join(sorted(dp.style_fonts))}")

    template_docx_env = template.resolve()
    if args.copy_reference:
        reference_path = args.out_dir / "reference.docx"
        shutil.copy2(template, reference_path)
        print(f"Copied reference: {reference_path}")
        template_docx_env = reference_path.resolve()

    print()
    print("Consumer wrapper env:")
    print(f"TEMPLATE_DOCX={template_docx_env}")
    print(f"TEMPLATE_PROFILE={profile_path.resolve()}")

    print()
    if args.guidance_doc is not None:
        if not args.guidance_doc.exists():
            print(f"ERROR: --guidance-doc not found: {args.guidance_doc}", file=sys.stderr)
            return 2
        try:
            scan = scan_guidance_doc(args.guidance_doc)
        except GuidanceScanError as exc:
            print(f"ERROR: {exc}", file=sys.stderr)
            return 2
        print(format_guidance_scan_report(scan))
    else:
        # A printed reminder, not a blocking stdin prompt (this CLI has no other
        # interactive input anywhere, and a blocking prompt would hang CI/scripted
        # runs) -- but still an active nudge at the one moment an operator has both
        # artifacts in hand and is thinking about this template, per issue #100.
        print("No --guidance-doc given. If this template ships with a separate style/usage "
              "guide (a policy paper explaining what each section is for), re-run with "
              "--guidance-doc <path> for a structural pointer toward seeding "
              "editorial-doctrine.yaml (issue #100/#84) -- easy to forget once the template "
              "itself has been imported.")

    if args.check is not None:
        if not args.check.exists():
            print(f"ERROR: --check probe not found: {args.check}", file=sys.stderr)
            return 2
        print()
        print(f"Running idempotency check against probe {args.check} ...")
        # Resolve all three paths BEFORE handing off: run_idempotency_check
        # invokes render-doc.sh with cwd=REPO_ROOT, so a probe path relative to
        # the user's own cwd would otherwise resolve against the wrong base
        # (found by an end-to-end run from a scratch directory).
        return run_idempotency_check(args.check.resolve(), template.resolve(),
                                     profile_path.resolve(), dp.derived)

    return 0


if __name__ == "__main__":
    sys.exit(main())

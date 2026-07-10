"""
Post-processor: apply a configurable house style to pandoc-rendered DOCX.

Default look (every value overridable via --template-profile <yaml>, see
template-profile-example.yaml):
- Font: Arial throughout (per-named-style overrides via the profile's optional
  `styles:` block, issue #97, falling back to this global font otherwise)
- Body text: 8.5pt (#595959), justified (compact profile)
- H1: 12pt bold accent colour, page break before
- H2: 10.5pt bold accent colour
- H3: 9pt bold accent colour
- Tables: thin borders, header row filled with the accent colour, white bold text
- Table body: 8.5pt #333333, first col bold
- Tables: centered, content-hugging widths
- Page: A4, 2cm all margins
- Headers/footers: centered, from the pandoc reference doc

The profile mechanism: everything organisation-specific (palette, font, geometry,
header/footer marking replacements, cover-banner cleanup, cover labels,
punctuation normalization) is plain data in an optional profile yaml. With no
profile, the neutral defaults below apply and no marking edits are made.

Usage:
  python style_postprocess.py <input.docx> [output.docx]
      [--profile compact|reference] [--template-profile <yaml>]
      [--table-widths <yaml>] [--cover-version <v>] [--cover-date <d>]

`render docx` (render-doc.sh) invokes this module directly as a subprocess as the
house-style pass of its own pipeline. The same flags are also reachable standalone
via `render docstyle <input.docx> [output.docx] ...` (see render.py) for callers
who want this post-processor's capabilities (e.g. --table-widths) without going
through the full docx pipeline (issue #74).
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree

def _rgb(hexstr):
    h = str(hexstr).lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# Themeable house-style defaults (a neutral default look). Override the whole
# palette / font / geometry with --template-profile <yaml> (e.g. a profile derived
# from a corporate Word template by a template-analysis stage). When no profile is
# passed these defaults apply, so existing callers render identically: the profile
# mechanism is purely additive.
THEME = {
    'font': 'Arial',
    'accent': '1F3864',        # navy: headings, table header-row fill, cover line
    'body': '262626',          # reference-profile body (near-black, reader-friendly)
    'body_muted': '595959',    # subtitle / compact body / figure captions
    'table_body': '333333',    # table cell text
    'zebra': 'F2F4F7',         # subtle alternate-row shading
    'page_width_cm': 21.0,     # A4
    'page_height_cm': 29.7,
    'margin_cm': 2.0,
}

NAVY = _rgb(THEME['accent'])
GREY_BODY = _rgb(THEME['body_muted'])
GREY_TABLE = _rgb(THEME['table_body'])
BODY_DARK = _rgb(THEME['body'])          # near-black, reader-friendly body
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = THEME['font']
ZEBRA_FILL = THEME['zebra']              # subtle alternate-row shading
HDR_FILL = THEME['accent']               # table header-row fill (hex string)

# Per-named-style font overrides (issue #97): a template-profile's optional
# `styles: {StyleName: {font: ...}}` block, keyed by paragraph style NAME.
# Empty by default (no profile applied): every paragraph then uses FONT_NAME,
# unchanged from before this key existed. Populated by apply_template_profile.
STYLE_FONT_OVERRIDES = {}

# Callout (info-box) palette: light fill + an accent left bar + a label colour.
CALLOUT_INFO_FILL,   CALLOUT_INFO_BORDER,   CALLOUT_INFO_LABEL   = 'EAF3FB', '2F6FB6', '1F4E79'  # blue
CALLOUT_WARN_FILL,   CALLOUT_WARN_BORDER,   CALLOUT_WARN_LABEL   = 'FDF3E7', 'C77700', '8A5300'  # amber

# Headings that should get a page break before them (section-level H1s)
PAGE_BREAK_PREFIXES = []  # applied to ALL H1 by default

# Style profiles. 'compact' is a dense annex-style profile (8.5pt, justified) and
# is the default, so existing renders are untouched. 'reference' is a
# reader-friendly profile for widely-shared reference documents (larger body,
# left-aligned, generous spacing, zebra tables, a cover page). Select via --profile.
PROFILES = {
    'compact': dict(
        body=8.5, body_color=GREY_BODY, justify=True, line=None, space_after=None,
        title=14, sub=11, h1=12, h2=10.5, h3=9, h4=8.5,
        table_hdr=8.5, table_body=8.5, zebra=False, cover=False),
    'reference': dict(
        body=10.5, body_color=BODY_DARK, justify=False, line=1.18, space_after=6,
        title=22, sub=12.5, h1=16, h2=13, h3=11, h4=10,
        table_hdr=9.5, table_body=9.5, zebra=True, cover=True),
}

# Optional marking behaviour, all profile-driven with neutral/empty defaults.
# Each *_replacements entry is a mapping {find: [str, ...], replace: str}.
CLASSIFICATION = {
    'header_footer_replacements': [],   # applied in fix_header_footer_text
    'brief_replacements': [],           # applied by apply_brief_classification_marking
    'strip_cover_banner_prefixes': ['Classification:', 'Distribution:'],
    'strip_cover_banner_contains': [],
}

# Cover-page behaviour (reference profile).
COVER = {
    'part_heading_prefix': 'Part',      # H1 prefix that marks the document body start
    'version_label': 'Version {version} - {date}',   # cover line template
}

# House punctuation normalization (unicode dashes/quotes/NBSP to ASCII) on the
# body and headers/footers. Disable with `normalize_punctuation: false`.
NORMALIZE_PUNCTUATION = True


def apply_template_profile(path):
    """Merge a template-profile yaml over the built-in defaults and re-derive the
    theme globals + the per-profile body colours. Theme keys: font, accent, body,
    body_muted, table_body, zebra, page_width_cm, page_height_cm, margin_cm.
    Optional top-level blocks:
      classification: header_footer_replacements / brief_replacements
                      (each a list of {find: [str, ...], replace: str}) +
                      strip_cover_banner_prefixes / strip_cover_banner_contains
      cover:          part_heading_prefix, version_label
      normalize_punctuation: bool (default true)
      styles:         {StyleName: {font: ...}}, per-named-style font override
                      (issue #97), falling back to 'font' for any style not
                      listed. Wholly replaces any previously-applied profile's
                      style overrides (not merged), same replace-on-apply
                      semantics as the theme keys above.
    Unknown keys are ignored. Call BEFORE styling. With no profile, the neutral
    built-in defaults apply."""
    global NAVY, GREY_BODY, GREY_TABLE, BODY_DARK, FONT_NAME, ZEBRA_FILL, HDR_FILL
    global NORMALIZE_PUNCTUATION, STYLE_FONT_OVERRIDES
    import yaml
    with open(path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f) or {}
    THEME.update({k: v for k, v in data.items() if k in THEME})
    NAVY = _rgb(THEME['accent'])
    GREY_BODY = _rgb(THEME['body_muted'])
    GREY_TABLE = _rgb(THEME['table_body'])
    BODY_DARK = _rgb(THEME['body'])
    FONT_NAME = THEME['font']
    ZEBRA_FILL = THEME['zebra']
    HDR_FILL = THEME['accent']
    # PROFILES was built at import time with the original colour objects; re-point the
    # body-colour entries at the (possibly re-themed) globals.
    PROFILES['compact']['body_color'] = GREY_BODY
    PROFILES['reference']['body_color'] = BODY_DARK
    # Marking / cover / punctuation blocks (only keys present override defaults).
    cls = data.get('classification') or {}
    if isinstance(cls, dict):
        for key in CLASSIFICATION:
            if key in cls:
                CLASSIFICATION[key] = list(cls[key] or [])
    cov = data.get('cover') or {}
    if isinstance(cov, dict):
        for key in COVER:
            if key in cov and cov[key] is not None:
                COVER[key] = cov[key]
    if 'normalize_punctuation' in data:
        NORMALIZE_PUNCTUATION = bool(data['normalize_punctuation'])
    styles = data.get('styles') or {}
    overrides = {}
    if isinstance(styles, dict):
        for style_name, val in styles.items():
            if isinstance(val, dict) and val.get('font'):
                overrides[style_name] = val['font']
    STYLE_FONT_OVERRIDES = overrides


def set_para_spacing(para, line=None, space_after=None, space_before=None):
    pf = para.paragraph_format
    if line is not None:
        pf.line_spacing = line
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)


def set_repeat_header_row(table):
    """Mark row 0 as a repeating header (w:trPr/w:tblHeader) so it reprints
    at the top of every page a long table spans."""
    if not table.rows:
        return
    tr = table.rows[0]._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = etree.SubElement(tr, qn('w:trPr'))
        tr.insert(0, trPr)
    if trPr.find(qn('w:tblHeader')) is None:
        etree.SubElement(trPr, qn('w:tblHeader')).set(qn('w:val'), 'true')


def set_table_cell_margins(table, top=40, bottom=40, left=90, right=90):
    """Set uniform cell padding (twips) for breathing room in cells."""
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn('w:tblCellMar'))
    if old is not None:
        tblPr.remove(old)
    mar = etree.SubElement(tblPr, qn('w:tblCellMar'))
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = etree.SubElement(mar, qn(f'w:{side}'))
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')


def center_figures_and_captions(doc):
    """Centre any paragraph that holds an inline image; style the caption that
    pandoc emits (alt text) as centred italic small muted-grey, visually
    distinct from body text.

    Backported 2026-07-03 from the reference consumer's fix (found there the
    same day): style-name matching must be NORMALISED and must also check the
    raw pStyle id. Pandoc references 'ImageCaption' (no space), a style the
    reference.docx typically does not define, so python-docx resolves p.style
    to the document default and a name-based exact match never fires, leaving
    captions indistinguishable from body text. The whole caption line gets one
    uniform font so a 'Figure:' prefix cannot diverge."""
    muted = _rgb(THEME['body_muted'])
    for p in doc.paragraphs:
        if p._element.findall('.//' + qn('w:drawing')):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(p, space_before=6, space_after=2)
        pPr = p._element.find(qn('w:pPr'))
        ps = pPr.find(qn('w:pStyle')) if pPr is not None else None
        raw_id = (ps.get(qn('w:val')) if ps is not None else '') or ''
        style_name = (p.style.name if p.style and p.style.name else '')
        norm_ids = {raw_id.replace(' ', '').lower(), style_name.replace(' ', '').lower()}
        if norm_ids & {'imagecaption', 'caption', 'captionedfigure', 'figurecaption'}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(p, space_before=2, space_after=10)
            for r in p.runs:
                set_run_style(r, 8, muted, bold=False)
                r.font.italic = True


def _cover_version_label(version, date_str):
    """Build the cover line from cover.version_label, substituting {version} and
    {date}. Segments (split on ' - ') whose placeholder value is None are dropped,
    so a missing version or date degrades gracefully."""
    import re as _re
    template = COVER.get('version_label') or 'Version {version} - {date}'
    values = {'version': version, 'date': date_str}
    parts = []
    for seg in template.split(' - '):
        fields = _re.findall(r'\{(\w+)\}', seg)
        if any(values.get(f) is None for f in fields):
            continue
        parts.append(seg.format(**{f: values.get(f) for f in fields}))
    return ' - '.join(parts)


def build_reference_cover(doc, version=None, date_str=None):
    """Turn page 1 into a clean cover: drop the duplicate title H1, set a single
    version/date line, move the (sdt-wrapped) table of contents onto its own page
    just before the first part heading, and keep the running header/footer on the
    cover. The banner-cleanup literals and the part-heading prefix come from the
    profile (classification.strip_cover_banner_* and cover.part_heading_prefix)."""
    body = doc.element.body
    part_prefix = COVER['part_heading_prefix']
    prefixes = CLASSIFICATION['strip_cover_banner_prefixes'] or []
    contains = CLASSIFICATION['strip_cover_banner_contains'] or []

    # 0. Remove verbose / duplicate marking banners from the front matter
    #    (the page marking lives in the header/footer, not the cover body).
    for p in list(doc.paragraphs):
        nm = p.style.name if p.style and p.style.name else ''
        if nm == 'Heading 1' and p.text.strip().startswith(part_prefix):
            break  # reached the body
        t = p.text.strip()
        if any(t.startswith(pre) for pre in prefixes) or any(c in t for c in contains):
            p._element.getparent().remove(p._element)

    # 1. Remove the duplicate title H1 (first Heading 1 that is not a part divider).
    for p in list(doc.paragraphs):
        nm = p.style.name if p.style and p.style.name else ''
        if nm == 'Heading 1' and not p.text.strip().startswith(part_prefix):
            p._element.getparent().remove(p._element)
            break

    # 2. Replace the Date line with a clean version/date cover line.
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Date':
            label = _cover_version_label(version, date_str)
            if p.runs:
                p.runs[0].text = label
                for r in p.runs[1:]:
                    r.text = ''
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_style(r, 12, NAVY, bold=True)
            break

    # 3. Move the TOC sdt to just before the first part heading, on its own page.
    toc_sdt = None
    for el in body.iterchildren():
        if el.tag == qn('w:sdt') and b'TOC' in etree.tostring(el):
            toc_sdt = el
            break
    part1 = None
    for p in doc.paragraphs:
        if p.style and p.style.name == 'Heading 1' and p.text.strip().startswith(part_prefix):
            part1 = p._element
            break
    if toc_sdt is not None and part1 is not None:
        toc_sdt.getparent().remove(toc_sdt)
        part1.addprevious(toc_sdt)
        # page-break before the TOC heading (first <w:p> inside the sdtContent)
        inner_p = toc_sdt.find('.//' + qn('w:sdtContent') + '/' + qn('w:p'))
        if inner_p is not None:
            pPr = inner_p.find(qn('w:pPr'))
            if pPr is None:
                pPr = inner_p.makeelement(qn('w:pPr'), {})
                inner_p.insert(0, pPr)
            if pPr.find(qn('w:pageBreakBefore')) is None:
                etree.SubElement(pPr, qn('w:pageBreakBefore'))

    # 4. Keep the running header/footer on the cover too (page markings must
    #    appear on every page, cover included).
    doc.sections[0].different_first_page_header_footer = False


def set_run_style(run, size_pt, color, bold=None, font_name=None):
    if font_name is None:                # resolve at call time so --template-profile
        font_name = FONT_NAME            # font overrides are picked up
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.name = font_name
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(run._element, qn('w:rPr'))
    el = rPr.find(qn('w:rFonts'))
    if el is None:
        el = etree.SubElement(rPr, qn('w:rFonts'))
    el.set(qn('w:ascii'), font_name)
    el.set(qn('w:hAnsi'), font_name)


def set_para_font(para, size_pt, color, bold=None, alignment=None, font_name=None):
    for run in para.runs:
        set_run_style(run, size_pt, color, bold, font_name)
    if alignment is not None:
        para.alignment = alignment


def _style_font(style_name):
    """The per-style font override (issue #97) for style_name, or None to fall
    back to the profile's single global FONT_NAME (set_run_style's own default
    when font_name=None). Resolved per-call, same "read at call time" pattern
    set_run_style already uses for FONT_NAME, so a --template-profile applied
    just before styling is picked up."""
    return STYLE_FONT_OVERRIDES.get(style_name)


def add_page_break_before(para):
    """Add a page break before a paragraph via w:pPr/w:pageBreakBefore."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(para._element, qn('w:pPr'))
    # Use pageBreakBefore
    existing = pPr.find(qn('w:pageBreakBefore'))
    if existing is None:
        etree.SubElement(pPr, qn('w:pageBreakBefore'))


def make_thin_borders():
    borders = etree.SubElement(etree.Element('dummy'), qn('w:tblBorders'))
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(borders, qn(f'w:{side}'))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
    return borders


def set_cell_shading(cell, fill_color):
    tcPr = cell._element.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = etree.SubElement(cell._element, qn('w:tcPr'))
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = etree.SubElement(tcPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)


def style_table(table, first_col_bold=True, hdr_size=8.5, body_size=8.5, zebra=False):
    """Apply the house table style: accent-filled header row, thin borders,
    centered. Optionally zebra-stripe body rows + repeat the header row across
    pages."""
    tbl_pr = table._tbl.find(qn('w:tblPr'))
    if tbl_pr is not None:
        # Remove Word table style
        for tag in [qn('w:tblStyle'), qn('w:tblLook')]:
            el = tbl_pr.find(tag)
            if el is not None:
                tbl_pr.remove(el)
        # Replace borders
        old = tbl_pr.find(qn('w:tblBorders'))
        if old is not None:
            tbl_pr.remove(old)
        tbl_pr.append(make_thin_borders())

        # Center table
        jc = tbl_pr.find(qn('w:jc'))
        if jc is None:
            jc = etree.SubElement(tbl_pr, qn('w:jc'))
        jc.set(qn('w:val'), 'center')

        # Keep pandoc-generated column widths (content-hugging)
        # Do NOT override to auto-width: that stretches tables to full page

    set_table_cell_margins(table)
    set_repeat_header_row(table)

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if ri == 0:
                set_cell_shading(cell, HDR_FILL)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_run_style(r, hdr_size, WHITE, bold=True)
            else:
                if zebra and (ri % 2 == 1):
                    set_cell_shading(cell, ZEBRA_FILL)
                for p in cell.paragraphs:
                    for r in p.runs:
                        is_first = (ci == 0) and first_col_bold
                        set_run_style(r, body_size, GREY_TABLE, bold=True if is_first else False)


def _section_text_width_twips(doc):
    """Available text width (twips) of the first section: pgSz width minus L+R margins.

    EMU -> twips: 1 twip = 635 EMU. python-docx exposes Emu objects (.twips attr in
    1.x). Fall back to arithmetic if .twips is unavailable.
    """
    sec = doc.sections[0]
    def tw(emu):
        if emu is None:
            return 0
        try:
            return emu.twips
        except AttributeError:
            return int(round(int(emu) / 635))
    return tw(sec.page_width) - tw(sec.left_margin) - tw(sec.right_margin)


def _set_tbl_layout_fixed(table):
    tblPr = table._tbl.tblPr
    lay = tblPr.find(qn('w:tblLayout'))
    if lay is None:
        lay = etree.SubElement(tblPr, qn('w:tblLayout'))
    lay.set(qn('w:type'), 'fixed')


def _set_tbl_width(table, twips):
    tblPr = table._tbl.tblPr
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tw = etree.SubElement(tblPr, qn('w:tblW'))
    tw.set(qn('w:w'), str(twips))
    tw.set(qn('w:type'), 'dxa')


def _set_tbl_grid(table, widths):
    """Replace the tblGrid with explicit gridCol widths (twips)."""
    tbl = table._tbl
    old = tbl.find(qn('w:tblGrid'))
    if old is not None:
        tbl.remove(old)
    # tblGrid must come right after tblPr
    grid = etree.Element(qn('w:tblGrid'))
    for w in widths:
        gc = etree.SubElement(grid, qn('w:gridCol'))
        gc.set(qn('w:w'), str(w))
    tblPr = tbl.tblPr
    tblPr.addnext(grid)


def _set_cell_widths(table, widths):
    """Set per-cell w:tcW (dxa) for every row to the given column widths (twips)."""
    for row in table.rows:
        cells = row.cells
        # Guard against merged/uneven rows: only set as many cells as we have widths
        for ci, cell in enumerate(cells):
            if ci >= len(widths):
                break
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = etree.SubElement(tcPr, qn('w:tcW'))
            tcW.set(qn('w:w'), str(widths[ci]))
            tcW.set(qn('w:type'), 'dxa')


def apply_table_widths(doc, table_specs):
    """Apply operator-fitted column widths to tables, full-width + fixed layout.

    GENERIC + parametrisable so any document can inherit the "compact,
    content-fitted, full-width" table doctrine. `table_specs` is a list of
    column-width lists (twips), matched to document tables by ORDINAL (0-based).
    Each spec's relative proportions are preserved and scaled to fill the actual
    section text-width, so tables stay full-width regardless of the margin spec.
    A spec whose column count does not match the table is skipped (a guard
    against source table-order drift).

    Returns a list of (table_index, applied_widths) for reporting.
    """
    text_w = _section_text_width_twips(doc)
    applied = []
    for ti, table in enumerate(doc.tables):
        if ti >= len(table_specs):
            break
        spec = table_specs[ti]
        if not spec:
            continue
        ncols = len(table.columns)
        if len(spec) != ncols:
            # column count mismatch -> skip rather than mis-fit
            continue
        total = sum(spec)
        if total <= 0:
            continue
        # Scale proportionally to the actual text width (full-width intent).
        scaled = [int(round(w * text_w / total)) for w in spec]
        # Fix rounding drift so the columns sum exactly to text_w.
        drift = text_w - sum(scaled)
        if scaled:
            scaled[-1] += drift
        _set_tbl_layout_fixed(table)
        _set_tbl_width(table, text_w)
        _set_tbl_grid(table, scaled)
        _set_cell_widths(table, scaled)
        applied.append((ti, scaled))
    return applied


def _global_punctuation_fix(text):
    """House punctuation normalization: unicode dashes/quotes/NBSP to ASCII.

    Rules:
      - em-dash (U+2014) and en-dash (U+2013) become an ASCII hyphen
      - a spaced run of 2-3 ASCII hyphens collapses to a single spaced hyphen
      - smart single/double quotes become straight quotes
      - NBSP + narrow NBSP become a regular space
    """
    import re as _re
    t = text
    t = t.replace('\u2014', '-').replace('\u2013', '-')
    t = _re.sub(r' -{2,3} ', ' - ', t)
    t = t.replace('\u2019', "'").replace('\u2018', "'")
    t = t.replace('\u201C', '"').replace('\u201D', '"')
    t = t.replace('\u00A0', ' ').replace('\u202F', ' ')
    return t


def fix_header_footer_text(section, brief=False):
    """Center all header/footer paragraphs and normalize their text.
    Works at paragraph level to handle run-splitting correctly.

    - House punctuation normalization (unicode dash/quote/NBSP to ASCII) via
      _global_punctuation_fix on the paragraph text, gated by the profile's
      normalize_punctuation flag
    - Marking upgrades from classification.header_footer_replacements: for each
      rule, when any of its find-strings is present and its replace-string is not
      already present, the find-string is replaced. Empty list (the default):
      no marking edits. Skipped when brief=True: the brief path applies
      classification.brief_replacements instead (apply_brief_classification_marking).
    """
    for rel_type in ['header', 'footer']:
        for attr_name in [f'{rel_type}', f'first_page_{rel_type}', f'even_page_{rel_type}']:
            try:
                hf = getattr(section, attr_name, None)
                if hf is not None:
                    for p in hf.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        full_text = p.text
                        new_text = _global_punctuation_fix(full_text) if NORMALIZE_PUNCTUATION else full_text
                        if not brief:
                            for rule in CLASSIFICATION['header_footer_replacements']:
                                repl = (rule or {}).get('replace') or ''
                                if not repl:
                                    continue
                                for find in (rule.get('find') or []):
                                    if find and find in new_text and repl not in new_text:
                                        new_text = new_text.replace(find, repl)
                        if new_text != full_text and p.runs:
                            first_run = p.runs[0]
                            first_run.text = new_text
                            for run in p.runs[1:]:
                                run.text = ''
            except Exception:
                pass


def apply_brief_classification_marking(doc):
    """Apply the brief/external marking replacements in headers/footers, per-run
    so the PAGE field is preserved. Driven by classification.brief_replacements
    (a list of {find: [str, ...], replace: str}); empty list = no-op.

    After a replacement, any leftover text sitting in the runs BETWEEN the marking
    run and the following Page-field run is blanked when it consists solely of
    optional whitespace around a parenthesized fragment of one of the configured
    find-strings (fragments are collected from ALL find-strings across all rules).
    This generalizes the original single-literal between-run cleanup."""
    import re as _re
    rules = [r for r in CLASSIFICATION['brief_replacements'] if r and r.get('replace')]
    if not rules:
        return
    fragments = set()
    for rule in rules:
        for find in (rule.get('find') or []):
            for m in _re.finditer(r'\([^)]*\)', find or ''):
                fragments.add(m.group(0))

    def fixpart(part):
        for para in part.paragraphs:
            for run in para.runs:
                replaced = False
                for rule in rules:
                    for find in (rule.get('find') or []):
                        if find and find in run.text:
                            run.text = run.text.replace(find, rule['replace'])
                            replaced = True
                            break
                    if replaced:
                        break
            runs = para.runs
            for rule in rules:
                mark = rule['replace']
                mi = next((i for i, r in enumerate(runs) if mark in r.text), None)
                if mi is not None:
                    stop = next((i for i in range(mi + 1, len(runs)) if 'Page' in runs[i].text), len(runs))
                    between = ''.join(runs[i].text for i in range(mi + 1, stop))
                    core = (between or '').strip()
                    if core and core in fragments:
                        for i in range(mi + 1, stop):
                            runs[i].text = ''

    for s in doc.sections:
        for nm in ('header', 'first_page_header', 'even_page_header',
                   'footer', 'first_page_footer', 'even_page_footer'):
            hf = getattr(s, nm, None)
            if hf is not None:
                try:
                    fixpart(hf)
                except Exception:
                    pass


def apply_global_punctuation_to_body(doc):
    """Apply the house punctuation rules to every paragraph text in the body.
    Sweeps paragraphs + tables + nested tables.

    2026-07-03 fix: per-RUN replacement, never collapse-into-run[0]. The old
    collapse pattern (borrowed from fix_header_footer_text) destroyed inline
    formatting: any paragraph pandoc gave a smart quote or dash was rewritten
    into its first run, so a paragraph opening with a bold phrase rendered
    fully bold. All rules are context-free character substitutions, so per-run
    application is equivalent except for a spaced 2-3 hyphen sequence split
    across two adjacent runs; that rare boundary case is handled by merging
    just that run pair (blast radius: one run's formatting, not the whole
    paragraph's)."""
    import re as _re

    def _fix_para(p):
        runs = p.runs
        for r in runs:
            new = _global_punctuation_fix(r.text)
            if new != r.text:
                r.text = new
        # Boundary case: after the per-run pass, only the spaced-hyphens rule
        # can still match ACROSS a run border (every other rule is a single
        # character substitution). Merge exactly that adjacent pair and re-fix.
        for i in range(len(runs) - 1):
            joined = runs[i].text + runs[i + 1].text
            if _re.search(r' -{2,3} ', joined):
                runs[i].text = _global_punctuation_fix(joined)
                runs[i + 1].text = ''

    for p in doc.paragraphs:
        _fix_para(p)

    def _walk_tables(tbl_iter):
        for tbl in tbl_iter:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _fix_para(p)
                    _walk_tables(cell.tables)
    _walk_tables(doc.tables)


def set_h3_out_of_toc(doc):
    """Remove H3+ headings from TOC by setting outlineLvl to 9 (body text).
    This keeps them as visual headings but excludes from TOC."""
    for p in doc.paragraphs:
        style_name = p.style.name if p.style and p.style.name else ''
        if style_name in ('Heading 3', 'Heading 4', 'Heading 5', 'Heading 6'):
            pPr = p._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(p._element, qn('w:pPr'))
            outline = pPr.find(qn('w:outlineLvl'))
            if outline is None:
                outline = etree.SubElement(pPr, qn('w:outlineLvl'))
            outline.set(qn('w:val'), '9')  # 9 = body text level, excluded from TOC


def _load_table_widths(path):
    """Load a table-widths YAML config -> list of column-width lists (twips)."""
    import yaml
    with open(path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f) or {}
    return data.get('tables', [])


def _set_para_shading(para, fill_hex):
    """Solid background fill on a paragraph (w:pPr/w:shd)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(para._element, qn('w:pPr'))
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)


def _set_para_left_bar(para, color_hex, sz=18, space=10):
    """A thick coloured LEFT border (accent bar) on a paragraph (w:pPr/w:pBdr/w:left)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(para._element, qn('w:pPr'))
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    left = pBdr.find(qn('w:left'))
    if left is None:
        left = etree.SubElement(pBdr, qn('w:left'))
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))       # eighths of a point
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), color_hex)


def style_callouts(doc, prof):
    """Render paragraphs styled 'Callout' (from a callouts lua filter) as a real
    info box: a light fill + a coloured left accent bar + smaller font, with a
    leading glyph on the first line. Warning-class callouts get an amber box + a
    warning glyph. Runs after the global punctuation sweep so the glyph survives."""
    size = max(8.0, prof['body'] - 1.0)
    prev = False
    for p in doc.paragraphs:
        name = p.style.name if p.style and p.style.name else ''
        if name != 'Callout':
            prev = False
            continue
        head = p.text.lstrip().lower()
        warn = head.startswith(('warning', 'important', 'caution', 'failure'))
        fill   = CALLOUT_WARN_FILL   if warn else CALLOUT_INFO_FILL
        border = CALLOUT_WARN_BORDER if warn else CALLOUT_INFO_BORDER
        label  = CALLOUT_WARN_LABEL  if warn else CALLOUT_INFO_LABEL
        _set_para_shading(p, fill)
        _set_para_left_bar(p, border)
        set_para_font(p, size, _rgb(THEME['body']))
        set_para_spacing(p, line=1.12, space_before=(6 if not prev else 0), space_after=2)
        if not prev and p.runs:
            # warning-sign / information-source glyphs (unicode escapes keep this file ASCII)
            glyph = '\u26A0' if warn else '\u2139'
            p.runs[0].text = glyph + '  ' + p.runs[0].text
            for r in p.runs:                          # colour + bold the label line
                r.font.color.rgb = _rgb(label)
                r.font.bold = True
        prev = True


def main(argv=None):
    """Entry point. `argv` defaults to sys.argv[1:] when run as a script (also how
    render-doc.sh invokes this module: as a subprocess, unchanged); an explicit list
    lets a caller (e.g. `render docstyle`, see render.py) invoke it in-process with
    its own argument list, without touching sys.argv."""
    # Optional flag: --table-widths <yaml> applies operator-fitted, full-width,
    # fixed-layout column widths (generic; any document can pass its own config).
    args = sys.argv[1:] if argv is None else list(argv)
    if '-h' in args or '--help' in args:
        print("Usage: python style_postprocess.py <input.docx> [output.docx]\n"
              "    [--profile compact|reference] [--template-profile <yaml>]\n"
              "    [--table-widths <yaml>] [--cover-version <v>] [--cover-date <d>]\n\n"
              "Applies the house style (font, headings, tables, page geometry, header/\n"
              "footer handling, punctuation normalization) to a pandoc-rendered DOCX, in\n"
              "place unless [output.docx] is given.\n\n"
              "  --profile compact|reference   style profile (default: compact)\n"
              "  --template-profile <yaml>     override theme/marking/cover from a profile yaml\n"
              "  --table-widths <yaml>         operator-fitted column widths (see apply_table_widths)\n"
              "  --cover-version <v>           cover version-line value (reference profile)\n"
              "  --cover-date <d>              cover date-line value (reference profile)")
        return 0
    table_widths_path = None
    profile_name = 'compact'
    cover_version = None
    cover_date = None
    positional = []
    i = 0
    while i < len(args):
        if args[i] == '--template-profile':
            apply_template_profile(args[i + 1]); i += 2; continue
        if args[i] == '--table-widths':
            table_widths_path = args[i + 1]; i += 2; continue
        if args[i] == '--profile':
            profile_name = args[i + 1]; i += 2; continue
        if args[i] == '--cover-version':
            cover_version = args[i + 1]; i += 2; continue
        if args[i] == '--cover-date':
            cover_date = args[i + 1]; i += 2; continue
        positional.append(args[i]); i += 1
    prof = PROFILES.get(profile_name, PROFILES['compact'])

    if len(positional) < 1:
        print("Usage: python style_postprocess.py <input.docx> [output.docx] "
              "[--table-widths <config.yaml>]")
        sys.exit(1)

    src = os.path.abspath(positional[0])
    dst = os.path.abspath(positional[1]) if len(positional) > 1 else src

    doc = Document(src)

    # --- Page setup (theme-driven; default A4, 2cm margins) ---
    for section in doc.sections:
        section.page_width = Cm(THEME['page_width_cm'])
        section.page_height = Cm(THEME['page_height_cm'])
        section.top_margin = Cm(THEME['margin_cm'])
        section.bottom_margin = Cm(THEME['margin_cm'])
        section.left_margin = Cm(THEME['margin_cm'])
        section.right_margin = Cm(THEME['margin_cm'])

    # --- Fix and center headers and footers ---
    _brief = (profile_name == 'reference')
    for section in doc.sections:
        fix_header_footer_text(section, brief=_brief)
    if _brief:
        apply_brief_classification_marking(doc)

    # --- Global punctuation fix (house rule, profile-gated) ---
    # Sweeps body paragraphs + table cells + nested table cells. Replaces em-dash /
    # en-dash / smart-quotes / NBSP with ASCII equivalents.
    if NORMALIZE_PUNCTUATION:
        apply_global_punctuation_to_body(doc)

    # --- Remove H3+ from TOC ---
    set_h3_out_of_toc(doc)

    # --- Style paragraphs (profile-driven) ---
    is_ref = (profile_name == 'reference')
    h1_count = 0
    for p in doc.paragraphs:
        style_name = p.style.name if p.style and p.style.name else ''
        text = p.text.strip()

        if style_name == 'Title':
            set_para_font(p, prof['title'], NAVY, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          font_name=_style_font(style_name))
            if is_ref:
                set_para_spacing(p, space_before=72, space_after=8)

        elif style_name == 'Subtitle':
            set_para_font(p, prof['sub'], GREY_BODY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          font_name=_style_font(style_name))
            if is_ref:
                set_para_spacing(p, space_after=18)

        elif style_name == 'Heading 1':
            set_para_font(p, prof['h1'], NAVY, bold=True, font_name=_style_font(style_name))
            h1_count += 1
            # Page break before every H1 except the very first
            if h1_count > 1:
                add_page_break_before(p)
            if is_ref:
                set_para_spacing(p, space_before=2, space_after=6)

        elif style_name == 'Heading 2':
            set_para_font(p, prof['h2'], NAVY, bold=True, font_name=_style_font(style_name))
            if is_ref:
                set_para_spacing(p, space_before=12, space_after=4)

        elif style_name == 'Heading 3':
            set_para_font(p, prof['h3'], NAVY, bold=True, font_name=_style_font(style_name))
            if is_ref:
                set_para_spacing(p, space_before=8, space_after=3)

        elif style_name == 'Heading 4':
            set_para_font(p, prof['h4'], NAVY, bold=True, font_name=_style_font(style_name))

        elif text and not style_name.startswith('Heading'):
            set_para_font(p, prof['body'], prof['body_color'], font_name=_style_font(style_name))
            if prof['justify'] and len(text) > 80:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif not prof['justify']:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, line=prof['line'], space_after=prof['space_after'])

    # --- Style callout boxes (info / warning) ---
    style_callouts(doc, prof)

    # --- Style ALL tables uniformly (profile-driven) ---
    for table in doc.tables:
        style_table(table, first_col_bold=True,
                    hdr_size=prof['table_hdr'], body_size=prof['table_body'], zebra=prof['zebra'])

    # --- Reference-profile extras: centred figures + cover page ---
    if is_ref:
        center_figures_and_captions(doc)
        build_reference_cover(doc, version=cover_version, date_str=cover_date)

    # --- Apply operator-fitted column widths (optional, parametrisable) ---
    width_report = []
    if table_widths_path:
        specs = _load_table_widths(table_widths_path)
        width_report = apply_table_widths(doc, specs)

    doc.save(dst)

    # Report
    headings = [p for p in doc.paragraphs if p.style and p.style.name and 'Heading' in p.style.name]
    print(f"Styled: {dst}")
    print(f"  {len(doc.paragraphs)} paragraphs, {len(headings)} headings, {len(doc.tables)} tables")
    if table_widths_path:
        print(f"  Table widths applied (fixed layout, full text-width) from "
              f"{os.path.basename(table_widths_path)}:")
        for ti, w in width_report:
            print(f"    table {ti}: {w}  (sum={sum(w)})")

    try:
        size = os.path.getsize(dst)
        print(f"  File size: {size/1024:.0f} KB")
    except OSError:
        print("  (file size check skipped: file not stat-able yet)")

    return 0


if __name__ == '__main__':
    sys.exit(main())

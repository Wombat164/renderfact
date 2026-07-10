"""
Tests for lint/comprehension_review_contract.py (issue #84): the fresh-reader
comprehension gate for rendered TEXT documents, the text peer of
lint/vision_review_contract.py's diagram vision-review gate.

Covers: fence-aware ATX heading chunking (leaf sections, code-fence and
div-fence awareness, paragraph sub-splitting, frontmatter stripping); DOCX
extraction (Heading-N / localized Kop-N styles) feeding the SAME chunker;
input assembly + schema validation; the D16 gate's deliberate "always
escalate" confidence (confidence() pinned at 0.0, the threshold<=0 escape
hatch); the deterministic unreviewed stub; the D8 contract (the generic
copy_paste driver drives this step from an injected paste, forcing
reviewer_mode); the CLI end to end for the needs-review path (no --escalate),
the threshold<=0 accept path, and the copy-paste escalate path; the
render.py copy-paste redirect (HAS_OWN_GATE); and init-ai exposing the step
to a harness.
"""

from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(REPO_ROOT / "lint"))

import comprehension_review_contract as cr  # noqa: E402

RENDER_PY = REPO_ROOT / "render.py"


# ------------------------------------------------------------------ chunking --

def test_chunk_by_heading_boundary():
    text = "# Title\n\n## One\n\nFirst body.\n\n## Two\n\nSecond body.\n"
    chunks = cr.chunk_document(text, target_words=1000)
    headings = [c["heading"] for c in chunks]
    assert headings == ["Title", "One", "Two"]
    assert chunks[1]["text"] == "First body."
    assert chunks[2]["text"] == "Second body."


def test_preamble_before_first_heading():
    text = "Some intro text.\n\n# First Heading\n\nBody.\n"
    chunks = cr.chunk_document(text)
    assert chunks[0]["heading"] == "(preamble)"
    assert chunks[0]["text"] == "Some intro text."
    assert chunks[1]["heading"] == "First Heading"


def test_blank_preamble_is_dropped():
    text = "\n\n# First Heading\n\nBody.\n"
    chunks = cr.chunk_document(text)
    assert [c["heading"] for c in chunks] == ["First Heading"]


def test_no_headings_at_all_is_one_leaf_section():
    text = "Just a single paragraph with no headings anywhere in this document."
    chunks = cr.chunk_document(text, target_words=1000)
    assert len(chunks) == 1
    assert chunks[0]["heading"] == "(preamble)"


def test_code_fence_hash_is_not_a_heading():
    text = (
        "## Real Section\n\n"
        "```python\n"
        "# not a heading, just a comment\n"
        "## also not a heading\n"
        "```\n\n"
        "Trailing text.\n"
    )
    chunks = cr.chunk_document(text, target_words=1000)
    headings = [c["heading"] for c in chunks]
    assert headings == ["Real Section"]
    assert "not a heading" in chunks[0]["text"]
    assert "```" in chunks[0]["text"]


def test_div_fence_hash_is_not_a_heading():
    text = (
        "## Real Section\n\n"
        ':::{clearance="internal"}\n'
        "# not a heading either\n"
        ":::\n\n"
        "Trailing text.\n"
    )
    chunks = cr.chunk_document(text, target_words=1000)
    assert [c["heading"] for c in chunks] == ["Real Section"]
    assert "not a heading either" in chunks[0]["text"]


def test_paragraph_sub_split_respects_target_words_and_never_splits_mid_paragraph():
    body_paras = [f"Paragraph {i} has exactly six words here." for i in range(6)]
    text = "## Long Section\n\n" + "\n\n".join(body_paras) + "\n"
    chunks = cr.chunk_document(text, target_words=15)
    long_section_chunks = [c for c in chunks if c["heading"].startswith("Long Section")]
    assert len(long_section_chunks) > 1
    # every sub-chunk after the first is labeled "(cont.)"
    assert long_section_chunks[0]["heading"] == "Long Section"
    assert all(c["heading"] == "Long Section (cont.)" for c in long_section_chunks[1:])
    # no paragraph was split mid-sentence: every original paragraph string
    # appears whole in exactly one chunk
    joined = "\n\n".join(c["text"] for c in long_section_chunks)
    for p in body_paras:
        assert p in joined


def test_oversized_single_paragraph_stays_whole():
    """A paragraph over budget on its own is never sliced mid-sentence -- the
    deterministic `render qa paras` check already flags overweight paragraphs;
    this gate is not the place to split one open."""
    long_para = " ".join(f"word{i}" for i in range(80))
    text = f"## Section\n\n{long_para}\n"
    chunks = cr.chunk_document(text, target_words=20)
    matching = [c for c in chunks if c["heading"] == "Section"]
    assert len(matching) == 1
    assert matching[0]["text"] == long_para


def test_frontmatter_is_stripped_before_chunking():
    text = "---\ntitle: X\ntags: [a, b]\n---\n\n## Section\n\nBody text.\n"
    chunks = cr.chunk_document(text)
    joined = " ".join(c["text"] for c in chunks) + " " + " ".join(c["heading"] for c in chunks)
    assert "title: X" not in joined
    assert [c["heading"] for c in chunks] == ["Section"]


def test_empty_document_produces_one_placeholder_chunk():
    chunks = cr.chunk_document("")
    assert len(chunks) == 1
    assert chunks[0]["text"] == ""


def test_indices_are_sequential_in_document_order():
    text = "# A\n\nx\n\n# B\n\ny\n\n# C\n\nz\n"
    chunks = cr.chunk_document(text, target_words=1000)
    assert [c["index"] for c in chunks] == list(range(len(chunks)))


# --------------------------------------------------------- docx extraction --

def test_extract_docx_converts_heading_styles_to_atx(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_heading("Report Title", level=1)
    doc.add_heading("Overview", level=2)
    doc.add_paragraph("Body text under overview.")
    p = tmp_path / "r.docx"
    doc.save(str(p))

    text = cr.extract_text(p)
    assert "# Report Title" in text
    assert "## Overview" in text
    chunks = cr.chunk_document(text)
    assert "Overview" in [c["heading"] for c in chunks]
    assert any("Body text under overview." in c["text"] for c in chunks)


def test_extract_text_reads_markdown_directly(tmp_path):
    p = tmp_path / "d.md"
    p.write_text("# Title\n\nBody.\n", encoding="utf-8")
    assert cr.extract_text(p) == "# Title\n\nBody.\n"


# --------------------------------------------------------- assemble_input --

def test_assemble_input_shape_and_schema():
    chunks = cr.chunk_document("## One\n\nBody.\n", target_words=1000)
    obj = cr.assemble_input(chunks, "My Doc")
    assert obj["task_intent"] == cr.TASK_INTENT
    assert obj["doc_title"] == "My Doc"
    assert obj["chunks"] == chunks


# ------------------------------------------------------ confidence + gate --

def test_confidence_is_always_zero_regardless_of_document_shape():
    """The load-bearing D16 decision for this step (docs/DECISIONS.md D20):
    there is no deterministic sufficiency proxy for comprehension, so
    confidence() is pinned at 0.0 no matter how short, long, or well-structured
    the input is."""
    tiny = cr.assemble_input(cr.chunk_document("## A\n\nx\n", target_words=1000), "tiny")
    big_text = "\n\n".join(f"## Section {i}\n\nSome body text here." for i in range(40))
    big = cr.assemble_input(cr.chunk_document(big_text, target_words=1000), "big")

    for obj in (tiny, big):
        conf = cr.confidence(obj)
        assert conf.score == 0.0
        assert set(conf.signals) == {"chunk_count", "word_count"}


def test_gate_always_escalates_at_any_positive_threshold():
    obj = cr.assemble_input(cr.chunk_document("## A\n\nx\n", target_words=1000), "d")
    assert cr.gate(obj, threshold=0.01)[0] == "escalate"
    assert cr.gate(obj, threshold=0.6)[0] == "escalate"
    assert cr.gate(obj, threshold=1.0)[0] == "escalate"


def test_gate_threshold_zero_or_below_accepts_the_stub():
    obj = cr.assemble_input(cr.chunk_document("## A\n\nx\n", target_words=1000), "d")
    assert cr.gate(obj, threshold=0.0)[0] == "accept"
    assert cr.gate(obj, threshold=-1.0)[0] == "accept"


# ------------------------------------------------- deterministic entries --

def test_deterministic_entry_is_schema_valid_and_honest():
    chunks = cr.chunk_document("## A\n\nx\n\n## B\n\ny\n", target_words=1000)
    obj = cr.assemble_input(chunks, "d")
    entry = cr.deterministic_entry(obj)
    ok, errors = cr.validate_output(entry)
    assert ok, errors
    assert entry["status"] == "WARN"
    assert entry["chunk_findings"] == []
    assert entry["reviewer_mode"] == "deterministic"
    assert "not" in entry["doc_purpose"]  # "not analyzed"
    assert str(len(chunks)) in entry["summary"]


# --------------------------------------------------------------- D8 contract --

def test_copy_paste_driver_drives_comprehension_review():
    """The generic D8 driver must run comprehension-review from a human paste
    and force reviewer_mode -- proving it is a valid, mode-uniform step
    contract, same as vision-review / decision-capture / contextualize."""
    from contracts import copy_paste

    chunks = cr.chunk_document("## Intro\n\nHello reader.\n", target_words=1000)
    obj = cr.assemble_input(chunks, "d")
    pasted = json.dumps({
        "status": "OK",
        "chunk_findings": [
            {"index": 0, "purpose": "Greet the reader", "confusing": "", "fluff": "", "cuttable": ""},
        ],
        "doc_purpose": "A greeting document",
        "worst_snippet": "none",
        "cut_first": "nothing",
        "summary": "Reads cleanly.",
        "reviewer_mode": "harness",  # deliberately wrong: the driver must overwrite it
    })
    lines = iter(pasted.splitlines() + ["END"])
    result = copy_paste.run_copy_paste_step(
        "comprehension-review", cr, obj, lines_source=lines,
        out=open(__import__("os").devnull, "w"))
    assert result["reviewer_mode"] == "copy-paste"  # forced, not trusted
    assert result["doc_purpose"] == "A greeting document"
    ok, errors = cr.validate_output(result)
    assert ok, errors


def test_mode_field_declared():
    assert cr.MODE_FIELD == "reviewer_mode"


# --------------------------------------------------------------------- CLI --

def _write_doc(tmp_path) -> Path:
    p = tmp_path / "doc.md"
    p.write_text(
        "# Sample Report\n\n"
        "## Overview\n\n"
        "This report describes a system for a new reader.\n\n"
        "## Details\n\n"
        "Padding filler text that could be considered fluff for this section.\n",
        encoding="utf-8",
    )
    return p


def test_cli_no_escalate_emits_needs_review_stub(tmp_path):
    doc = _write_doc(tmp_path)
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "comprehension-review", str(doc), "--json"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(REPO_ROOT))
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["decision"] == "escalate"
    assert payload["confidence"] == 0.0
    assert payload["needs_review"] is True
    assert payload["entry"]["reviewer_mode"] == "deterministic"
    assert payload["chunk_count"] >= 3


def test_cli_threshold_zero_accepts_stub_without_prompting(tmp_path):
    doc = _write_doc(tmp_path)
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "comprehension-review", str(doc),
         "--threshold", "0", "--json"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(REPO_ROOT), input="")
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["decision"] == "accept"
    assert payload["needs_review"] is False
    assert "Paste the LLM" not in r.stdout and "Paste the LLM" not in r.stderr


def test_cli_escalate_copy_paste_runs_the_llm(tmp_path):
    doc = _write_doc(tmp_path)
    pasted = json.dumps({
        "status": "WARN",
        "chunk_findings": [
            {"index": 0, "purpose": "Title", "confusing": "", "fluff": "", "cuttable": ""},
            {"index": 1, "purpose": "Overview", "confusing": "", "fluff": "", "cuttable": ""},
            {"index": 2, "purpose": "Details", "confusing": "no transition", "fluff": "padding",
             "cuttable": "the filler text"},
        ],
        "doc_purpose": "Describe a system",
        "worst_snippet": "Details",
        "cut_first": "the filler text in Details",
        "summary": "Mostly clear; Details section pads.",
        "reviewer_mode": "harness",
    })
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "comprehension-review", str(doc), "--escalate", "copy-paste"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(REPO_ROOT), input=pasted + "\nEND\n")
    assert r.returncode == 0, r.stderr
    assert "-> escalate" in r.stderr
    assert "mode: copy-paste" in r.stdout
    assert "Details" in r.stdout
    assert "the filler text in Details" in r.stdout


def test_cli_missing_document_is_a_clean_error(tmp_path):
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "comprehension-review", str(tmp_path / "nope.md")],
        capture_output=True, text=True, encoding="utf-8", cwd=str(REPO_ROOT))
    assert r.returncode == 1
    assert "not found" in r.stderr


def test_render_copy_paste_redirects_comprehension_review():
    """The vision-shaped `render copy-paste` CLI must not mis-drive
    comprehension-review; it points the user at the step's own gated command
    (via the declared HAS_OWN_GATE flag)."""
    assert cr.HAS_OWN_GATE is True
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "copy-paste", "comprehension-review"],
        capture_output=True, text=True, encoding="utf-8", cwd=str(REPO_ROOT), input="")
    assert r.returncode == 2
    assert "comprehension-review --escalate copy-paste" in r.stderr


# --------------------------------------------------------------- harness --

def test_init_ai_exposes_comprehension_review():
    from contracts import init_ai

    names = [n for n, _ in init_ai.step_contracts()]
    assert "comprehension-review" in names
    skill = init_ai.render_claude_skill()
    assert "comprehension-review" in skill
    assert cr.TASK_INTENT[:40] in skill

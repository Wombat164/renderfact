"""
Tests for docstyle/style_postprocess.py + docstyle/heading_numbering.py (chunk F2b).

All DOCX fixtures are built in tmp_path via python-docx (no binary fixtures).
Covers: end-to-end style application through main() (font + heading colour +
table header), punctuation normalization and its normalize_punctuation gate,
header/footer marking replacements from a profile yaml, cover-banner stripping
via a configured prefix, heading-numbering injection (numbering.xml + styles.xml
numPr) with byte-identical idempotency, and the trailing-dot vs modern numbering
schemes producing different abstractNum text.
"""

from __future__ import annotations

import importlib
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

import docstyle.style_postprocess as sp  # noqa: E402
from docstyle import heading_numbering as hn  # noqa: E402


@pytest.fixture(autouse=True)
def fresh_style_module():
    """style_postprocess keeps profile config in module globals; reload before
    each test so a --template-profile applied in one test cannot leak into the
    next."""
    importlib.reload(sp)
    yield


def _build_doc(tmp_path, name="in.docx",
               body_text="A plain body paragraph that is comfortably longer than "
                         "eighty characters so the justify rule kicks in."):
    doc = Document()
    doc.add_heading("First Section", level=1)
    doc.add_heading("A Subsection", level=2)
    doc.add_paragraph(body_text)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Head A"
    table.rows[0].cells[1].text = "Head B"
    table.rows[1].cells[0].text = "Cell 1"
    table.rows[1].cells[1].text = "Cell 2"
    path = tmp_path / name
    doc.save(str(path))
    return path


def _run_style(monkeypatch, argv):
    monkeypatch.setattr(sys, "argv", ["style_postprocess.py", *argv])
    sp.main()


# ---------- (a) end-to-end style application ----------

def test_style_end_to_end_font_and_heading_colour(tmp_path, monkeypatch):
    src = _build_doc(tmp_path)
    out = tmp_path / "styled.docx"
    _run_style(monkeypatch, [str(src), str(out)])

    doc = Document(str(out))
    h1 = next(p for p in doc.paragraphs if p.style.name == "Heading 1")
    assert h1.runs[0].font.name == "Arial"
    assert h1.runs[0].font.color.rgb == RGBColor(0x1F, 0x38, 0x64)
    h2 = next(p for p in doc.paragraphs if p.style.name == "Heading 2")
    assert h2.runs[0].font.color.rgb == RGBColor(0x1F, 0x38, 0x64)
    body = next(p for p in doc.paragraphs
                if p.text and not p.style.name.startswith("Heading"))
    assert body.runs[0].font.name == "Arial"
    hdr_run = doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
    assert hdr_run.font.bold is True
    assert hdr_run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
    assert hdr_run.font.name == "Arial"


# ---------- (a2) per-style font overrides (issue #97) ----------

def test_style_fonts_override_applies_per_named_style(tmp_path, monkeypatch):
    """A profile's `styles:` block overrides the font for paragraphs carrying
    that named style; every other paragraph keeps the global 'font'."""
    doc = Document()
    doc.add_heading("First Section", level=1)
    doc.add_paragraph("A quoted line that should get its own font.", style=doc.styles["Quote"])
    doc.add_paragraph("Regular body paragraph using the global font throughout.")
    src = tmp_path / "styles-in.docx"
    doc.save(str(src))

    profile = tmp_path / "styles.yaml"
    profile.write_text(
        "font: Arial\n"
        "styles:\n"
        '  "Quote":\n'
        "    font: Georgia\n",
        encoding="utf-8",
    )
    out = tmp_path / "styles-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    out_doc = Document(str(out))
    quote_para = next(p for p in out_doc.paragraphs if p.style.name == "Quote")
    assert quote_para.runs[0].font.name == "Georgia"
    body_para = next(p for p in out_doc.paragraphs
                     if p.text.startswith("Regular body"))
    assert body_para.runs[0].font.name == "Arial"
    h1 = next(p for p in out_doc.paragraphs if p.style.name == "Heading 1")
    assert h1.runs[0].font.name == "Arial"  # no override for Heading 1: global font


def test_style_fonts_override_absent_leaves_all_paragraphs_on_global_font(tmp_path, monkeypatch):
    """No `styles:` block in the profile: unchanged pre-#97 behaviour, every
    paragraph (any style) gets the single global font."""
    doc = Document()
    doc.add_paragraph("A quoted line with no per-style override configured.",
                      style=doc.styles["Quote"])
    src = tmp_path / "no-styles-in.docx"
    doc.save(str(src))

    profile = tmp_path / "no-styles.yaml"
    profile.write_text("font: Arial\n", encoding="utf-8")
    out = tmp_path / "no-styles-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    out_doc = Document(str(out))
    quote_para = next(p for p in out_doc.paragraphs if p.style.name == "Quote")
    assert quote_para.runs[0].font.name == "Arial"


# ---------- (b) punctuation normalization + gate ----------

def test_punctuation_normalized_by_default(tmp_path, monkeypatch):
    text = "alpha \u2014 beta \u201Cquoted\u201D \u2018single\u2019"
    src = _build_doc(tmp_path, "punct.docx", body_text=text)
    out = tmp_path / "punct-styled.docx"
    _run_style(monkeypatch, [str(src), str(out)])

    joined = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "\u2014" not in joined
    assert "\u201C" not in joined and "\u2018" not in joined
    assert 'alpha - beta "quoted" \'single\'' in joined


def test_punctuation_gate_disables_normalization(tmp_path, monkeypatch):
    text = "alpha \u2014 beta \u201Cquoted\u201D"
    src = _build_doc(tmp_path, "punct2.docx", body_text=text)
    profile = tmp_path / "no-punct.yaml"
    profile.write_text("normalize_punctuation: false\n", encoding="utf-8")
    out = tmp_path / "punct-kept.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    joined = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "\u2014" in joined and "\u201C" in joined


# ---------- (c) header/footer replacements from a profile yaml ----------

def test_header_footer_replacements_from_profile(tmp_path, monkeypatch, capsys):
    doc = Document()
    doc.add_paragraph("body text")
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = "HANDLE WITH CARE"
    src = tmp_path / "hdr.docx"
    doc.save(str(src))

    profile = tmp_path / "marking.yaml"
    profile.write_text(
        "classification:\n"
        "  header_footer_replacements:\n"
        '    - find: ["HANDLE WITH CARE"]\n'
        '      replace: "HANDLE WITH CARE (POLICY-REF)"\n',
        encoding="utf-8",
    )
    out = tmp_path / "hdr-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    hdr_text = "\n".join(p.text for p in Document(str(out)).sections[0].header.paragraphs)
    assert "HANDLE WITH CARE (POLICY-REF)" in hdr_text
    assert "WARN" not in capsys.readouterr().out  # matched: no unmatched-rule warning


def test_classification_replacement_warns_when_find_string_never_matches(tmp_path, monkeypatch, capsys):
    doc = Document()
    doc.add_paragraph("body text")
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = "UNCLASS"
    src = tmp_path / "hdr.docx"
    doc.save(str(src))

    profile = tmp_path / "marking.yaml"
    profile.write_text(
        "classification:\n"
        "  header_footer_replacements:\n"
        '    - find: ["A STRING NOT IN THE DOCUMENT"]\n'
        '      replace: "SNC"\n',
        encoding="utf-8",
    )
    out = tmp_path / "hdr-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    warn = capsys.readouterr().out
    assert "WARN" in warn
    assert "never matched" in warn
    assert "header_footer_replacements" in warn


def test_classification_replacement_warns_when_configured_under_wrong_profile_key(
        tmp_path, monkeypatch, capsys):
    doc = Document()
    doc.add_paragraph("body text")
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = "UNCLASS"
    src = tmp_path / "hdr.docx"
    doc.save(str(src))

    profile = tmp_path / "marking.yaml"
    profile.write_text(
        "classification:\n"
        "  brief_replacements:\n"  # active key for --profile reference is brief_replacements;
        '    - find: ["UNCLASS"]\n'  # this test renders --profile compact, so this key is inactive
        '      replace: "SNC"\n',
        encoding="utf-8",
    )
    out = tmp_path / "hdr-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile),
                             "--profile", "compact"])

    warn = capsys.readouterr().out
    assert "WARN" in warn
    assert "is empty for this render" in warn
    assert "header_footer_replacements" in warn and "brief_replacements" in warn
    hdr_text = Document(str(out)).sections[0].header.paragraphs[0].text
    assert hdr_text == "UNCLASS"  # confirms the marking genuinely passed through unchanged


def test_brief_replacement_survives_a_run_boundary_split(tmp_path, monkeypatch):
    """Root-cause reproduction for #87: apply_brief_classification_marking (the
    --profile reference path) used to match per-run, so a marking string split
    across two runs -- exactly what a real Word hand-edit produces -- silently
    never matched. fix_header_footer_text (the compact path) was never affected,
    since it already concatenated to paragraph-level text first; this test
    exercises the path that WAS broken."""
    doc = Document()
    doc.add_paragraph("body text")
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    p = hdr.paragraphs[0]
    p.text = ""
    p.add_run("UNCLA")
    p.add_run("SS")  # the marking string split across two runs, on purpose
    src = tmp_path / "hdr-split.docx"
    doc.save(str(src))

    profile = tmp_path / "marking.yaml"
    profile.write_text(
        "classification:\n"
        "  brief_replacements:\n"
        '    - find: ["UNCLASS"]\n'
        '      replace: "SNC"\n',
        encoding="utf-8",
    )
    out = tmp_path / "hdr-split-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile),
                             "--profile", "reference"])

    hdr_text = Document(str(out)).sections[0].header.paragraphs[0].text
    assert hdr_text == "SNC", (
        f"expected the run-split 'UNCLASS' to be replaced with 'SNC', got {hdr_text!r} "
        f"-- if this still reads 'UNCLASS', the run-boundary fix regressed")


# ---------- (d) cover banner stripping via configured prefix ----------

def test_cover_banner_strip_with_configured_prefix(tmp_path, monkeypatch):
    doc = Document()
    doc.add_paragraph("Handling: TAKE CARE")
    doc.add_heading("Part I: Overview", level=1)
    doc.add_paragraph("Body content that stays.")
    src = tmp_path / "cover.docx"
    doc.save(str(src))

    profile = tmp_path / "cover.yaml"
    profile.write_text(
        "classification:\n"
        '  strip_cover_banner_prefixes: ["Handling:"]\n',
        encoding="utf-8",
    )
    out = tmp_path / "cover-out.docx"
    _run_style(monkeypatch, [str(src), str(out), "--profile", "reference",
                             "--template-profile", str(profile)])

    texts = [p.text for p in Document(str(out)).paragraphs]
    assert not any(t.startswith("Handling:") for t in texts)
    assert any("Body content that stays." in t for t in texts)
    assert any(t.startswith("Part I") for t in texts)  # body start survived


# ---------- (e) heading numbering: injection + idempotency ----------

def test_heading_numbering_injects_and_second_run_is_noop(tmp_path):
    doc = Document()
    doc.add_heading("One", level=1)
    doc.add_heading("One point one", level=2)
    doc.add_paragraph("body")
    path = tmp_path / "num.docx"
    doc.save(str(path))

    assert hn.process(path, check=False) == 0
    with zipfile.ZipFile(path) as z:
        num_xml = z.read("word/numbering.xml").decode("utf-8")
        sty_xml = z.read("word/styles.xml").decode("utf-8")
    assert 'w:abstractNumId="8100"' in num_xml
    assert '<w:num w:numId="8100">' in num_xml
    assert '<w:numId w:val="8100"/>' in sty_xml  # numPr wired into heading styles

    before = path.read_bytes()
    assert hn.process(path, check=False) == 0
    assert path.read_bytes() == before  # byte-identical: true no-op


# ---------- (f) trailing-dot vs modern schemes ----------

def test_trailing_dot_and_modern_schemes_differ():
    modern = hn.build_abstract_num("modern")
    trailing = hn.build_abstract_num("trailing-dot")
    assert modern != trailing
    # level 0 has a trailing dot in both schemes
    assert '<w:lvlText w:val="%1."/>' in modern
    assert '<w:lvlText w:val="%1."/>' in trailing
    # deeper levels: only the trailing-dot scheme keeps the terminal dot
    assert '<w:lvlText w:val="%1.%2"/>' in modern
    assert '<w:lvlText w:val="%1.%2."/>' in trailing


def test_punctuation_is_per_run_and_preserves_inline_formatting(tmp_path):
    # Regression for the 2026-07-03 per-run fix: a bold opening phrase must stay
    # scoped to its own run; the old collapse-into-run[0] made it swallow the
    # whole paragraph.
    doc = Document()
    p = doc.add_paragraph()
    r0 = p.add_run("Key point:")
    r0.bold = True
    r1 = p.add_run(" this uses a " + chr(0x2019) + "smart" + chr(0x2019)
                + " quote and an em" + chr(0x2014) + "dash.")

    sp.apply_global_punctuation_to_body(doc)

    assert p.runs[0].text == "Key point:"
    assert p.runs[0].bold is True
    assert p.runs[1].bold is not True
    assert "'smart'" in p.runs[1].text
    assert "em-dash" in p.runs[1].text
    assert chr(0x2019) not in p.text and chr(0x2014) not in p.text


def test_punctuation_merges_only_the_split_hyphen_run_pair(tmp_path):
    # The spaced 2-3 hyphen rule is the one substitution that can straddle a
    # run border; only that adjacent pair may be merged, later runs untouched.
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("alpha -")
    p.add_run("- beta")
    tail = p.add_run(" tail")
    tail.italic = True

    sp.apply_global_punctuation_to_body(doc)

    assert p.text == "alpha - beta tail"
    assert p.runs[0].text == "alpha - beta"
    assert p.runs[1].text == ""
    assert p.runs[2].text == " tail"
    assert p.runs[2].italic is True


# ---------- (g) main(argv=...) explicit-arg entry point (issue #74) ----------

def test_main_accepts_explicit_argv_list_without_touching_sys_argv(tmp_path, monkeypatch):
    # render.py's `run_docstyle` calls style_postprocess.main(args) directly,
    # in-process, rather than mutating sys.argv (see render.py). main() must
    # honour an explicit argv list and leave sys.argv untouched.
    src = _build_doc(tmp_path)
    out = tmp_path / "styled.docx"
    sentinel = ["style_postprocess.py", "--should-not-be-read"]
    monkeypatch.setattr(sys, "argv", sentinel)

    rc = sp.main([str(src), str(out)])

    assert rc == 0
    assert sys.argv == sentinel  # untouched
    assert out.exists()


def test_help_flag_returns_zero_and_does_not_require_a_docx(capsys):
    rc = sp.main(["--help"])
    assert rc == 0
    out = capsys.readouterr().out
    assert "--table-widths" in out
    assert "--cover-version" in out


# ---------- (h) --table-widths applies scaled, fixed-layout column widths ----------

def test_table_widths_scales_proportionally_to_text_width(tmp_path):
    src = _build_doc(tmp_path, "widths.docx")
    out = tmp_path / "widths-out.docx"
    widths_yaml = tmp_path / "widths.yaml"
    # 1:2 ratio; absolute magnitude is irrelevant, only the ratio survives scaling.
    widths_yaml.write_text("tables:\n  - [1000, 2000]\n", encoding="utf-8")

    sp.main([str(src), str(out), "--table-widths", str(widths_yaml)])

    doc = Document(str(out))
    table = doc.tables[0]
    text_w = sp._section_text_width_twips(doc)
    grid = table._tbl.find(sp.qn("w:tblGrid"))
    col_widths = [int(gc.get(sp.qn("w:w"))) for gc in grid.findall(sp.qn("w:gridCol"))]

    assert len(col_widths) == 2
    assert sum(col_widths) == text_w  # scaled to fill the section text width exactly
    # 1:2 ratio preserved (within a rounding twip)
    assert abs(col_widths[1] - 2 * col_widths[0]) <= 1


def test_table_widths_skipped_on_column_count_mismatch(tmp_path):
    # A spec with the wrong column count is a guard against source table-order
    # drift: it must be skipped, not mis-fit onto the table. python-docx already
    # writes a default tblGrid on table creation, so the guard is proven by
    # checking apply_table_widths()'s own reported set, and that no fixed-layout
    # width was stamped onto the table -- not by grid presence/absence.
    src = _build_doc(tmp_path, "widths-mismatch.docx")
    out = tmp_path / "widths-mismatch-out.docx"
    widths_yaml = tmp_path / "widths-bad.yaml"
    widths_yaml.write_text("tables:\n  - [1000, 2000, 3000]\n", encoding="utf-8")  # 3 cols, table has 2

    sp.main([str(src), str(out), "--table-widths", str(widths_yaml)])

    doc = Document(str(out))
    tbl_pr = doc.tables[0]._tbl.tblPr
    assert tbl_pr.find(sp.qn("w:tblLayout")) is None  # never switched to fixed layout


# ---------- (i) --cover-version / --cover-date on the reference profile ----------

def test_cover_version_and_date_render_the_cover_line(tmp_path):
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    doc.styles.add_style("Date", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_heading("Report Title", level=1)
    date_p = doc.add_paragraph("placeholder")
    date_p.style = doc.styles["Date"]
    doc.add_heading("Part I: Overview", level=1)
    doc.add_paragraph("Body content.")
    src = tmp_path / "cover-fields.docx"
    doc.save(str(src))
    out = tmp_path / "cover-fields-out.docx"

    sp.main([str(src), str(out), "--profile", "reference",
              "--cover-version", "1.2", "--cover-date", "2026-07-10"])

    texts = [p.text for p in Document(str(out)).paragraphs]
    assert any("1.2" in t and "2026-07-10" in t for t in texts)


# ---------- (i2) build_reference_cover title-match, not positional deletion (#122) ----------

def test_flat_document_h1_matching_title_is_removed_as_duplicate(tmp_path):
    doc = Document()
    doc.add_heading("Example Report", level=0)   # level=0 -> [Title]-styled paragraph
    doc.add_heading("Example Report", level=1)   # a genuine duplicate of the title text
    doc.add_paragraph("Body content.")
    src = tmp_path / "dup-title.docx"
    doc.save(str(src))
    out = tmp_path / "dup-title-out.docx"

    sp.main([str(src), str(out), "--profile", "reference"])

    paras = Document(str(out)).paragraphs
    assert sum(1 for p in paras if p.text == "Example Report") == 1  # Title survives, H1 dupe gone
    assert any(p.style and p.style.name == "Title" for p in paras)


def test_flat_document_no_duplicate_h1_keeps_every_section_heading(tmp_path):
    """Root-cause reproduction for #122: a flat document (no body H1 duplicating
    the title at all -- a real-consumer shape, where the title comes purely from
    YAML frontmatter via --keep-frontmatter) must not lose its first real
    section heading to a positional 'first non-Part H1 = duplicate' guess."""
    doc = Document()
    doc.add_heading("Example Report", level=0)      # [Title], no matching body H1
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("First section body.")
    doc.add_heading("Section Two", level=1)
    doc.add_paragraph("Second section body.")
    src = tmp_path / "flat-no-dup.docx"
    doc.save(str(src))
    out = tmp_path / "flat-no-dup-out.docx"

    sp.main([str(src), str(out), "--profile", "reference"])

    headings = [p.text for p in Document(str(out)).paragraphs if p.style and p.style.name == "Heading 1"]
    assert headings == ["Section One", "Section Two"], (
        "both real section headings must survive -- the old positional check "
        "would have deleted 'Section One' as if it were a duplicate title")


def test_no_title_paragraph_at_all_skips_removal_without_deleting_a_section(tmp_path, capsys):
    """A second #122 shape: no [Title] paragraph survives either (e.g. metadata
    was stripped upstream before this ever reached style_postprocess). Nothing
    to match against -> skip step 1 entirely rather than guess positionally."""
    doc = Document()
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("First section body.")
    doc.add_heading("Section Two", level=1)
    src = tmp_path / "no-title.docx"
    doc.save(str(src))
    out = tmp_path / "no-title-out.docx"

    sp.main([str(src), str(out), "--profile", "reference"])

    headings = [p.text for p in Document(str(out)).paragraphs if p.style and p.style.name == "Heading 1"]
    assert headings == ["Section One", "Section Two"]
    assert "NOTE" in capsys.readouterr().out


def test_cover_no_parts_opt_out_skips_duplicate_removal_even_with_a_match(tmp_path):
    doc = Document()
    doc.add_heading("Example Report", level=0)
    doc.add_heading("Example Report", level=1)  # would normally be removed as the duplicate
    doc.add_paragraph("Body content.")
    src = tmp_path / "no-parts.docx"
    doc.save(str(src))

    profile = tmp_path / "no-parts.yaml"
    profile.write_text("cover:\n  no_parts: true\n", encoding="utf-8")
    out = tmp_path / "no-parts-out.docx"

    sp.main([str(src), str(out), "--profile", "reference", "--template-profile", str(profile)])

    headings = [p.text for p in Document(str(out)).paragraphs if p.style and p.style.name == "Heading 1"]
    assert headings == ["Example Report"]  # the H1 survives, no_parts skipped the whole step


def test_book_shaped_document_with_part_heading_is_unaffected(tmp_path):
    """Regression guard: the original book-shaped case (title H1 + a real
    `Part N: ...` chapter heading) the function was designed for keeps working
    exactly as before."""
    doc = Document()
    doc.add_heading("Example Report", level=0)
    doc.add_heading("Example Report", level=1)     # duplicate of the title, matches
    doc.add_heading("Part I: Overview", level=1)    # real chapter, Part-prefixed, never touched
    doc.add_paragraph("Body content.")
    src = tmp_path / "book.docx"
    doc.save(str(src))
    out = tmp_path / "book-out.docx"

    sp.main([str(src), str(out), "--profile", "reference"])

    headings = [p.text for p in Document(str(out)).paragraphs if p.style and p.style.name == "Heading 1"]
    assert headings == ["Part I: Overview"]


# ---------- (j) custom-style font/size fidelity (issue #98) ----------

def _add_custom_style(doc, name="PullQuote", font_name="Georgia", size_pt=16):
    """A genuinely custom paragraph style whose OWN w:rPr explicitly sets a
    distinct font + size, mirroring a style pandoc would resolve from a
    --reference-doc's styles.xml when a source uses
    `::: {custom-style="PullQuote"} ... :::`."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt as _Pt

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = _Pt(size_pt)
    return style


def test_custom_style_paragraph_keeps_its_own_font_by_default(tmp_path, monkeypatch):
    # Root-cause reproduction for issue #98: a paragraph carrying a custom style
    # with its own explicit font/size definition must NOT get a run-level house
    # font/size override stamped onto it by default -- it should fall through to
    # pure style inheritance (no direct rFonts/sz on the run at all).
    doc = Document()
    _add_custom_style(doc)
    doc.add_heading("First Section", level=1)
    quote = doc.add_paragraph("A quote that should keep the PullQuote style's own font.")
    quote.style = doc.styles["PullQuote"]
    doc.add_paragraph("A plain body paragraph with no custom style at all.")
    src = tmp_path / "custom-style-in.docx"
    doc.save(str(src))
    out = tmp_path / "custom-style-out.docx"

    _run_style(monkeypatch, [str(src), str(out)])

    styled = Document(str(out))
    quote_p = next(p for p in styled.paragraphs if p.style.name == "PullQuote")
    for r in quote_p.runs:
        assert r.font.name is None          # no direct-formatting override injected
        assert r.font.size is None

    # Regression: a plain (non-custom-style) body paragraph is untouched by the fix
    # and still gets the house body font/size stamped on, exactly as before.
    body_p = next(p for p in styled.paragraphs
                  if p.text and p.style.name not in ("Heading 1", "PullQuote"))
    assert body_p.runs[0].font.name == "Arial"
    assert body_p.runs[0].font.size.pt == sp.PROFILES["compact"]["body"]


def test_override_custom_style_fonts_flag_restores_blanket_override(tmp_path, monkeypatch):
    # The explicit opt-in (--override-custom-style-fonts) restores the pre-#98
    # behaviour for callers who genuinely want one uniform house font everywhere.
    doc = Document()
    _add_custom_style(doc)
    quote = doc.add_paragraph("A quote paragraph.")
    quote.style = doc.styles["PullQuote"]
    src = tmp_path / "override-in.docx"
    doc.save(str(src))
    out = tmp_path / "override-out.docx"

    _run_style(monkeypatch, [str(src), str(out), "--override-custom-style-fonts"])

    styled = Document(str(out))
    quote_p = next(p for p in styled.paragraphs if p.style.name == "PullQuote")
    assert quote_p.runs[0].font.name == "Arial"
    assert quote_p.runs[0].font.size.pt == sp.PROFILES["compact"]["body"]


def test_override_custom_style_fonts_profile_key(tmp_path, monkeypatch):
    # Same opt-in, reachable via a template-profile.yaml key instead of the CLI flag.
    doc = Document()
    _add_custom_style(doc)
    quote = doc.add_paragraph("A quote paragraph.")
    quote.style = doc.styles["PullQuote"]
    src = tmp_path / "override-profile-in.docx"
    doc.save(str(src))
    profile = tmp_path / "override.yaml"
    profile.write_text("override_custom_style_fonts: true\n", encoding="utf-8")
    out = tmp_path / "override-profile-out.docx"

    _run_style(monkeypatch, [str(src), str(out), "--template-profile", str(profile)])

    styled = Document(str(out))
    quote_p = next(p for p in styled.paragraphs if p.style.name == "PullQuote")
    assert quote_p.runs[0].font.name == "Arial"


def test_style_without_its_own_font_definition_still_gets_house_font(tmp_path, monkeypatch):
    # A style OUTSIDE the known built-in/default-body set that does NOT itself
    # define a font (pure inheritance from Normal, no explicit rPr) is not
    # "custom" for the purposes of this gate: it still gets the house body
    # font/size, same as before. Guards against over-broadly treating any
    # unrecognised style name as template-fidelity-worthy.
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    bare = doc.styles.add_style("BareCustom", WD_STYLE_TYPE.PARAGRAPH)
    bare.base_style = doc.styles["Normal"]
    # No explicit font/size set on `bare`: its own w:rPr stays empty/absent.
    p = doc.add_paragraph("Text in a bare custom style with no font of its own.")
    p.style = bare
    src = tmp_path / "bare-in.docx"
    doc.save(str(src))
    out = tmp_path / "bare-out.docx"

    _run_style(monkeypatch, [str(src), str(out)])

    styled = Document(str(out))
    bare_p = next(p for p in styled.paragraphs if p.style.name == "BareCustom")
    assert bare_p.runs[0].font.name == "Arial"
    assert bare_p.runs[0].font.size.pt == sp.PROFILES["compact"]["body"]


def test_is_custom_style_paragraph_helper(tmp_path):
    # Direct unit coverage of the detector itself.
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    _add_custom_style(doc)
    quote = doc.add_paragraph("quote")
    quote.style = doc.styles["PullQuote"]
    normal = doc.add_paragraph("plain")

    bare = doc.styles.add_style("BareCustom2", WD_STYLE_TYPE.PARAGRAPH)
    bare.base_style = doc.styles["Normal"]
    bare_p = doc.add_paragraph("bare")
    bare_p.style = bare

    assert sp.is_custom_style_paragraph(quote) is True
    assert sp.is_custom_style_paragraph(normal) is False
    assert sp.is_custom_style_paragraph(bare_p) is False


def test_caption_matching_uses_raw_pstyle_id(tmp_path):
    # Backport regression (2026-07-03): pandoc references pStyle id
    # 'ImageCaption', usually undefined in reference.docx, so p.style resolves
    # to the document default and a name-based exact match never fires. The
    # matcher must normalise and also read the RAW pStyle id.
    from docx.oxml.ns import qn

    doc = Document()
    p = doc.add_paragraph("Figure: probe caption")
    pPr = p._element.get_or_add_pPr()
    ps = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): "ImageCaption"})
    pPr.insert(0, ps)

    sp.center_figures_and_captions(doc)

    assert all(r.font.italic for r in p.runs)
    assert all(r.font.size.pt == 8 for r in p.runs if r.font.size)

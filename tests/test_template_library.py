"""
Tests for api/templates.py, the Track J template library (chunk 6.3). Distinct
from tests/test_templates.py, which covers the pre-existing flat templates/*.md
genre pack (cv, cover-letter, pitch-*, purchase-request, executive-summary,
external-party-brief) -- an unrelated, untouched surface this module composes
with (project creation's best-effort seed still reads it by filename) but does
not replace.

Unit tests drive TemplateLibrary directly (built-in scan/get against the real
templates/library/ entries shipped in this repo, custom-root scan/get/import
against temporary directories); integration tests drive GET /templates,
GET /templates/{name}, POST /templates/import through the HTTP API in-process
via the same environ builder as test_store.py (no socket).

Covers: built-in entries are discovered and readable (plain-report, plain-deck
ship with the repo), invalid name rejected, unknown name 404s, a custom entry
shadows a built-in of the same name, a broken custom entry (missing
template.yaml) is skipped by scan() but not fatal; import_docx() happy path
(profile + metadata written, scaffold/profile round-trip through get()),
refuse-existing, invalid name, missing source; and the D15 guard set on
POST /templates/import (CSRF required, cross-origin rejected, path jail, 409
propagates).
"""

from __future__ import annotations

import io
import json
import sys
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from api import app as api_app  # noqa: E402
from api import templates as templates_mod  # noqa: E402


def _plain_docx(tmp_path: Path, name: str = "corporate.docx") -> Path:
    """A minimal but real .docx: python-docx's Document() derives from its own
    built-in default template, which carries a real (if generic) theme part,
    so read_theme() succeeds -- this only needs to exercise the import
    plumbing end-to-end, not theme-derivation fidelity (test_template_import.py
    already covers that in depth)."""
    doc = Document()
    doc.add_paragraph("Body text for the template-import wrapper test fixture.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def _entry(root: Path, name: str, **meta) -> Path:
    d = root / name
    d.mkdir(parents=True)
    payload = {"name": name, **meta}
    lines = "\n".join(f"{k}: {v!r}" if isinstance(v, str) else f"{k}: {v}"
                      for k, v in payload.items())
    (d / templates_mod.METADATA_NAME).write_text(lines + "\n", encoding="utf-8")
    return d


# ---------- built-in library (shipped with the repo) ----------

def test_builtin_entries_are_discovered(tmp_path):
    lib = templates_mod.TemplateLibrary(tmp_path / "no-custom-root-yet")
    rows = {r["name"]: r for r in lib.scan()}
    assert "plain-report" in rows and "plain-deck" in rows
    assert rows["plain-report"]["doc_type"] == "report"
    assert rows["plain-report"]["builtin"] is True
    assert rows["plain-deck"]["doc_type"] == "deck"


def test_builtin_get_includes_scaffold(tmp_path):
    lib = templates_mod.TemplateLibrary(tmp_path / "no-custom-root-yet")
    detail = lib.get("plain-report")
    assert detail["has_scaffold"] is True
    assert "[Report title]" in detail["scaffold"]
    assert detail["has_profile"] is False  # built-ins ship no derived profile
    assert "profile" not in detail


def test_get_rejects_invalid_name_and_unknown(tmp_path):
    lib = templates_mod.TemplateLibrary(tmp_path)
    with pytest.raises(templates_mod.TemplateError):
        lib.get("Not A Slug")
    with pytest.raises(templates_mod.TemplateError):
        lib.get("no-such-template")


# ---------- custom root: scan / shadowing / broken entries ----------

def test_custom_entry_shadows_builtin_of_same_name(tmp_path):
    custom = tmp_path / "custom"
    _entry(custom, "plain-report", doc_type="report",
          description="operator override of the built-in")
    lib = templates_mod.TemplateLibrary(custom)
    rows = {r["name"]: r for r in lib.scan()}
    assert rows["plain-report"]["builtin"] is False
    assert rows["plain-report"]["description"] == "operator override of the built-in"
    detail = lib.get("plain-report")
    assert detail["builtin"] is False


def test_scan_skips_broken_custom_entry(tmp_path):
    custom = tmp_path / "custom"
    _entry(custom, "good", doc_type="report")
    broken = custom / "broken"
    broken.mkdir(parents=True)  # a directory with no template.yaml at all
    lib = templates_mod.TemplateLibrary(custom)
    names = {r["name"] for r in lib.scan()}
    assert "good" in names
    assert "broken" not in names


def test_scan_merges_custom_and_builtin(tmp_path):
    custom = tmp_path / "custom"
    _entry(custom, "acme-report", doc_type="report", description="a custom one")
    lib = templates_mod.TemplateLibrary(custom)
    names = {r["name"] for r in lib.scan()}
    assert names == {"plain-report", "plain-deck", "acme-report"}


# ---------- import_docx() ----------

def test_import_docx_happy_path(tmp_path):
    docx = _plain_docx(tmp_path)
    lib = templates_mod.TemplateLibrary(tmp_path / "custom")
    result = lib.import_docx("acme-report", docx, doc_type="report",
                             description="Imported for a test", copy_reference=True)
    assert result["name"] == "acme-report"
    assert result["doc_type"] == "report"
    assert result["derived_from"] == "corporate.docx"
    assert result["has_profile"] is True
    assert "profile" in result
    assert result["idempotency_check_passed"] is None  # no --check requested
    assert (tmp_path / "custom" / "acme-report" / "reference.docx").is_file()

    # round-trips through a fresh get()
    detail = lib.get("acme-report")
    assert detail["description"] == "Imported for a test"


def test_import_docx_refuses_existing(tmp_path):
    docx = _plain_docx(tmp_path)
    lib = templates_mod.TemplateLibrary(tmp_path / "custom")
    lib.import_docx("acme-report", docx)
    with pytest.raises(templates_mod.TemplateExistsError):
        lib.import_docx("acme-report", docx)


def test_import_docx_rejects_invalid_name(tmp_path):
    docx = _plain_docx(tmp_path)
    lib = templates_mod.TemplateLibrary(tmp_path / "custom")
    with pytest.raises(templates_mod.TemplateError):
        lib.import_docx("Not A Slug", docx)


def test_import_docx_missing_source(tmp_path):
    lib = templates_mod.TemplateLibrary(tmp_path / "custom")
    with pytest.raises(templates_mod.TemplateError):
        lib.import_docx("acme-report", tmp_path / "does-not-exist.docx")


def test_import_docx_defaults_when_metadata_omitted(tmp_path):
    docx = _plain_docx(tmp_path, name="unbranded.docx")
    lib = templates_mod.TemplateLibrary(tmp_path / "custom")
    result = lib.import_docx("plain-import", docx)
    assert result["doc_type"] == "report"  # default
    assert "unbranded.docx" in result["description"]
    assert result["diagram_scaffolds"] == []


# ---------- HTTP API integration ----------

def _call(api, method, path, body=None, extra=None):
    app = api_app.make_wsgi_app(api)
    raw = json.dumps(body).encode("utf-8") if body is not None else b""
    environ = {
        "REQUEST_METHOD": method, "PATH_INFO": path, "HTTP_HOST": "127.0.0.1:8385",
        "REMOTE_ADDR": "127.0.0.1", "CONTENT_LENGTH": str(len(raw)),
        "wsgi.input": io.BytesIO(raw),
    }
    for k, v in (extra or {}).items():
        environ[k] = v
    captured = {}

    def start_response(status, headers):
        captured["status"] = status

    payload = b"".join(app(environ, start_response)).decode("utf-8")
    return int(captured["status"].split()[0]), json.loads(payload)


def _csrf_token(api):
    _, data = _call(api, "GET", "/session")
    return data["csrf_token"]


def test_api_templates_list_includes_builtins(tmp_path):
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    code, data = _call(api, "GET", "/templates")
    assert code == 200
    names = {t["name"] for t in data["templates"]}
    assert {"plain-report", "plain-deck"} <= names


def test_api_template_detail_and_404(tmp_path):
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    code, detail = _call(api, "GET", "/templates/plain-report")
    assert code == 200 and detail["doc_type"] == "report"
    code, _ = _call(api, "GET", "/templates/does-not-exist")
    assert code == 404


def test_api_import_requires_csrf(tmp_path):
    docx = _plain_docx(tmp_path)
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    code, data = _call(api, "POST", "/templates/import",
                       {"name": "acme-report", "source": docx.name})
    assert code == 403
    assert "CSRF" in data["error"]


def test_api_import_with_csrf_creates_entry(tmp_path):
    docx = _plain_docx(tmp_path)
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    token = _csrf_token(api)
    code, data = _call(api, "POST", "/templates/import",
                       {"name": "acme-report", "source": docx.name, "doc_type": "report"},
                       extra={"HTTP_X_CSRF_TOKEN": token})
    assert code == 200
    assert data["name"] == "acme-report"
    code, _ = _call(api, "GET", "/templates/acme-report")
    assert code == 200


def test_api_import_duplicate_is_409(tmp_path):
    docx = _plain_docx(tmp_path)
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    token = _csrf_token(api)
    body = {"name": "acme-report", "source": docx.name}
    _call(api, "POST", "/templates/import", body, extra={"HTTP_X_CSRF_TOKEN": token})
    code, _ = _call(api, "POST", "/templates/import", body,
                    extra={"HTTP_X_CSRF_TOKEN": token})
    assert code == 409


def test_api_import_cross_origin_rejected(tmp_path):
    docx = _plain_docx(tmp_path)
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    token = _csrf_token(api)
    code, _ = _call(api, "POST", "/templates/import",
                    {"name": "acme-report", "source": docx.name},
                    extra={"HTTP_X_CSRF_TOKEN": token, "HTTP_ORIGIN": "https://evil.example"})
    assert code == 403


def test_api_import_source_path_jail(tmp_path):
    api = api_app.RenderfactApi(root=tmp_path, templates_root=tmp_path / "custom")
    token = _csrf_token(api)
    code, data = _call(api, "POST", "/templates/import",
                       {"name": "acme-report", "source": "../outside.docx"},
                       extra={"HTTP_X_CSRF_TOKEN": token})
    assert code == 403
    assert "escapes" in data["error"]

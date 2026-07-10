"""
Tests for gates/content_scan.py: the generic post-render content-safety gate
(issue #71, D18). Fixtures are built programmatically with python-docx (no
binary fixtures, per CONTRIBUTING.md).

Covers: the paragraph + table-cell (incl. nested-table) scan itself, both a
clean pass and a deliberate hit; the required-pattern refusal (no default
pattern ships, D18); pattern resolution order (CLI --pattern/--pattern-file
beats the RENDERFACT_GATE_PATTERN* env fallback, which exists so this script
also works as a zero-arg render-doc.sh hook); and the CLI exit-code contract.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from gates import content_scan  # noqa: E402


def _build_clean_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("This is a perfectly ordinary paragraph about logistics.")
    doc.add_paragraph("Another paragraph, still nothing sensitive here.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Line item"
    table.cell(0, 1).text = "Quantity"
    table.cell(1, 0).text = "Widgets"
    table.cell(1, 1).text = "42"
    doc.save(str(path))
    return path


def _build_hit_in_paragraph_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Nothing to see here.")
    doc.add_paragraph("The internal ceiling is EUR 450.000, do not disclose it.")
    doc.save(str(path))
    return path


def _build_hit_in_table_cell_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Body text is clean.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Budget ceiling"
    table.cell(0, 1).text = "EUR 12.500"
    doc.save(str(path))
    return path


def _build_hit_in_nested_table_doc(path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Body text is clean.")
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.cell(0, 0).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "Nested EUR 9.999 figure"
    doc.save(str(path))
    return path


CURRENCY_PATTERN = r"EUR\s*[\d.,]+"


# ---- scan_document: the core paragraph + table-cell scan ----

def test_clean_document_produces_no_findings(tmp_path):
    doc_path = _build_clean_doc(tmp_path / "clean.docx")
    findings = content_scan.scan_document(doc_path, [re.compile(CURRENCY_PATTERN)])
    assert findings == []


def test_hit_in_body_paragraph_is_found(tmp_path):
    doc_path = _build_hit_in_paragraph_doc(tmp_path / "hit-para.docx")
    findings = content_scan.scan_document(doc_path, [re.compile(CURRENCY_PATTERN)])
    assert len(findings) == 1
    assert "paragraph" in findings[0]
    assert "EUR 450.000" in findings[0]


def test_hit_in_table_cell_is_found(tmp_path):
    doc_path = _build_hit_in_table_cell_doc(tmp_path / "hit-cell.docx")
    findings = content_scan.scan_document(doc_path, [re.compile(CURRENCY_PATTERN)])
    assert len(findings) == 1
    assert "table 0 cell-paragraph" in findings[0]
    assert "EUR 12.500" in findings[0]


def test_hit_in_nested_table_is_found(tmp_path):
    doc_path = _build_hit_in_nested_table_doc(tmp_path / "hit-nested.docx")
    findings = content_scan.scan_document(doc_path, [re.compile(CURRENCY_PATTERN)])
    assert len(findings) == 1
    assert "EUR 9.999" in findings[0]


def test_multiple_patterns_are_ored(tmp_path):
    doc_path = _build_clean_doc(tmp_path / "clean2.docx")
    patterns = [re.compile(r"logistics"), re.compile(r"nonexistent-xyz")]
    findings = content_scan.scan_document(doc_path, patterns)
    assert len(findings) == 1  # 'logistics' hits the first paragraph


# ---- pattern resolution: CLI beats env, env exists for zero-arg hook use ----

def test_resolve_patterns_cli_flag_wins_over_env(monkeypatch):
    monkeypatch.setenv("RENDERFACT_GATE_PATTERN", "should-not-be-used")
    result = content_scan.resolve_patterns(["from-cli"], None)
    assert result == ["from-cli"]


def test_resolve_patterns_falls_back_to_env_pattern(monkeypatch):
    monkeypatch.setenv("RENDERFACT_GATE_PATTERN", "from-env")
    result = content_scan.resolve_patterns([], None)
    assert result == ["from-env"]


def test_resolve_patterns_falls_back_to_env_pattern_file(tmp_path, monkeypatch):
    pf = tmp_path / "patterns.txt"
    pf.write_text("# a comment\nfoo\n\nbar\n", encoding="utf-8")
    monkeypatch.delenv("RENDERFACT_GATE_PATTERN", raising=False)
    monkeypatch.setenv("RENDERFACT_GATE_PATTERN_FILE", str(pf))
    result = content_scan.resolve_patterns([], None)
    assert result == ["foo", "bar"]


def test_resolve_patterns_none_configured_is_empty(monkeypatch):
    monkeypatch.delenv("RENDERFACT_GATE_PATTERN", raising=False)
    monkeypatch.delenv("RENDERFACT_GATE_PATTERN_FILE", raising=False)
    assert content_scan.resolve_patterns([], None) == []


def test_pattern_file_cli_reads_comments_and_blanks_correctly(tmp_path):
    pf = tmp_path / "patterns.txt"
    pf.write_text("# comment\n\n  \nEUR\\s*[\\d.,]+\n", encoding="utf-8")
    result = content_scan.resolve_patterns([], str(pf))
    assert result == [r"EUR\s*[\d.,]+"]


# ---- CLI: main() exit codes ----

def test_main_exits_2_when_no_pattern_configured(tmp_path, monkeypatch, capsys):
    monkeypatch.delenv("RENDERFACT_GATE_PATTERN", raising=False)
    monkeypatch.delenv("RENDERFACT_GATE_PATTERN_FILE", raising=False)
    doc_path = _build_clean_doc(tmp_path / "clean.docx")
    rc = content_scan.main([str(doc_path)])
    assert rc == 2
    assert "no pattern configured" in capsys.readouterr().err


def test_main_exits_0_on_clean_document_with_cli_pattern(tmp_path, capsys):
    doc_path = _build_clean_doc(tmp_path / "clean.docx")
    rc = content_scan.main([str(doc_path), "--pattern", CURRENCY_PATTERN])
    assert rc == 0
    assert "content-safety gate OK" in capsys.readouterr().out


def test_main_exits_1_on_hit_with_cli_pattern(tmp_path, capsys):
    doc_path = _build_hit_in_paragraph_doc(tmp_path / "hit.docx")
    rc = content_scan.main([str(doc_path), "--pattern", CURRENCY_PATTERN])
    assert rc == 1
    out = capsys.readouterr().out
    assert "CONTENT-SAFETY GATE FAIL" in out
    assert "EUR 450.000" in out


def test_main_exits_1_via_env_pattern_zero_arg_invocation(tmp_path, monkeypatch, capsys):
    """The exact call shape render-doc.sh's POSTRENDER_GATE_SCRIPT hook uses:
    only the docx path on the command line, pattern supplied via env."""
    monkeypatch.setenv("RENDERFACT_GATE_PATTERN", CURRENCY_PATTERN)
    doc_path = _build_hit_in_table_cell_doc(tmp_path / "hit-cell.docx")
    rc = content_scan.main([str(doc_path)])
    assert rc == 1


def test_main_exits_2_on_missing_file(tmp_path, capsys):
    rc = content_scan.main([str(tmp_path / "does-not-exist.docx"), "--pattern", "x"])
    assert rc == 2
    assert "not a file" in capsys.readouterr().err


def test_main_exits_2_on_invalid_regex(tmp_path, capsys):
    doc_path = _build_clean_doc(tmp_path / "clean.docx")
    rc = content_scan.main([str(doc_path), "--pattern", "("])  # unbalanced group
    assert rc == 2
    assert "invalid regex" in capsys.readouterr().err


def test_main_exits_2_on_missing_pattern_file(tmp_path, capsys):
    doc_path = _build_clean_doc(tmp_path / "clean.docx")
    rc = content_scan.main([str(doc_path), "--pattern-file", str(tmp_path / "nope.txt")])
    assert rc == 2
    assert "pattern file not found" in capsys.readouterr().err

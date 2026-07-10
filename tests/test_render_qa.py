"""
Tests for lint/render_qa.py, the deterministic post-render QA gate.

Covers: probe loading (defaults, consumer yaml merge, defaults-off), the leaks
scan (hit counting, fail-on-hits exit semantics), the paras overweight ranking
against a real python-docx document, the tables geometry scan (pressure and
slack ratios, including the over-allocated-column blind spot from issue #90),
the figs inventory MISSING path, and dispatch via render.py as a real
subprocess.
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "lint"))
sys.path.insert(0, str(REPO_ROOT))

import render_qa  # noqa: E402


def test_default_probes_are_generic_only():
    # consumer-specific probe terms are built dynamically so THIS file stays
    # clean under the repo's own denylist gate (same trick as test_denylist_scan)
    probes = render_qa.load_probes(None)
    joined = " ".join(probes.values()).lower()
    for consumer_term in ("ss" + "ot", "vau" + "lt", "rc" + "de", "ken" + "nis", ".cla" + "ude"):
        assert consumer_term not in joined


def test_probes_yaml_merges_over_defaults(tmp_path):
    p = tmp_path / "probes.yaml"
    p.write_text('probes:\n  "codename": "\\\\bNIGHTJAR\\\\b"\n', encoding="utf-8")
    probes = render_qa.load_probes(str(p))
    assert "codename" in probes
    assert "surviving wikilink brackets" in probes  # defaults kept


def test_probes_defaults_can_be_disabled(tmp_path):
    p = tmp_path / "probes.yaml"
    p.write_text('probes:\n  "codename": "NIGHTJAR"\n', encoding="utf-8")
    probes = render_qa.load_probes(str(p), use_defaults=False)
    assert list(probes) == ["codename"]


def test_leaks_scan_counts_and_fail_on_hits(tmp_path, capsys):
    txt = tmp_path / "full.txt"
    txt.write_text(
        "clean page one\n\fpage two has a [[wikilink]] leak\nand a title (2026-05) suffix\n",
        encoding="utf-8",
    )
    rc = render_qa.cmd_leaks(str(txt), render_qa.load_probes(None), fail_on_hits=True)
    out = capsys.readouterr().out
    assert rc == 1
    assert "TOTAL leak hits: 2" in out

    clean = tmp_path / "clean.txt"
    clean.write_text("nothing to see\n", encoding="utf-8")
    rc = render_qa.cmd_leaks(str(clean), render_qa.load_probes(None), fail_on_hits=True)
    assert rc == 0


def _make_docx(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Section One", level=1)
    doc.add_paragraph("short paragraph")
    doc.add_paragraph(" ".join(["word"] * 150))  # overweight
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "k"
    table.rows[0].cells[1].text = "a much longer content cell " * 4
    table.rows[1].cells[0].text = "k2"
    table.rows[1].cells[1].text = "more content here"
    path = tmp_path / "doc.docx"
    doc.save(str(path))
    return path


def test_paras_ranks_overweight_paragraph(tmp_path, capsys):
    path = _make_docx(tmp_path)
    rc = render_qa.cmd_paras(str(path), top=5, limit_words=110)
    out = capsys.readouterr().out
    assert rc == 0
    assert "1 paragraph(s) >= 110 words" in out
    assert "Section One" in out  # attributed to its nearest heading


def test_tables_scan_runs_on_real_table(tmp_path, capsys):
    path = _make_docx(tmp_path)
    rc = render_qa.cmd_tables(str(path), top=5)
    out = capsys.readouterr().out
    assert rc == 0
    assert "TABLES geometry" in out


def _make_lopsided_docx(tmp_path):
    """Worked example from issue #90: a 4-column table where an ordinal
    (row-number) column is given generous width for single-digit content,
    while a prose column carrying most of the real content is comparatively
    starved. The ordinal column's content share never clears the pressure
    eligibility floor, so the existing pressure/squeezed-col signal has
    nothing to say about it; the new slack/wasteful-col signal should.
    """
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    table = doc.add_table(rows=5, cols=4)
    table.autofit = False
    col_widths = [Inches(2.4), Inches(1.0), Inches(1.3), Inches(1.3)]
    for ci, w in enumerate(col_widths):
        table.columns[ci].width = w
    rows = [
        ("1", "a much longer prose sentence describing the finding for row one", "note a", "note b"),
        ("2", "another longer prose sentence describing the finding for row two", "note a", "note b"),
        ("3", "yet another longer prose sentence for row three findings detail", "note a", "note b"),
        ("4", "a fourth longer prose sentence describing row four in more detail", "note a", "note b"),
        ("5", "a fifth longer prose sentence with the row five description text", "note a", "note b"),
    ]
    for ri, rowdata in enumerate(rows):
        for ci, val in enumerate(rowdata):
            table.rows[ri].cells[ci].width = col_widths[ci]
            table.rows[ri].cells[ci].text = val
    path = tmp_path / "lopsided.docx"
    doc.save(str(path))
    return path


def test_slack_signal_flags_over_allocated_column_pressure_misses(tmp_path, capsys):
    path = _make_lopsided_docx(tmp_path)
    rc = render_qa.cmd_tables(str(path), top=5)
    out = capsys.readouterr().out
    assert rc == 0
    # the starved prose column (index 1) is what the existing pressure signal
    # catches
    assert "squeezed-col=1" in out
    # the over-allocated ordinal column (index 0) never clears the pressure
    # eligibility floor, so pressure alone never names it; the new slack
    # signal is the only one that does
    assert "squeezed-col=0" not in out
    assert "wasteful-col=0" in out


def test_slack_ratio_flags_over_allocated_tiny_content_column():
    # single-digit content (near-zero cshare) given a full quarter of the
    # table's width: clearly wasteful, should read well above 1.0
    ratio = render_qa._slack_ratio(wshare=0.25, cshare=0.01)
    assert ratio > 1.8


def test_slack_ratio_does_not_flag_genuinely_proportional_small_column():
    # a small column given a proportionally small width: the shared floor
    # keeps this at or below 1.0, not flagged
    ratio = render_qa._slack_ratio(wshare=0.03, cshare=0.03)
    assert ratio <= 1.0


def test_slack_ratio_floor_bounds_near_zero_content():
    # as cshare approaches zero the ratio must not blow up to infinity; it is
    # capped by wshare / SLACK_CSHARE_FLOOR
    ratio = render_qa._slack_ratio(wshare=0.10, cshare=0.0)
    assert ratio == pytest.approx(0.10 / render_qa.SLACK_CSHARE_FLOOR)


def test_slack_and_pressure_ratios_are_inverse_shaped():
    # same share pair: pressure rewards high cshare/low wshare, slack rewards
    # the opposite (high wshare/low cshare)
    wshare, cshare = 0.30, 0.02
    pressure = render_qa._pressure_ratio(wshare, cshare)
    slack = render_qa._slack_ratio(wshare, cshare)
    assert slack > 1.8
    assert pressure < 1.0  # this column is over-allocated, not squeezed


def test_single_column_table_is_never_flagged(tmp_path, capsys):
    # degenerate case: one column carries 100% of both width and content, so
    # both ratios must land exactly at 1.0 (never flagged)
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "only column, only content here"
    table.rows[1].cells[0].text = "more content in the only column"
    path = tmp_path / "single_col.docx"
    doc.save(str(path))

    rc = render_qa.cmd_tables(str(path), top=5)
    out = capsys.readouterr().out
    assert rc == 0
    assert "pressure= 1.0" in out
    assert "slack= 1.0" in out


def test_figs_reports_missing_reference(tmp_path, capsys):
    md = tmp_path / "src.md"
    md.write_text("intro\n\n![diagram](figures/missing.png)\n", encoding="utf-8")
    rc = render_qa.cmd_figs(str(md))
    out = capsys.readouterr().out
    assert rc == 0
    assert "MISSING" in out and "figures/missing.png" in out


def test_render_entrypoint_dispatches_qa_mode(tmp_path):
    txt = tmp_path / "full.txt"
    txt.write_text("has a [[leak]]\n", encoding="utf-8")
    result = subprocess.run(
        [sys.executable, str(REPO_ROOT / "render.py"), "qa", "leaks", str(txt),
         "--fail-on-hits"],
        capture_output=True, text=True,
    )
    assert result.returncode == 1
    assert "TOTAL leak hits: 1" in result.stdout

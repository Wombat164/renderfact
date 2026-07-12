"""
Tests for docstyle/ooxml_theme.py + docstyle/template_import.py (chunk C7a: DOCX
template import, style axis).

All DOCX fixtures are built in tmp_path via python-docx, then a KNOWN DrawingML
theme is injected by rewriting word/theme/theme1.xml inside the zip directly
(python-docx has no theme-writing API; this mirrors how docstyle/
heading_numbering.py's own tests rewrite a docx zip in place). Covers: theme
resolution (srgbClr AND sysClr/lastClr), ThemeError on a package with no theme
part, the derived profile's provenance header + only-derivable-keys-uncommented
shape, section geometry mapping, the --check comparison function in isolation,
one real --check integration test through render-doc.sh (skipif pandoc/bash
absent), and one real subprocess dispatch test through render.py.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(REPO_ROOT / "docstyle"))

from docstyle import template_import as ti  # noqa: E402
from docstyle.ooxml_theme import Theme, ThemeError, read_theme  # noqa: E402

RENDER_PY = REPO_ROOT / "render.py"

KNOWN_ACCENT1 = "FF6600"
KNOWN_DK1 = "222222"  # via sysClr lastClr
KNOWN_MAJOR_FONT = "Georgia"
KNOWN_MINOR_FONT = "Verdana"

_THEME_XML = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<a:theme xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" name="Test Theme">
  <a:themeElements>
    <a:clrScheme name="Test">
      <a:dk1><a:sysClr val="windowText" lastClr="{KNOWN_DK1}"/></a:dk1>
      <a:lt1><a:sysClr val="window" lastClr="FFFFFF"/></a:lt1>
      <a:dk2><a:srgbClr val="1F1F1F"/></a:dk2>
      <a:lt2><a:srgbClr val="EEEEEE"/></a:lt2>
      <a:accent1><a:srgbClr val="{KNOWN_ACCENT1}"/></a:accent1>
      <a:accent2><a:srgbClr val="336699"/></a:accent2>
      <a:accent3><a:srgbClr val="669933"/></a:accent3>
      <a:accent4><a:srgbClr val="993366"/></a:accent4>
      <a:accent5><a:srgbClr val="999933"/></a:accent5>
      <a:accent6><a:srgbClr val="336633"/></a:accent6>
      <a:hlink><a:srgbClr val="0563C1"/></a:hlink>
      <a:folHlink><a:srgbClr val="954F72"/></a:folHlink>
    </a:clrScheme>
    <a:fontScheme name="Test">
      <a:majorFont><a:latin typeface="{KNOWN_MAJOR_FONT}"/></a:majorFont>
      <a:minorFont><a:latin typeface="{KNOWN_MINOR_FONT}"/></a:minorFont>
    </a:fontScheme>
  </a:themeElements>
</a:theme>"""


def _rewrite_zip_entry(docx_path: Path, entry: str, data: bytes | None) -> None:
    """Rewrite one zip entry's content (data=None deletes the entry). Mirrors
    docstyle/heading_numbering.process()'s own in-place zip rewrite pattern."""
    with zipfile.ZipFile(docx_path) as zin:
        items = {i.filename: zin.read(i.filename) for i in zin.infolist()}
    if data is None:
        items.pop(entry, None)
    else:
        items[entry] = data
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in items.items():
            zout.writestr(name, blob)


def _build_branded_template(tmp_path, name="corporate.docx") -> Path:
    """A synthetic 'corporate template': python-docx styles (Normal font/colour,
    Heading 1 colour) + explicit section geometry, with a KNOWN theme injected
    (accent1 srgbClr, dk1 sysClr/lastClr, known major/minor fonts)."""
    path = tmp_path / name
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = KNOWN_MINOR_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    h1 = doc.styles["Heading 1"]
    h1.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

    doc.add_heading("Corporate Title", level=1)
    doc.add_paragraph("Body paragraph text for the corporate template test fixture.")
    doc.save(str(path))

    _rewrite_zip_entry(path, "word/theme/theme1.xml", _THEME_XML.encode("utf-8"))
    return path


def _template_with_no_theme(tmp_path, name="no-theme.docx") -> Path:
    doc = Document()
    doc.add_paragraph("no theme here")
    path = tmp_path / name
    doc.save(str(path))
    _rewrite_zip_entry(path, "word/theme/theme1.xml", None)
    return path


# ---------- (a) ooxml_theme.py: theme resolution incl. sysClr/lastClr ----------

def test_read_theme_resolves_srgb_and_sysclr(tmp_path):
    template = _build_branded_template(tmp_path)
    theme = read_theme(template)
    assert isinstance(theme, Theme)
    assert theme.colors["accent1"] == KNOWN_ACCENT1     # srgbClr
    assert theme.colors["dk1"] == KNOWN_DK1              # sysClr/lastClr
    assert theme.colors["lt1"] == "FFFFFF"               # sysClr/lastClr, second role
    assert theme.fonts["major"] == KNOWN_MAJOR_FONT
    assert theme.fonts["minor"] == KNOWN_MINOR_FONT


def test_read_theme_raises_on_missing_theme_part(tmp_path):
    path = _template_with_no_theme(tmp_path)
    with pytest.raises(ThemeError):
        read_theme(path)


def test_read_theme_raises_on_non_ooxml_zip(tmp_path):
    path = tmp_path / "not-office.docx"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("hello.txt", "not an office package")
    with pytest.raises(ThemeError):
        read_theme(path)


# ---------- (b) derived profile: injected values + provenance + only-derivable-uncommented ----------

def test_derive_profile_and_yaml_carry_injected_values(tmp_path):
    template = _build_branded_template(tmp_path)
    theme = read_theme(template)
    doc = Document(str(template))
    dp = ti.derive_profile(doc, theme)

    assert dp.derived["accent"] == KNOWN_ACCENT1
    assert dp.derived["body"] == "222222"
    assert dp.derived["font"] == KNOWN_MINOR_FONT
    assert dp.derived["page_width_cm"] == 21.0
    assert dp.derived["page_height_cm"] == 29.7
    assert dp.derived["margin_cm"] == 2.0
    # never derivable from a DOCX template in v1 scope
    assert set(dp.not_derived) >= {"body_muted", "table_body", "zebra"}

    prov = ti.build_import_provenance(template, date_str="2026-07-03")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert f"source template: {template.name}" in yaml_text
    m = re.search(r"source sha256:\s+([0-9a-f]{64})", yaml_text)
    assert m, yaml_text
    assert "import date:     2026-07-03" in yaml_text
    assert "tool_version:" in yaml_text

    # derivable keys present UNCOMMENTED as real yaml assignments
    assert f'accent: "{KNOWN_ACCENT1}"' in yaml_text
    assert f"font: {KNOWN_MINOR_FONT}" in yaml_text
    assert "page_width_cm: 21.0" in yaml_text
    assert "margin_cm: 2.0" in yaml_text
    # never a bare (uncommented) line for the never-derivable keys
    for key in ("body_muted", "table_body", "zebra"):
        assert re.search(rf"^{key}:", yaml_text, re.MULTILINE) is None
        assert re.search(rf"^# {key}: not derivable", yaml_text, re.MULTILINE) is not None


# ---------- (k) header/footer marking detection at import time (#123) ----------

def _template_with_header_marking(tmp_path, marking_text, name="marked.docx") -> Path:
    doc = Document()
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = marking_text
    doc.add_paragraph("Body paragraph.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def test_derive_profile_detects_marking_like_header_text(tmp_path):
    template = _template_with_header_marking(tmp_path, "UNCLASS")
    theme = read_theme(_build_branded_template(tmp_path))  # any valid theme; not under test here
    doc = Document(str(template))
    dp = ti.derive_profile(doc, theme)

    assert "UNCLASS" in dp.marking_findings


def test_derive_profile_finds_no_marking_when_header_is_plain(tmp_path):
    template = _template_with_header_marking(tmp_path, "Acme Corp Quarterly Report")
    theme = read_theme(_build_branded_template(tmp_path))
    doc = Document(str(template))
    dp = ti.derive_profile(doc, theme)

    assert dp.marking_findings == []


def test_render_profile_yaml_flags_detected_marking(tmp_path):
    template = _template_with_header_marking(tmp_path, "CONFIDENTIAL")
    theme = read_theme(_build_branded_template(tmp_path))
    doc = Document(str(template))
    dp = ti.derive_profile(doc, theme)
    prov = ti.build_import_provenance(template, date_str="2026-07-12")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert "NOT DERIVED, but flagged (#123)" in yaml_text
    assert "CONFIDENTIAL" in yaml_text
    assert "compliance issue, not a cosmetic one" in yaml_text


def test_render_profile_yaml_no_marking_block_when_nothing_detected(tmp_path):
    template = _build_branded_template(tmp_path)  # header text is plain, no marking
    theme = read_theme(template)
    doc = Document(str(template))
    dp = ti.derive_profile(doc, theme)
    prov = ti.build_import_provenance(template, date_str="2026-07-12")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert "flagged (#123)" not in yaml_text
    assert "Optional marking / cover behaviour: authoring hints" in yaml_text  # the generic block instead


def test_derive_profile_falls_back_to_theme_when_style_unset(tmp_path):
    # Normal/Heading 1 carry NO explicit colour/font -> theme accent1/dk1/minor used.
    path = tmp_path / "theme-only.docx"
    doc = Document()
    doc.add_heading("Heading with no explicit colour", level=1)
    doc.add_paragraph("Body with no explicit colour or font.")
    doc.save(str(path))
    _rewrite_zip_entry(path, "word/theme/theme1.xml", _THEME_XML.encode("utf-8"))

    theme = read_theme(path)
    doc2 = Document(str(path))
    dp = ti.derive_profile(doc2, theme)
    assert dp.derived["accent"] == KNOWN_ACCENT1
    assert dp.derived["body"] == KNOWN_DK1
    assert dp.derived["font"] == KNOWN_MINOR_FONT


# ---------- (b2) per-style font overrides (issue #97) ----------

def test_derive_style_font_overrides_only_genuine_differences(tmp_path):
    """Two paragraph styles carrying genuinely different fonts: the derived
    profile must capture the DISTINCT one (Quote: Georgia, an existing
    built-in python-docx style repurposed here rather than a custom one, so
    the walk is proven to cover built-ins too) as a per-style override, and
    must NOT emit a redundant entry for a style (Caption) whose explicit font
    just happens to match the global default (Verdana, from Normal). Word's
    own default template also ships one intrinsic deviation of its own ('macro'
    / Macro Text -> Courier), asserted here too so the test stays honest about
    what the walk actually returns rather than special-casing it away."""
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    doc.styles["Quote"].font.name = "Georgia"
    doc.styles["Caption"].font.name = KNOWN_MINOR_FONT  # matches the global default
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)

    assert dp.derived["font"] == KNOWN_MINOR_FONT
    assert dp.style_fonts == {"Quote": "Georgia", "macro": "Courier"}
    assert "Caption" not in dp.style_fonts


def test_derive_style_font_overrides_filters_out_when_global_matches(tmp_path):
    """When the global font itself equals python-docx's one intrinsic built-in
    deviation (the 'macro' style's Courier), that style is correctly filtered
    as redundant, proving the "only genuine differences" rule applies
    symmetrically to built-in noise, not just custom styles."""
    path = tmp_path / "uniform.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Courier"
    doc.add_paragraph("body")
    doc.save(str(path))

    theme = Theme(colors={}, fonts={})
    doc2 = Document(str(path))
    dp = ti.derive_profile(doc2, theme)
    assert dp.derived["font"] == "Courier"
    assert dp.style_fonts == {}


def test_render_profile_yaml_emits_styles_block_when_overrides_found(tmp_path):
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    doc.styles["Quote"].font.name = "Georgia"
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert "styles:" in yaml_text
    assert '"Quote":' in yaml_text
    assert "font: Georgia" in yaml_text


def test_render_profile_yaml_notes_no_style_overrides_when_none_found(tmp_path):
    # Force the global font to equal python-docx's one intrinsic built-in
    # deviation (the 'macro' style ships Courier by default) so every
    # paragraph style in this template resolves to the SAME font.
    path = tmp_path / "uniform.docx"
    doc = Document()
    doc.styles["Normal"].font.name = "Courier"
    doc.add_paragraph("body")
    doc.save(str(path))

    theme = Theme(colors={}, fonts={})
    doc2 = Document(str(path))
    dp = ti.derive_profile(doc2, theme)
    prov = ti.build_import_provenance(path, date_str="2026-07-10")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert "styles:" not in yaml_text
    assert "No per-style font overrides found" in yaml_text


def test_derived_style_fonts_round_trip_through_consumer(tmp_path, monkeypatch):
    """End-to-end (derivation + consumption): derive a profile from a template
    with two distinct paragraph-style fonts, then run style_postprocess on a
    document using both styles, and assert the OUTPUT carries the correct
    per-style font (not the global one) where a style overrides it, and the
    global font everywhere else (issue #97)."""
    import importlib

    from docstyle import style_postprocess as sp

    importlib.reload(sp)  # isolate module globals from any earlier test

    template = _build_branded_template(tmp_path, name="two-font-template.docx")
    doc = Document(str(template))
    doc.styles["Quote"].font.name = "Georgia"
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    assert dp.derived["font"] == KNOWN_MINOR_FONT
    assert dp.style_fonts["Quote"] == "Georgia"

    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    profile_path = tmp_path / "template-profile.yaml"
    profile_path.write_text(ti.render_profile_yaml(prov, dp), encoding="utf-8")

    consumer_doc = Document()
    consumer_doc.add_heading("Body Section", level=1)
    consumer_doc.add_paragraph(
        "A quoted line long enough to carry a run.", style=consumer_doc.styles["Quote"])
    consumer_doc.add_paragraph("A plain body paragraph using the global font.")
    src = tmp_path / "consumer-in.docx"
    consumer_doc.save(str(src))

    out = tmp_path / "consumer-out.docx"
    monkeypatch.setattr(sys, "argv",
                        ["style_postprocess.py", str(src), str(out),
                         "--template-profile", str(profile_path)])
    sp.main()

    out_doc = Document(str(out))
    quote_para = next(p for p in out_doc.paragraphs if p.style.name == "Quote")
    assert quote_para.runs[0].font.name == "Georgia"
    body_para = next(p for p in out_doc.paragraphs
                     if p.text.startswith("A plain body"))
    assert body_para.runs[0].font.name == KNOWN_MINOR_FONT


# ---------- body_style detection + consumption (issue #101) ----------

def _add_custom_style(doc, name, font_name, based_on="Normal"):
    from docx.enum.style import WD_STYLE_TYPE
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles[based_on]
    style.font.name = font_name
    return style


def test_derive_profile_detects_body_style_named_body(tmp_path):
    """A template defining its own "Body" paragraph style with a font distinct
    from Normal (the real-world shape: an institutional template author adds a
    dedicated body style rather than styling Normal directly) is detected as
    the template's body_style."""
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    _add_custom_style(doc, "Body", "Calibri")
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)

    assert dp.body_style == "Body"
    assert dp.style_fonts.get("Body") == "Calibri"


def test_derive_profile_falls_back_to_body_text(tmp_path):
    """No style literally named "Body", but Word's own built-in "Body Text"
    style carries a distinct font -- still detected."""
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    doc.styles["Body Text"].font.name = "Calibri"
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)

    assert dp.body_style == "Body Text"


def test_derive_profile_prefers_body_over_body_text(tmp_path):
    """If a template somehow defines distinct fonts for BOTH "Body" and "Body
    Text", "Body" wins -- it's the more specific, deliberately-authored name
    seen in real institutional templates; "Body Text" is Word's generic
    built-in gallery entry."""
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    _add_custom_style(doc, "Body", "Calibri")
    doc.styles["Body Text"].font.name = "Arial"
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)

    assert dp.body_style == "Body"


def test_derive_profile_no_body_style_when_absent(tmp_path):
    """No "Body" or "Body Text" style with a distinguishing font: body_style
    is None, exactly the pre-#101 behaviour (no remap, unchanged output)."""
    template = _build_branded_template(tmp_path)
    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)

    assert dp.body_style is None


def test_render_profile_yaml_emits_body_style(tmp_path):
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    _add_custom_style(doc, "Body", "Calibri")
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert 'body_style: "Body"' in yaml_text


def test_render_profile_yaml_notes_no_body_style_when_absent(tmp_path):
    template = _build_branded_template(tmp_path)
    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    yaml_text = ti.render_profile_yaml(prov, dp)

    assert "body_style:" not in yaml_text.replace("# body_style:", "")
    assert "not derivable" in yaml_text


def test_body_style_round_trip_through_consumer(tmp_path, monkeypatch):
    """End-to-end: a template with its own "Body" style (Calibri, distinct from
    Normal's global font) derives body_style="Body", and style_postprocess then
    remaps an ordinary Normal-styled body paragraph to it -- the paragraph's
    style NAME changes AND its run carries no direct-formatting override
    (issue #98's existing custom-style-preservation logic takes over once the
    remap happens), so it resolves Body's own Calibri via pure style
    inheritance rather than the house/global font (issue #101)."""
    import importlib

    from docstyle import style_postprocess as sp

    importlib.reload(sp)

    template = _build_branded_template(tmp_path, name="body-style-template.docx")
    doc = Document(str(template))
    _add_custom_style(doc, "Body", "Calibri")
    doc.save(str(template))

    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    assert dp.body_style == "Body"

    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    profile_path = tmp_path / "template-profile.yaml"
    profile_path.write_text(ti.render_profile_yaml(prov, dp), encoding="utf-8")

    # Base the consumer doc on a COPY of the template (not a bare Document()):
    # this is what pandoc's own --reference-doc mechanism actually does --
    # carries the reference doc's styles.xml, including the custom "Body"
    # style, into the rendered output. A bare Document() would only have
    # python-docx's own built-ins (fine for "Quote" in the sibling test above,
    # not fine for a style this test itself just defined on the template).
    src = tmp_path / "consumer-in.docx"
    shutil.copy(str(template), str(src))
    consumer_doc = Document(str(src))
    consumer_doc.add_heading("Body Section", level=1)
    consumer_doc.add_paragraph("An ordinary Normal-styled body paragraph.")
    consumer_doc.save(str(src))

    out = tmp_path / "consumer-out.docx"
    monkeypatch.setattr(sys, "argv",
                        ["style_postprocess.py", str(src), str(out),
                         "--template-profile", str(profile_path)])
    sp.main()

    out_doc = Document(str(out))
    body_para = next(p for p in out_doc.paragraphs
                     if p.text.startswith("An ordinary Normal"))
    assert body_para.style.name == "Body"
    assert body_para.runs[0].font.name is None  # no direct override -> pure style inheritance


def test_no_body_style_leaves_normal_paragraphs_unchanged(tmp_path, monkeypatch):
    """No body_style key in the profile (the common case: most templates don't
    define a distinct body style) -- Normal-styled paragraphs keep getting the
    house/global font exactly as before this feature existed, proving the
    change is a no-op by default."""
    import importlib

    from docstyle import style_postprocess as sp

    importlib.reload(sp)

    template = _build_branded_template(tmp_path, name="no-body-style-template.docx")
    theme = read_theme(template)
    doc2 = Document(str(template))
    dp = ti.derive_profile(doc2, theme)
    assert dp.body_style is None

    prov = ti.build_import_provenance(template, date_str="2026-07-10")
    profile_path = tmp_path / "template-profile.yaml"
    profile_path.write_text(ti.render_profile_yaml(prov, dp), encoding="utf-8")

    consumer_doc = Document()
    consumer_doc.add_paragraph("An ordinary Normal-styled body paragraph.")
    src = tmp_path / "consumer-in.docx"
    consumer_doc.save(str(src))

    out = tmp_path / "consumer-out.docx"
    monkeypatch.setattr(sys, "argv",
                        ["style_postprocess.py", str(src), str(out),
                         "--template-profile", str(profile_path)])
    sp.main()

    out_doc = Document(str(out))
    body_para = next(p for p in out_doc.paragraphs
                     if p.text.startswith("An ordinary Normal"))
    assert body_para.style.name == "Normal"
    assert body_para.runs[0].font.name == KNOWN_MINOR_FONT


def test_main_writes_profile_with_copy_reference(tmp_path):
    template = _build_branded_template(tmp_path)
    out_dir = tmp_path / "skin"
    rc = ti.main([str(template), "--out-dir", str(out_dir), "--copy-reference",
                  "--date", "2026-07-03"])
    assert rc == 0
    profile_path = out_dir / "template-profile.yaml"
    assert profile_path.exists()
    assert (out_dir / "reference.docx").exists()
    text = profile_path.read_text(encoding="utf-8")
    assert f'accent: "{KNOWN_ACCENT1}"' in text


# ---------- (c) section geometry mapping ----------

def test_section_geometry_uniform_margin(tmp_path):
    template = _build_branded_template(tmp_path)
    doc = Document(str(template))
    geom = ti.section_geometry_cm(doc.sections[0])
    assert geom["page_width_cm"] == 21.0
    assert geom["page_height_cm"] == 29.7
    assert geom["margin_cm"] == 2.0


def test_section_geometry_asymmetric_margin_not_derived(tmp_path):
    path = tmp_path / "asym.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)   # deliberately different
    sec.right_margin = Cm(2.0)
    doc.add_paragraph("body")
    doc.save(str(path))

    theme = Theme(colors={}, fonts={})
    doc2 = Document(str(path))
    dp = ti.derive_profile(doc2, theme)
    assert "margin_cm" not in dp.derived
    assert "margin_cm" in dp.not_derived
    assert "not uniform" in dp.not_derived["margin_cm"]


# ---------- (d) --check comparison function, unit-tested directly ----------

def test_compare_properties_pass_and_drift():
    expected = {"font": "Arial", "accent": "1F3864", "body": "262626",
               "page_width_cm": 21.0, "page_height_cm": 29.7, "margin_cm": 2.0}
    actual_matching = dict(expected)
    rows = ti.compare_properties(expected, actual_matching)
    assert all(r.status == "PASS" for r in rows)
    assert {r.key for r in rows} == set(expected)

    actual_drifted = dict(expected)
    actual_drifted["accent"] = "000000"
    actual_drifted["page_width_cm"] = 25.0
    rows2 = ti.compare_properties(expected, actual_drifted)
    by_key = {r.key: r.status for r in rows2}
    assert by_key["accent"] == "DRIFT"
    assert by_key["page_width_cm"] == "DRIFT"
    assert by_key["font"] == "PASS"
    assert by_key["body"] == "PASS"


def test_compare_properties_skips_non_derivable_expected():
    expected = {"font": "Arial", "accent": None}  # accent not derivable from the template
    actual = {"font": "Arial", "accent": "ANYTHING"}
    rows = ti.compare_properties(expected, actual)
    by_key = {r.key: r.status for r in rows}
    assert by_key["accent"] == "SKIP"
    assert by_key["font"] == "PASS"


def test_format_comparison_table_reads_cleanly():
    rows = ti.compare_properties(
        {"font": "Arial", "accent": None},
        {"font": "Arial", "accent": "X"},
    )
    table = ti.format_comparison_table(rows)
    assert "font" in table and "PASS" in table
    assert "(not derivable)" in table and "SKIP" in table


# ---------- (e) real --check integration through render-doc.sh ----------

_PANDOC_AND_BASH_AVAILABLE = shutil.which("pandoc") is not None and shutil.which("bash") is not None


@pytest.mark.skipif(not _PANDOC_AND_BASH_AVAILABLE, reason="requires pandoc and bash on PATH")
def test_check_gate_clean_render_through_real_pipeline(tmp_path, capsys):
    template = _build_branded_template(tmp_path)
    probe = tmp_path / "probe.md"
    probe.write_text(
        "---\n"
        "title: Probe Document\n"
        "version: v1\n"
        "---\n\n"
        "# Part I: Probe Heading\n\n"
        "This is a probe body paragraph long enough to exercise the justify/style "
        "rules the same way a real document's body text would.\n",
        encoding="utf-8",
    )
    out_dir = tmp_path / "skin"
    rc = ti.main([str(template), "--out-dir", str(out_dir), "--check", str(probe)])
    out = capsys.readouterr().out
    assert rc == 0, out
    assert "PASS" in out
    assert "DRIFT" not in out


# ---------- (f) real subprocess dispatch through render.py ----------

def test_render_py_dispatches_import_template(tmp_path):
    template = _build_branded_template(tmp_path)
    out_dir = tmp_path / "skin"
    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "import-template", str(template),
         "--out-dir", str(out_dir), "--date", "2026-07-03"],
        capture_output=True, text=True, timeout=60,
    )
    assert result.returncode == 0, result.stdout + result.stderr
    profile_path = out_dir / "template-profile.yaml"
    assert profile_path.exists()
    text = profile_path.read_text(encoding="utf-8")
    assert f'accent: "{KNOWN_ACCENT1}"' in text
    assert "TEMPLATE_DOCX=" in result.stdout
    assert "TEMPLATE_PROFILE=" in result.stdout


# ---------- (g) guidance-doc scan (issue #100) ----------

def _guidance_docx(tmp_path, name="guidance.docx") -> Path:
    doc = Document()
    doc.add_heading("Scope", level=1)
    doc.add_paragraph("This section explains what the template covers.")
    doc.add_heading("Out of scope", level=1)
    doc.add_paragraph("This section explains what is explicitly excluded.")
    doc.add_paragraph("A second paragraph under the same heading.")
    path = tmp_path / name
    doc.save(str(path))
    return path


def test_scan_guidance_doc_docx_counts_headings_and_paragraphs(tmp_path):
    path = _guidance_docx(tmp_path)
    scan = ti.scan_guidance_doc(path)
    assert scan.heading_count == 2
    assert scan.paragraph_count == 3
    assert scan.headings == ["Scope", "Out of scope"]
    assert scan.headings_truncated is False


def test_scan_guidance_doc_docx_truncates_heading_preview(tmp_path):
    doc = Document()
    for i in range(15):
        doc.add_heading(f"Section {i}", level=1)
        doc.add_paragraph("body")
    path = tmp_path / "many-headings.docx"
    doc.save(str(path))

    scan = ti.scan_guidance_doc(path, heading_preview_cap=12)
    assert scan.heading_count == 15
    assert len(scan.headings) == 12
    assert scan.headings_truncated is True


def test_scan_guidance_doc_markdown_counts_headings_and_paragraphs(tmp_path):
    path = tmp_path / "guidance.md"
    path.write_text(
        "# Scope\n\nThis section explains what the template covers.\n\n"
        "## Out of scope\n\nThis section explains what is excluded.\n\n"
        "Another standalone paragraph with no heading of its own.\n",
        encoding="utf-8",
    )
    scan = ti.scan_guidance_doc(path)
    assert scan.heading_count == 2
    assert scan.headings == ["Scope", "Out of scope"]
    # 3 body blocks: one under Scope, one under Out of scope, plus the trailing
    # standalone paragraph -- blank-line-separated blocks, not sentences.
    assert scan.paragraph_count == 3


def test_scan_guidance_doc_txt_uses_same_atx_heading_convention(tmp_path):
    path = tmp_path / "guidance.txt"
    path.write_text("# Purpose\n\nWhy this template exists.\n", encoding="utf-8")
    scan = ti.scan_guidance_doc(path)
    assert scan.heading_count == 1
    assert scan.headings == ["Purpose"]
    assert scan.paragraph_count == 1


def test_scan_guidance_doc_raises_on_unsupported_extension(tmp_path):
    path = tmp_path / "guidance.pdf"
    path.write_bytes(b"%PDF-1.4 not a real pdf")
    with pytest.raises(ti.GuidanceScanError, match="unsupported"):
        ti.scan_guidance_doc(path)


def test_format_guidance_scan_report_includes_counts_and_headings():
    scan = ti.GuidanceScan(
        path=Path("guidance.docx"), heading_count=2, paragraph_count=3,
        headings=["Scope", "Out of scope"], headings_truncated=False,
    )
    report = ti.format_guidance_scan_report(scan)
    assert "2 section heading(s), 3 paragraph(s)" in report
    assert "Scope; Out of scope" in report
    assert "editorial-doctrine.yaml" in report


def test_main_with_guidance_doc_prints_scan_report(tmp_path, capsys):
    template = _build_branded_template(tmp_path)
    guidance = _guidance_docx(tmp_path)
    rc = ti.main([str(template), "--out-dir", str(tmp_path / "skin"),
                  "--date", "2026-07-03", "--guidance-doc", str(guidance)])
    assert rc == 0
    out = capsys.readouterr().out
    assert "Guidance-doc scan:" in out
    assert "2 section heading(s), 3 paragraph(s)" in out
    assert "No --guidance-doc given" not in out


def test_main_without_guidance_doc_prints_reminder(tmp_path, capsys):
    template = _build_branded_template(tmp_path)
    rc = ti.main([str(template), "--out-dir", str(tmp_path / "skin"), "--date", "2026-07-03"])
    assert rc == 0
    out = capsys.readouterr().out
    assert "No --guidance-doc given" in out
    assert "--guidance-doc <path>" in out
    assert "Guidance-doc scan:" not in out


def test_main_with_missing_guidance_doc_path_errors(tmp_path, capsys):
    template = _build_branded_template(tmp_path)
    rc = ti.main([str(template), "--out-dir", str(tmp_path / "skin"), "--date", "2026-07-03",
                  "--guidance-doc", str(tmp_path / "does-not-exist.docx")])
    assert rc == 2
    err = capsys.readouterr().err
    assert "ERROR" in err and "does-not-exist.docx" in err


def test_main_with_unsupported_guidance_doc_type_errors(tmp_path, capsys):
    template = _build_branded_template(tmp_path)
    bad = tmp_path / "guidance.pdf"
    bad.write_bytes(b"%PDF-1.4 not a real pdf")
    rc = ti.main([str(template), "--out-dir", str(tmp_path / "skin"), "--date", "2026-07-03",
                  "--guidance-doc", str(bad)])
    assert rc == 2
    err = capsys.readouterr().err
    assert "unsupported" in err

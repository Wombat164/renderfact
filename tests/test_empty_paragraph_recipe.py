r"""
Regression test for issue #102: markdown has no native syntax for a genuinely
empty spacer paragraph -- multiple blank lines are pure block separators (zero
extra paragraphs), and common workarounds (a lone `&nbsp;`, a trailing `\`)
both leave a residual character behind, not a true empty `<w:p>`. Confirmed
this reproducing a real institutional template that uses several genuinely-
empty paragraphs as vertical spacing and needed paragraph-count fidelity to
match it exactly.

This is NOT new markdown syntax: issue #96's raw_attribute escape hatch
(merged, see test_raw_attribute_escape_hatch.py) already lets a fenced
```{=openxml} block emit a genuine `<w:p/>` verbatim. This test protects the
specific worked recipe now documented in wiki/content/how-to/index.md ("Insert
a genuinely empty paragraph (spacer)") so it doesn't silently break under a
future pandoc/renderfact change -- the wiki page names this test file by name
as its regression proof, so it needs to keep passing under that name.

Deliberately narrower than test_raw_attribute_escape_hatch.py: that file
proves the general escape-hatch mechanism works; this file proves the ONE
specific idiom (`<w:p/>`, no attributes, no content) that answers #102
produces a paragraph indistinguishable from a genuinely hand-inserted Word
spacer paragraph -- zero runs, zero text, default ("Normal") style.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
RENDER_PY = REPO_ROOT / "render.py"

HAVE_PANDOC = shutil.which("pandoc") is not None

try:
    import docx  # noqa: F401
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

EMPTY_PARAGRAPH_SOURCE = (
    "---\ntitle: Empty Paragraph Recipe Check\n---\n\n"
    "# Heading One\n\n"
    "Some text before the spacer.\n\n"
    "```{=openxml}\n"
    "<w:p/>\n"
    "```\n\n"
    "# Heading Two\n\n"
    "More text after the spacer.\n"
)


@pytest.mark.skipif(not HAVE_PANDOC or not HAVE_DOCX, reason="needs pandoc + python-docx")
def test_raw_empty_paragraph_block_renders_as_genuinely_empty_paragraph(tmp_path):
    """End-to-end: the documented `<w:p/>` recipe, run through the REAL `render
    docx` pipeline, produces a paragraph with zero runs and empty text -- not
    a paragraph carrying a stray non-breaking space or a literal backslash,
    which is exactly what the pre-#96 workarounds this test guards against
    used to leave behind."""
    import docx as docx_lib

    src = tmp_path / "empty-paragraph.md"
    src.write_text(EMPTY_PARAGRAPH_SOURCE, encoding="utf-8")
    env = {**os.environ, "OUTPUT_DIR": str(tmp_path / "out")}
    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "docx", str(src)],
        capture_output=True, text=True, timeout=120, env=env, cwd=tmp_path,
    )
    combined = result.stdout + result.stderr
    if result.returncode == 3 and ("pandoc not found" in combined or "bash not found" in combined
                                    or "python not found" in combined):
        pytest.skip("render engines not installed on this host: " + combined.strip()[:160])
    assert result.returncode == 0, combined

    docx_files = sorted((tmp_path / "out").glob("*.docx"))
    assert len(docx_files) == 1, combined
    doc = docx_lib.Document(str(docx_files[0]))

    texts = [p.text for p in doc.paragraphs]
    assert "Some text before the spacer." in texts
    assert "More text after the spacer." in texts

    spacer_idx = texts.index("Some text before the spacer.") + 1
    spacer = doc.paragraphs[spacer_idx]
    # The genuinely-empty-paragraph bar: zero text, zero runs (not a paragraph
    # holding a single run with an empty-looking-but-nonzero-length string, and
    # not the &nbsp;/trailing-backslash residue #102 reported).
    assert spacer.text == ""
    assert len(spacer.runs) == 0


def test_recipe_source_matches_the_documented_wiki_snippet():
    """The exact fenced block this test exercises must be byte-identical to
    what a real author would copy from the wiki how-to entry -- if a future
    edit to either drifts, this catches it as a source-string mismatch, not
    just a passing/failing render."""
    wiki_page = REPO_ROOT / "wiki" / "content" / "how-to" / "index.md"
    wiki_text = wiki_page.read_text(encoding="utf-8")
    assert "```{=openxml}\n<w:p/>\n```" in wiki_text

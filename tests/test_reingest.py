"""
Tests for roundtrip/reingest.py (D11 part 3a, chunk 4.4) and the source_commit
provenance hardening (D11 part 4).

Covers: provenance verdicts (FAST_FORWARD, DIVERGED, UID mismatch, artifact
without provenance, source without a UID: all fail closed with clean messages);
extraction of Word comments and tracked changes from a hand-built minimal OOXML
zip (the extractor reads raw XML, so the fixture does not need to be a full
valid docx); structure walking and the normalized delta; the fast-forward plan
discipline (plain unique lines apply, inline-markup / duplicate / heading /
add-delete cases defer to manual); apply preserving leading list markers; the
render.py CLI end to end including the DIVERGED --apply refusal; and
source_commit (sha in a git repo, -dirty suffix, None outside a repo, and
legacy payloads without the field extracting cleanly).
"""

from __future__ import annotations

import json
import subprocess
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(REPO_ROOT / "roundtrip"))

import reingest  # noqa: E402
from roundtrip import provenance  # noqa: E402

RENDER_PY = REPO_ROOT / "render.py"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _make_source(tmp_path: Path, body: str = "The plan starts in March.\n") -> Path:
    src = tmp_path / "doc.md"
    src.write_text(f"---\ntitle: T\n---\n\n# Overview\n\n{body}", encoding="utf-8")
    return src


def _make_rendered_docx(tmp_path: Path, source: Path, paragraphs: list[str]) -> Path:
    doc = Document()
    doc.add_heading("Overview", level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    path = tmp_path / "rendered.docx"
    doc.save(str(path))
    provenance.embed(path, provenance.build_provenance(source))
    return path


# ---- provenance verdicts ----

def test_fast_forward_when_source_unchanged(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    prov, verdict = reingest.check_provenance(art, src)
    assert verdict == "FAST_FORWARD"
    assert prov.source_uid


def test_diverged_when_source_evolved(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    src.write_text(src.read_text(encoding="utf-8") + "\nNew paragraph.\n", encoding="utf-8")
    _prov, verdict = reingest.check_provenance(art, src)
    assert verdict == "DIVERGED"


def test_uid_mismatch_fails_closed(tmp_path):
    src_a = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src_a, ["x"])
    src_b = tmp_path / "other.md"
    src_b.write_text("---\ntitle: O\n---\n\nBody.\n", encoding="utf-8")
    provenance.build_provenance(src_b)  # gives other.md its own uid
    with pytest.raises(reingest.ReingestError, match="UID mismatch"):
        reingest.check_provenance(art, src_b)


def test_artifact_without_provenance_fails_closed(tmp_path):
    src = _make_source(tmp_path)
    provenance.build_provenance(src)
    doc = Document()
    doc.add_paragraph("x")
    bare = tmp_path / "bare.docx"
    doc.save(str(bare))
    with pytest.raises(reingest.ReingestError, match="no renderfact provenance"):
        reingest.check_provenance(bare, src)


def test_source_without_uid_fails_closed(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["x"])
    stranger = tmp_path / "stranger.md"
    stranger.write_text("---\ntitle: S\n---\n\nBody.\n", encoding="utf-8")
    with pytest.raises(reingest.ReingestError, match="no renderfact_uid"):
        reingest.check_provenance(art, stranger)


# ---- comments + tracked changes (hand-built OOXML: the extractor reads raw XML) ----

_DOC_XML = f"""<?xml version="1.0"?>
<w:document xmlns:w="{W_NS}"><w:body>
  <w:p><w:r><w:t>Kept text.</w:t></w:r></w:p>
  <w:p><w:ins w:author="Reviewer A"><w:r><w:t>inserted words</w:t></w:r></w:ins></w:p>
  <w:p><w:del w:author="Reviewer B"><w:r><w:delText>removed words</w:delText></w:r></w:del></w:p>
</w:body></w:document>"""

_COMMENTS_XML = f"""<?xml version="1.0"?>
<w:comments xmlns:w="{W_NS}">
  <w:comment w:id="1" w:author="Reviewer A" w:date="2026-07-04T10:00:00Z">
    <w:p><w:r><w:t>Please confirm the date.</w:t></w:r></w:p>
  </w:comment>
</w:comments>"""


def _minimal_ooxml(tmp_path: Path) -> Path:
    p = tmp_path / "tracked.docx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("word/document.xml", _DOC_XML)
        z.writestr("word/comments.xml", _COMMENTS_XML)
    return p


def test_extract_comments_and_tracked_changes(tmp_path):
    p = _minimal_ooxml(tmp_path)
    z = zipfile.ZipFile(p)
    comments = reingest.extract_comments(z)
    assert comments == [{"id": "1", "author": "Reviewer A",
                         "date": "2026-07-04T10:00:00Z", "text": "Please confirm the date."}]
    root = ET.fromstring(z.read("word/document.xml"))
    ins, dele = reingest.extract_tracked(root)
    assert ins == [("Reviewer A", "inserted words")]
    assert dele == [("Reviewer B", "removed words")]


def test_embedded_objects_are_inventoried_not_ignored(tmp_path):
    p = tmp_path / "with-embed.docx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("word/document.xml", _DOC_XML)
        z.writestr("word/embeddings/oleObject1.xlsx", b"PK-fake-workbook-bytes")
        z.writestr("word/embeddings/blob.bin", b"\x00\x01")
    z = zipfile.ZipFile(p)
    embedded = reingest.extract_embedded(z)
    assert {e["name"] for e in embedded} == {"oleObject1.xlsx", "blob.bin"}
    xlsx = next(e for e in embedded if e["kind"] == "xlsx")
    assert xlsx["bytes"] == len(b"PK-fake-workbook-bytes")
    # OOXML suffix without provenance = unknown-origin, never silently "foreign"
    assert xlsx["status"] == "unknown-origin"
    assert "adopt" in xlsx["note"]


def _zip_with_embedding(tmp_path: Path, name: str, payload: bytes) -> zipfile.ZipFile:
    p = tmp_path / "host.docx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("word/document.xml", _DOC_XML)
        z.writestr(f"word/embeddings/{name}", payload)
    return zipfile.ZipFile(p)


def test_embedded_renderfact_tracked_doc_is_recognised(tmp_path):
    inner_src = tmp_path / "inner.md"
    inner_src.write_text("---\ntitle: Inner\n---\n\nBody.\n", encoding="utf-8")
    inner = tmp_path / "inner.docx"
    doc = Document()
    doc.add_paragraph("Embedded report body.")
    doc.save(str(inner))
    provenance.embed(inner, provenance.build_provenance(inner_src))

    z = _zip_with_embedding(tmp_path, "report.docx", inner.read_bytes())
    entry = reingest.extract_embedded(z)[0]
    assert entry["status"] == "renderfact-tracked"
    assert entry["provenance"]["source_uid"]
    assert "its own per-format path" in entry["note"]


def test_embedded_vsdx_shaped_opc_is_provenance_readable(tmp_path):
    """The generic OPC reader must cover .vsdx bytes BEFORE the full vsdx
    adapter (C8.2) lands: docProps/core.xml is shared across all OPC formats."""
    import io

    core = ('<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/'
            'package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/">'
            '<dc:identifier>renderfact:v1:{"source_uid": "vsdx-uid", '
            '"source_version": "v", "rendered_at": "t", "tool_version": "x"}'
            '</dc:identifier></cp:coreProperties>')
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as inner:
        inner.writestr("docProps/core.xml", core)
        inner.writestr("visio/document.xml", "<x/>")
    z = _zip_with_embedding(tmp_path, "drawing.vsdx", buf.getvalue())
    entry = reingest.extract_embedded(z)[0]
    assert entry["status"] == "renderfact-tracked"
    assert entry["provenance"]["source_uid"] == "vsdx-uid"


def test_embedded_foreign_type_gets_markitdown_preview(tmp_path, monkeypatch):
    class FakeResult:
        text_content = "Extracted PDF text for the preview."

    class FakeMarkItDown:
        def convert_stream(self, stream, file_extension):
            assert file_extension == ".pdf"
            return FakeResult()

    import types
    fake = types.ModuleType("markitdown")
    fake.MarkItDown = FakeMarkItDown
    monkeypatch.setitem(sys.modules, "markitdown", fake)

    z = _zip_with_embedding(tmp_path, "annex.pdf", b"%PDF-1.4 fake")
    entry = reingest.extract_embedded(z)[0]
    assert entry["status"] == "foreign"
    assert entry["preview"].startswith("Extracted PDF text")


def test_real_markitdown_unsupported_format_is_contained(tmp_path):
    """markitdown's UnsupportedFormatException does NOT inherit from Exception
    (real-library quirk that crashed the report before this regression test):
    an unconvertible type must degrade to a note, never an escape."""
    pytest.importorskip("markitdown")
    preview, note = reingest._markitdown_preview(b"\x00\x01", "bin")
    assert preview is None
    assert "could not convert" in note


def test_embeddings_recognised_in_non_docx_hosts(tmp_path):
    """XLSX (xl/embeddings/), PPTX (ppt/embeddings/) and VSDX hosts carry
    embedded objects too: the matcher keys on the path segment."""
    p = tmp_path / "host.xlsx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("xl/embeddings/inner.docx", b"not-a-zip")
        z.writestr("ppt/embeddings/deck-annex.pdf", b"%PDF-1.4")
    entries = reingest.extract_embedded(zipfile.ZipFile(p))
    assert {e["name"] for e in entries} == {"inner.docx", "deck-annex.pdf"}


def test_embedded_foreign_type_notes_missing_markitdown(tmp_path, monkeypatch):
    monkeypatch.setitem(sys.modules, "markitdown", None)  # forces ImportError
    z = _zip_with_embedding(tmp_path, "annex.pdf", b"%PDF-1.4 fake")
    entry = reingest.extract_embedded(z)[0]
    assert entry["status"] == "foreign"
    assert entry["preview"] is None
    assert "pip install markitdown" in entry["note"]


# ---- fast-forward plan discipline ----

def _plan(md_text: str, edited_lines: list[str]):
    md_lines = [x for x in (reingest._norm(y) for y in reingest.md_plaintext(md_text)) if x]
    dx_raw = edited_lines
    dx_lines = [x for x in (reingest._norm(y) for y in dx_raw) if x]
    return reingest.plan_fast_forward(md_text, md_lines, dx_lines, dx_raw)


def test_plain_unique_rewording_is_safe():
    md = "---\ntitle: T\n---\n\n# H\n\nThe plan starts in March.\n"
    safe, manual = _plan(md, ["#> H", "The plan starts in May."])
    assert len(safe) == 1
    assert safe[0][2] == "The plan starts in May."
    assert manual == []


def test_inline_markup_line_defers_to_manual():
    md = "---\ntitle: T\n---\n\n# H\n\nThe **bold** plan starts in March.\n"
    safe, manual = _plan(md, ["#> H", "The bold plan starts in May."])
    assert safe == []
    assert any("inline markup" in (m[2] if len(m) > 2 else "") for m in manual)


def test_duplicate_normalized_line_defers_to_manual():
    md = "---\ntitle: T\n---\n\n# H\n\nSame line.\n\nOther.\n\nSame line.\n"
    safe, manual = _plan(md, ["#> H", "Same line.", "Other.", "Changed line."])
    assert safe == []
    assert manual  # ambiguity is never auto-applied


def test_heading_edits_defer_to_manual():
    md = "---\ntitle: T\n---\n\n# Old Heading\n\nBody.\n"
    safe, manual = _plan(md, ["#> New Heading", "Body."])
    assert safe == []
    assert any("heading" in (m[2] if len(m) > 2 else "") for m in manual)


def test_in_place_replacement_is_a_safe_rewording():
    """An equal-length replace opcode IS the mechanical rewording case, even when
    the whole paragraph changed: same position, plain text, unique in the source."""
    md = "---\ntitle: T\n---\n\n# H\n\nKept.\n\nDoomed.\n"
    safe, manual = _plan(md, ["#> H", "Kept.", "Brand new paragraph."])
    assert len(safe) == 1 and safe[0][2] == "Brand new paragraph."


def test_pure_additions_and_deletions_go_to_manual():
    md_add = "---\ntitle: T\n---\n\n# H\n\nKept.\n"
    safe, manual = _plan(md_add, ["#> H", "Kept.", "Brand new paragraph."])
    assert safe == []
    assert ("(added in the edited DOCX)", "Brand new paragraph.") in manual

    md_del = "---\ntitle: T\n---\n\n# H\n\nKept.\n\nDoomed.\n"
    safe, manual = _plan(md_del, ["#> H", "Kept."])
    assert safe == []
    assert ("Doomed.", "(deleted in the edited DOCX)") in manual


def test_apply_preserves_leading_list_marker(tmp_path):
    src = tmp_path / "doc.md"
    src.write_text("---\ntitle: T\n---\n\n# H\n\n- The plan starts in March.\n", encoding="utf-8")
    md_text = src.read_text(encoding="utf-8")
    safe, _ = _plan(md_text, ["#> H", "The plan starts in May."])
    assert len(safe) == 1
    reingest.apply_fast_forward(src, safe)
    assert "- The plan starts in May.\n" in src.read_text(encoding="utf-8")


def test_apply_preserves_leading_ordered_list_marker(tmp_path):
    """Regression guard (issue #72): ordered-list markers ('1.', '2.', ...) are
    NOT stripped by _norm (unlike bullet markers): the fast-forward planner
    already handles this correctly because _norm_source_line() and the diff's
    own normalization agree on keeping the digit marker, so the two sides still
    match uniquely; _LEADING_MARKER then strips it only for the write-back.
    Must keep working exactly as-is: this is the "already handled" case from
    the issue, not something this PR's fix touches."""
    src = tmp_path / "doc.md"
    src.write_text("---\ntitle: T\n---\n\n# H\n\n1. The plan starts in March.\n", encoding="utf-8")
    md_text = src.read_text(encoding="utf-8")
    safe, manual = _plan(md_text, ["#> H", "The plan starts in May."])
    assert len(safe) == 1 and manual == []
    reingest.apply_fast_forward(src, safe)
    assert "1. The plan starts in May.\n" in src.read_text(encoding="utf-8")


# ---- structural-noise stripping (issue #72) ----

def test_fenced_div_lines_produce_no_text(tmp_path):
    md = ('---\ntitle: T\n---\n\n::: {custom-style="Title"}\n# Overview\n:::\n\n'
         "Body paragraph.\n")
    lines = [x for x in (reingest._norm(y) for y in reingest.md_plaintext(md)) if x]
    assert lines == ["#> Overview", "Body paragraph."]
    assert not any(":::" in ln for ln in lines)


def test_raw_attribute_block_is_dropped_whole(tmp_path):
    md = ('---\ntitle: T\n---\n\nBefore.\n\n```{=openxml}\n'
         '<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\nAfter.\n')
    lines = [x for x in (reingest._norm(y) for y in reingest.md_plaintext(md)) if x]
    assert lines == ["Before.", "After."]
    assert not any("openxml" in ln or "w:br" in ln for ln in lines)


def test_blockquote_marker_is_dequoted_not_dropped():
    md = '---\ntitle: T\n---\n\n> A quoted tagline.\n'
    lines = [x for x in (reingest._norm(y) for y in reingest.md_plaintext(md)) if x]
    assert lines == ["A quoted tagline."]


def _make_structural_noise_source(tmp_path: Path, tagline: str = "The plan starts in March.") -> Path:
    src = tmp_path / "doc.md"
    src.write_text(
        '---\ntitle: T\n---\n\n::: {custom-style="Title"}\n# Overview\n:::\n\n'
        '```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n'
        f"> A quoted tagline.\n\n{tagline}\n",
        encoding="utf-8",
    )
    return src


def _make_structural_noise_docx(tmp_path: Path, source: Path, tagline: str) -> Path:
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("A quoted tagline.")
    doc.add_paragraph(tagline)
    path = tmp_path / "rendered.docx"
    doc.save(str(path))
    provenance.embed(path, provenance.build_provenance(source))
    return path


def test_reingest_report_has_no_structural_false_deletions(tmp_path):
    """End-to-end (issue #72 repro shape): a source using fenced-divs, a
    raw-attribute page-break block, and a blockquote, re-ingested against a
    DOCX that matches it exactly (no reviewer edit at all). Before the fix,
    '{=openxml}', the raw XML line, ':::', and '> ...' all showed up as
    spurious deletions. After the fix the delta must be completely empty."""
    src = _make_structural_noise_source(tmp_path)
    art = _make_structural_noise_docx(tmp_path, src, "The plan starts in March.")

    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["verdict"] == "FAST_FORWARD"
    assert payload["delta"] == []
    assert payload["manual"] == []
    assert payload["safe_edits"] == []


def test_genuine_edit_near_structural_noise_still_surfaces(tmp_path):
    """A real reviewer edit sitting right next to fenced-div/raw-attribute/
    blockquote noise must still be caught, and caught ALONE: no adjacent
    structural noise should also show up in the delta."""
    src = _make_structural_noise_source(tmp_path)
    art = _make_structural_noise_docx(tmp_path, src, "The plan starts in March.")

    doc = Document(str(art))
    for p in doc.paragraphs:
        if p.text == "The plan starts in March.":
            p.runs[0].text = "The plan starts in May."
    doc.save(str(art))

    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["verdict"] == "FAST_FORWARD"
    assert payload["delta"] == ["-The plan starts in March.", "+The plan starts in May."]
    assert len(payload["safe_edits"]) == 1
    edit = payload["safe_edits"][0]
    assert edit["old"] == "The plan starts in March."
    assert edit["new"] == "The plan starts in May."
    assert payload["manual"] == []


def test_strip_pattern_knob_removes_project_specific_sigil(tmp_path):
    """The --strip-pattern knob (issue #72) covers project-specific structural
    conventions renderfact itself has no reason to special-case, e.g. a custom
    heading-anchor sigil like '#>' that never renders as literal DOCX text."""
    src = tmp_path / "doc.md"
    src.write_text(
        "---\ntitle: T\n---\n\n# Overview\n\n#> Aanpak (kept).\n\nBody paragraph.\n",
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Aanpak (kept).")
    doc.add_paragraph("Body paragraph.")
    art = tmp_path / "rendered.docx"
    doc.save(str(art))
    provenance.embed(art, provenance.build_provenance(src))

    without_flag = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert without_flag.returncode == 0, without_flag.stderr
    payload_before = json.loads(without_flag.stdout)
    assert "-#> Aanpak (kept)." in payload_before["delta"]

    with_flag = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src),
         "--json", "--strip-pattern", r"^#>\s*"],
        capture_output=True, text=True, timeout=120,
    )
    assert with_flag.returncode == 0, with_flag.stderr
    payload_after = json.loads(with_flag.stdout)
    assert not any("Aanpak" in d for d in payload_after["delta"])


def test_strip_pattern_invalid_regex_fails_closed(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src),
         "--strip-pattern", "[unclosed"],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 1
    assert "invalid --strip-pattern regex" in result.stderr


# ---- CLI end to end ----

def test_cli_report_apply_and_diverged_refusal(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    doc = Document(str(art))
    for p in doc.paragraphs:
        if p.text == "The plan starts in March.":
            p.runs[0].text = "The plan starts in May."
    doc.save(str(art))

    report = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert report.returncode == 0, report.stderr
    payload = json.loads(report.stdout)
    assert payload["verdict"] == "FAST_FORWARD"
    assert payload["safe_edits"][0]["new"] == "The plan starts in May."
    assert payload["applied"] is None  # report-only by default

    applied = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--apply"],
        capture_output=True, text=True, timeout=120,
    )
    assert applied.returncode == 0, applied.stderr
    assert "The plan starts in May." in src.read_text(encoding="utf-8")

    # the source moved past the render: DIVERGED now, and --apply must refuse
    refusal = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--apply"],
        capture_output=True, text=True, timeout=120,
    )
    assert refusal.returncode == 1
    assert "DIVERGED" in refusal.stderr


# ---- G8: workflow surfacing (--contextualize chaining + next-step hint) ----

def _make_rendered_docx_with_extra_paragraph(tmp_path: Path, source: Path, extra: str) -> Path:
    """A reviewer-added paragraph with no source counterpart -- lands in the
    `manual` bucket (an 'add'), not `safe`, unlike a clean 1:1 reword."""
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("The plan starts in March.")
    doc.add_paragraph(extra)
    path = tmp_path / "rendered.docx"
    doc.save(str(path))
    provenance.embed(path, provenance.build_provenance(source))
    return path


def test_cli_contextualize_chains_when_manual_residue_exists(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx_with_extra_paragraph(tmp_path, src, "A brand new reviewer paragraph.")
    log = tmp_path / "out.decisions.md"
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json",
         "--contextualize", "--decision-log", str(log)],
        capture_output=True, text=True, timeout=120,
    )
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["manual"], "fixture should have manual-review residue"
    assert "contextualize" in payload
    assert payload["contextualize"]["round"] == 1
    assert log.exists() and "## " in log.read_text(encoding="utf-8")


def test_cli_contextualize_skipped_when_nothing_needs_it(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    log = tmp_path / "out.decisions.md"
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json",
         "--contextualize", "--decision-log", str(log)],
        capture_output=True, text=True, timeout=120,
    )
    assert r.returncode == 0, r.stderr
    payload = json.loads(r.stdout)
    assert payload["contextualize"] == {"skipped": "nothing needing a decision (no manual residue, FAST_FORWARD)"}
    assert not log.exists()


def test_cli_second_reingest_contextualize_increments_round(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx_with_extra_paragraph(tmp_path, src, "Round one reviewer content.")
    log = tmp_path / "out.decisions.md"
    r1 = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json",
         "--contextualize", "--decision-log", str(log)],
        capture_output=True, text=True, timeout=120,
    )
    assert r1.returncode == 0, r1.stderr
    assert json.loads(r1.stdout)["contextualize"]["round"] == 1

    # Re-render (fresh provenance against the now-unchanged source) then add a
    # second reviewer edit, so this is a genuine second round for the SAME doc.
    art2 = _make_rendered_docx_with_extra_paragraph(tmp_path, src, "Round two reviewer content.")
    r2 = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art2), "--source", str(src), "--json",
         "--contextualize", "--decision-log", str(log)],
        capture_output=True, text=True, timeout=120,
    )
    assert r2.returncode == 0, r2.stderr
    assert json.loads(r2.stdout)["contextualize"]["round"] == 2
    assert "Round 2:" in log.read_text(encoding="utf-8")


def test_cli_next_step_hint_printed_when_manual_residue_without_contextualize_flag(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx_with_extra_paragraph(tmp_path, src, "A brand new reviewer paragraph.")
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src)],
        capture_output=True, text=True, timeout=120,
    )
    assert r.returncode == 0, r.stderr
    assert "render contextualize" in r.stdout
    assert "--contextualize" in r.stdout


def test_cli_no_next_step_hint_when_nothing_needs_contextualize(tmp_path):
    src = _make_source(tmp_path)
    art = _make_rendered_docx(tmp_path, src, ["The plan starts in March."])
    r = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src)],
        capture_output=True, text=True, timeout=120,
    )
    assert r.returncode == 0, r.stderr
    assert "render contextualize" not in r.stdout


# ---- source_commit hardening ----

def test_source_commit_in_a_git_repo(tmp_path):
    subprocess.run(["git", "init", "-q", str(tmp_path)], check=True, timeout=30)
    subprocess.run(["git", "-C", str(tmp_path), "config", "user.email", "t@example.com"],
                   check=True, timeout=30)
    subprocess.run(["git", "-C", str(tmp_path), "config", "user.name", "t"], check=True, timeout=30)
    src = _make_source(tmp_path)
    subprocess.run(["git", "-C", str(tmp_path), "add", "-A"], check=True, timeout=30)
    subprocess.run(["git", "-C", str(tmp_path), "commit", "-qm", "x"], check=True, timeout=30)

    commit = provenance.source_commit(src)
    assert commit and "-dirty" not in commit

    src.write_text(src.read_text(encoding="utf-8") + "\nEdit.\n", encoding="utf-8")
    assert provenance.source_commit(src).endswith("-dirty")


def test_source_commit_none_outside_git(tmp_path, monkeypatch):
    src = _make_source(tmp_path)
    monkeypatch.setenv("GIT_CEILING_DIRECTORIES", str(tmp_path.parent))
    monkeypatch.setenv("GIT_DIR", str(tmp_path / "nonexistent-git"))
    assert provenance.source_commit(src) is None


def test_legacy_payload_without_source_commit_extracts(tmp_path):
    doc = Document()
    doc.add_paragraph("x")
    art = tmp_path / "legacy.docx"
    doc.save(str(art))
    legacy = ('renderfact:v1:{"source_uid": "abc", "source_version": "def",'
              ' "rendered_at": "2026-07-01T00:00:00Z", "tool_version": "123"}')
    d = Document(str(art))
    d.core_properties.identifier = legacy
    d.save(str(art))
    prov = provenance.extract(art)
    assert prov.source_uid == "abc"
    assert prov.source_commit is None


# ---- table-width sidecar (--apply-widths, issue #73) ----

def _add_table(doc, header, rows, col_widths_twips):
    from docx.oxml.ns import qn as _qn

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    for ci, h in enumerate(header):
        table.rows[0].cells[ci].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = cell
    grid = table._tbl.find(_qn("w:tblGrid"))
    for gc, w in zip(grid.findall(_qn("w:gridCol")), col_widths_twips):
        gc.set(_qn("w:w"), str(w))
    return table


def test_build_table_width_entries_keys_by_header_cols_rows():
    tables = [
        {"header": ["#", "Item"], "rows": [["#", "Item"], ["1", "Widget"]],
         "twips": [227, 5443], "cm": [0.4, 9.6], "pct": [4, 96]},
    ]
    entries = reingest.build_table_width_entries(tables)
    assert entries == [{"header": ["#", "Item"], "rows": 2, "cols": 2, "widths": [227, 5443]}]


def test_render_width_spec_yaml_matches_load_table_widths_shape(tmp_path):
    """The emitted sidecar must parse via _load_table_widths() into exactly the
    'tables: [[...], ...]' shape apply_table_widths() consumes: same repo,
    same function, not a parallel incompatible format."""
    sys.path.insert(0, str(REPO_ROOT / "docstyle"))
    import style_postprocess as sp

    entries = [
        {"header": ["#", "Item"], "rows": 9, "cols": 2, "widths": [227, 5443]},
        {"header": ["A", "B", "C"], "rows": 4, "cols": 3, "widths": [1000, 1000, 1000]},
    ]
    text = reingest.render_width_spec_yaml(entries)
    assert "T1 (9 rows) [# | Item]" in text
    assert "T2 (4 rows) [A | B | C]" in text

    out = tmp_path / "widths.yaml"
    out.write_text(text, encoding="utf-8")
    specs = sp._load_table_widths(str(out))
    assert specs == [[227, 5443], [1000, 1000, 1000]]


def test_render_width_spec_yaml_empty_tables():
    text = reingest.render_width_spec_yaml([])
    assert "tables: []" in text


def test_apply_widths_end_to_end_round_trips_through_style_postprocess(tmp_path):
    """Full proof, not just 'a YAML file was written': reingest --apply-widths
    on a reviewer-widened table emits a sidecar that, fed to
    'render docstyle --table-widths', actually reproduces the reviewer's
    column-width ratio on a real render."""
    sys.path.insert(0, str(REPO_ROOT / "docstyle"))
    import style_postprocess as sp

    src = tmp_path / "doc.md"
    src.write_text(
        "---\ntitle: T\n---\n\n# Overview\n\n| # | Item |\n|---|---|\n| 1 | Widget |\n",
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("Overview", level=1)
    # reviewer narrowed column 1, widened column 2 (9:1 ratio, mirrors the issue's
    # own worked example: widths %: [4, 96])
    _add_table(doc, ["#", "Item"], [["1", "Widget"]], [227, 5443])
    art = tmp_path / "rendered.docx"
    doc.save(str(art))
    provenance.embed(art, provenance.build_provenance(src))

    widths_out = tmp_path / "widths.yaml"
    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src),
         "--apply-widths", str(widths_out)],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    assert widths_out.exists()
    assert "wrote width spec to" in result.stdout
    assert "render docstyle --table-widths" in result.stdout

    # a NEW document, standing in for "the next render" of the same source
    next_render = Document()
    _add_table(next_render, ["#", "Item"], [["1", "Widget"]], [4500, 4500])  # default-ish 1:1
    next_path = tmp_path / "next.docx"
    next_render.save(str(next_path))
    styled_out = tmp_path / "next-styled.docx"

    rc = sp.main([str(next_path), str(styled_out), "--table-widths", str(widths_out)])
    assert rc == 0

    styled = Document(str(styled_out))
    grid = styled.tables[0]._tbl.find(sp.qn("w:tblGrid"))
    final_widths = [int(gc.get(sp.qn("w:w"))) for gc in grid.findall(sp.qn("w:gridCol"))]
    text_w = sp._section_text_width_twips(styled)
    assert sum(final_widths) == text_w  # scaled to fill the section, full-width intent preserved
    # the reviewer's 227:5443 ratio (~24x) survives the scale-to-text-width step
    assert final_widths[1] / final_widths[0] == pytest.approx(5443 / 227, rel=0.02)


# ---- page-break report section (## 3b, issue #73) ----

def _pagebreak_paragraph(doc):
    from docx.oxml.ns import qn as _qn
    from lxml import etree

    p = doc.add_paragraph()
    r = p.add_run()
    br = etree.SubElement(r._r, _qn("w:br"))
    br.set(_qn("w:type"), "page")
    return p


def test_source_page_breaks_finds_newpage_and_raw_openxml_lines():
    md = (
        "---\ntitle: T\n---\n\n# H\n\nBefore.\n\n\\newpage\n\nMiddle.\n\n"
        '```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\nAfter.\n'
    )
    offsets = reingest.source_page_breaks(md)
    lines = md.split("\n")
    assert [lines[i - 1].strip() for i in offsets] == [
        "\\newpage",
        '<w:p><w:r><w:br w:type="page"/></w:r></w:p>',
    ]


def test_docx_page_breaks_finds_offset_ignores_lastrenderedpagebreak():
    from docx.oxml.ns import qn as _qn
    from lxml import etree

    doc = Document()
    doc.add_paragraph("Kept text.")  # offset 0
    _pagebreak_paragraph(doc)  # offset 1: real manual break
    p2 = doc.add_paragraph("More text.")  # offset 2
    # Word's own layout-cache marker: must NOT be counted as a deliberate break
    r2 = p2.add_run()
    etree.SubElement(r2._r, _qn("w:lastRenderedPageBreak"))

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    z = zipfile.ZipFile(buf)
    root = ET.fromstring(z.read("word/document.xml"))
    body = root.find(f"{reingest.W}body")
    offsets = reingest.docx_page_breaks(body)
    assert offsets == [1]


def test_pagebreak_added_appears_in_3b_and_not_in_generic_manual_list(tmp_path):
    """A page break the reviewer added must show up distinctly in the new
    docx_page_breaks/source_page_breaks JSON fields (the ## 3b section's data),
    and must NOT also appear in the generic manual-review noise list: the
    page-break-only paragraph carries no visible text, so it was never able to
    reach the text-delta/manual path in the first place (see docx_page_breaks()
    docstring); this test locks that behaviour."""
    src = tmp_path / "doc.md"
    src.write_text(
        "---\ntitle: T\n---\n\n# H\n\nBefore.\n\nAfter.\n",
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("H", level=1)
    doc.add_paragraph("Before.")
    doc.add_paragraph("After.")
    art = tmp_path / "rendered.docx"
    doc.save(str(art))
    provenance.embed(art, provenance.build_provenance(src))

    # reviewer inserts a manual page break between the two paragraphs
    d = Document(str(art))
    body = d.element.body
    new_p = body.makeelement(f"{reingest.W}p", {})
    body.insert(2, new_p)  # after "Before." (index 1, heading is 0)
    from docx.oxml.ns import qn as _qn
    from lxml import etree
    r = etree.SubElement(new_p, f"{reingest.W}r")
    br = etree.SubElement(r, f"{reingest.W}br")
    br.set(f"{reingest.W}type", "page")
    d.save(str(art))

    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert payload["verdict"] == "FAST_FORWARD"
    assert payload["source_page_breaks"] == []
    assert len(payload["docx_page_breaks"]) == 1
    # not folded into the generic noise list
    assert payload["delta"] == []
    assert payload["manual"] == []


def test_pagebreak_removed_reports_source_count_higher(tmp_path):
    """Source has a manual page break the edited DOCX no longer carries: the
    ## 3b section must show source > docx (a removal), independent of the
    generic text delta."""
    src = tmp_path / "doc.md"
    src.write_text(
        "---\ntitle: T\n---\n\n# H\n\nBefore.\n\n" + "\\newpage" + "\n\nAfter.\n",
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("H", level=1)
    doc.add_paragraph("Before.")
    doc.add_paragraph("After.")  # reviewer deleted the break paragraph entirely
    art = tmp_path / "rendered.docx"
    doc.save(str(art))
    provenance.embed(art, provenance.build_provenance(src))

    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src), "--json"],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    payload = json.loads(result.stdout)
    assert len(payload["source_page_breaks"]) == 1
    assert payload["docx_page_breaks"] == []


def test_report_has_3b_section_with_counts_and_offsets(tmp_path):
    src = tmp_path / "doc.md"
    src.write_text(
        "---\ntitle: T\n---\n\n# H\n\nBefore.\n\nAfter.\n",
        encoding="utf-8",
    )
    doc = Document()
    doc.add_heading("H", level=1)
    doc.add_paragraph("Before.")
    _pagebreak_paragraph(doc)
    doc.add_paragraph("After.")
    art = tmp_path / "rendered.docx"
    doc.save(str(art))
    provenance.embed(art, provenance.build_provenance(src))

    result = subprocess.run(
        [sys.executable, str(RENDER_PY), "reingest", str(art), "--source", str(src)],
        capture_output=True, text=True, timeout=120,
    )
    assert result.returncode == 0, result.stderr
    assert "## 3b. Page breaks" in result.stdout
    assert "source markdown: 0 manual page break(s)" in result.stdout
    assert "edited DOCX: 1 manual page break(s)" in result.stdout
    assert "delta: +1" in result.stdout

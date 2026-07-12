"""
Tests for docstyle/marking_lint.py (#123): the default POSTRENDER_GATE_SCRIPT-shaped
check for a template-inherited classification marking shipped with no covering
classification.* replacement rule.
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))
sys.path.insert(0, str(REPO_ROOT / "docstyle"))

from docstyle import marking_lint as ml  # noqa: E402


def _docx_with_header(tmp_path, text, name="probe.docx") -> Path:
    doc = Document()
    hdr = doc.sections[0].header
    hdr.is_linked_to_previous = False
    hdr.paragraphs[0].text = text
    doc.add_paragraph("body")
    path = tmp_path / name
    doc.save(str(path))
    return path


def test_check_flags_marking_with_no_configured_rule(tmp_path):
    docx = _docx_with_header(tmp_path, "UNCLASS")
    uncovered = ml.check(docx, template_profile_path=None)
    assert uncovered == ["UNCLASS"]


def test_check_clean_when_rule_covers_the_marking(tmp_path):
    docx = _docx_with_header(tmp_path, "UNCLASS")
    profile = tmp_path / "profile.yaml"
    profile.write_text(
        "classification:\n"
        "  brief_replacements:\n"
        '    - find: ["UNCLASS"]\n'
        '      replace: "SNC"\n',
        encoding="utf-8",
    )
    uncovered = ml.check(docx, str(profile))
    assert uncovered == []


def test_check_clean_when_no_marking_present(tmp_path):
    docx = _docx_with_header(tmp_path, "Acme Corp")
    assert ml.check(docx, template_profile_path=None) == []


def test_check_covers_via_either_classification_key(tmp_path):
    """A rule under header_footer_replacements (the compact-profile key) still
    covers a marking found in a docx actually rendered under --profile reference
    -- this script doesn't know which profile produced the file, so it unions
    both keys rather than picking one."""
    docx = _docx_with_header(tmp_path, "CONFIDENTIAL")
    profile = tmp_path / "profile.yaml"
    profile.write_text(
        "classification:\n"
        "  header_footer_replacements:\n"
        '    - find: ["CONFIDENTIAL"]\n'
        '      replace: "GENERAL"\n',
        encoding="utf-8",
    )
    uncovered = ml.check(docx, str(profile))
    assert uncovered == []


def test_main_exit_code_1_on_finding(tmp_path, capsys):
    docx = _docx_with_header(tmp_path, "UNCLASS")
    rc = ml.main([str(docx)])
    assert rc == 1
    assert "FINDING" in capsys.readouterr().out


def test_main_exit_code_0_when_clean(tmp_path, capsys):
    docx = _docx_with_header(tmp_path, "Acme Corp")
    rc = ml.main([str(docx)])
    assert rc == 0
    assert "OK" in capsys.readouterr().out


def test_main_reads_template_profile_from_env(tmp_path, monkeypatch):
    docx = _docx_with_header(tmp_path, "UNCLASS")
    profile = tmp_path / "profile.yaml"
    profile.write_text(
        "classification:\n"
        "  brief_replacements:\n"
        '    - find: ["UNCLASS"]\n'
        '      replace: "SNC"\n',
        encoding="utf-8",
    )
    monkeypatch.setenv("TEMPLATE_PROFILE", str(profile))
    rc = ml.main([str(docx)])
    assert rc == 0


def test_main_errors_on_missing_docx(tmp_path, capsys):
    rc = ml.main([str(tmp_path / "does-not-exist.docx")])
    assert rc == 2
    assert "ERROR" in capsys.readouterr().err

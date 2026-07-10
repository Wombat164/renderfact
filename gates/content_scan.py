#!/usr/bin/env python3
"""content_scan.py: the generic post-render content-safety gate (issue #71, D18).

Opens a DOCX with python-docx and regex-scans every paragraph and every table
cell (recursively, including nested tables) for a caller-supplied pattern.
Exits 1 on any hit.

This module ships with NO default pattern. The regex is a REQUIRED parameter
so the module stays domain-neutral (D3: generic core, private skin): a
consumer skin with a hard content-safety requirement (for example "never let
a currency figure reach an artifact that goes to an external vendor") points
this gate at its OWN pattern. The issue's own currency-figure example is
motivation only; it is never baked in here as a default.

This is the generic, in-repo reference implementation of the "open docx with
python-docx, regex over every paragraph + every table cell, configurable
pattern, exit 1 on hit" building block described in issue #71. Point
container/render-doc.sh's POSTRENDER_GATE_SCRIPT at this script (directly, via
env-var pattern, or through a thin wrapper that hardcodes the consumer's own
pattern) to gate every render on it.

Pattern sources (all optional, combined; any single hit fails the gate):
    --pattern REGEX          repeatable
    --pattern-file PATH      one regex per non-blank, non-'#'-comment line
    RENDERFACT_GATE_PATTERN          env var, one regex (used when no --pattern
                              flag was given, so this script also works as a
                              zero-arg hook: render-doc.sh calls
                              POSTRENDER_GATE_SCRIPT with only the docx path)
    RENDERFACT_GATE_PATTERN_FILE     env var, same format as --pattern-file

At least one pattern must resolve from SOME source, or the gate refuses to
run (exit 2): a content-safety gate with no configured pattern would silently
pass everything, which is worse than not running at all.

Usage:
    python gates/content_scan.py <docx> --pattern 'REGEX' [--pattern REGEX ...]
    python gates/content_scan.py <docx> --pattern-file patterns.txt
    RENDERFACT_GATE_PATTERN='REGEX' python gates/content_scan.py <docx>

Exit codes: 0 clean, 1 pattern hit(s) found, 2 usage/environment error
(missing file, unreadable docx, bad regex, or no pattern configured at all).
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table


def iter_table_paragraphs(table: Table):
    """Yield every paragraph in every cell of `table`, recursing into any
    table nested inside a cell (python-docx's Document.tables only lists
    top-level tables; nested tables are reachable only via cell.tables)."""
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def scan_document(path: Path, patterns: list[re.Pattern]) -> list[str]:
    """Scan every body paragraph and every table-cell paragraph (recursively)
    against `patterns`. Returns one human-readable finding string per
    location that matches at least one pattern (first match wins per
    location; a location with two different pattern hits still reports once,
    which is enough to fail the gate and point a human at the paragraph)."""
    doc = Document(str(path))
    findings: list[str] = []

    def _check(label: str, text: str) -> None:
        if not text:
            return
        for pat in patterns:
            m = pat.search(text)
            if m:
                findings.append(f"{label}: pattern {pat.pattern!r} matched {m.group(0)!r}")
                return

    for i, para in enumerate(doc.paragraphs):
        _check(f"paragraph {i}", para.text)

    for ti, table in enumerate(doc.tables):
        for ci, para in enumerate(iter_table_paragraphs(table)):
            _check(f"table {ti} cell-paragraph {ci}", para.text)

    return findings


def _read_pattern_file(path: Path) -> list[str]:
    if not path.is_file():
        raise SystemExit(f"ERROR: pattern file not found: {path}")
    patterns = []
    for line in path.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if s and not s.startswith("#"):
            patterns.append(s)
    return patterns


def resolve_patterns(cli_patterns: list[str], cli_pattern_file: str | None,
                     env=os.environ) -> list[str]:
    """CLI flags win when given; otherwise fall back to env vars, so this
    script also works as a zero-arg render-doc.sh hook (POSTRENDER_GATE_SCRIPT
    is invoked with only the docx path, no room for --pattern on the call)."""
    raw: list[str] = list(cli_patterns)
    if cli_pattern_file:
        raw.extend(_read_pattern_file(Path(cli_pattern_file)))
    if not raw:
        env_pattern = env.get("RENDERFACT_GATE_PATTERN")
        if env_pattern:
            raw.append(env_pattern)
        env_file = env.get("RENDERFACT_GATE_PATTERN_FILE")
        if env_file:
            raw.extend(_read_pattern_file(Path(env_file)))
    return raw


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        prog="content_scan",
        description=__doc__.splitlines()[0] if __doc__ else "",
    )
    ap.add_argument("docx", help="path to the rendered DOCX to scan")
    ap.add_argument("--pattern", action="append", default=[],
                    help="regex to scan for (repeatable; any match fails the gate)")
    ap.add_argument("--pattern-file", default=None,
                    help="file of regexes, one per line ('#' comments and blank lines ignored)")
    args = ap.parse_args(argv)

    docx_path = Path(args.docx)
    if not docx_path.is_file():
        print(f"ERROR: not a file: {docx_path}", file=sys.stderr)
        return 2

    try:
        raw_patterns = resolve_patterns(args.pattern, args.pattern_file)
    except SystemExit as e:
        print(str(e), file=sys.stderr)
        return 2

    if not raw_patterns:
        print(
            "ERROR: no pattern configured. This is a generic content-safety gate with "
            "NO built-in pattern (docs/DECISIONS.md D18): pass --pattern/--pattern-file, "
            "or set RENDERFACT_GATE_PATTERN / RENDERFACT_GATE_PATTERN_FILE for zero-arg "
            "hook invocation (e.g. as render-doc.sh's POSTRENDER_GATE_SCRIPT).",
            file=sys.stderr,
        )
        return 2

    try:
        patterns = [re.compile(p) for p in raw_patterns]
    except re.error as e:
        print(f"ERROR: invalid regex pattern: {e}", file=sys.stderr)
        return 2

    try:
        findings = scan_document(docx_path, patterns)
    except Exception as e:  # a corrupt/non-docx input is a usage error, not a finding
        print(f"ERROR: could not scan {docx_path}: {e}", file=sys.stderr)
        return 2

    if findings:
        print(f"CONTENT-SAFETY GATE FAIL: {len(findings)} hit(s) in {docx_path.name}:")
        for f in findings[:100]:
            print(f"  {f}")
        return 1

    print(f"content-safety gate OK: {docx_path.name} clean against {len(patterns)} pattern(s)")
    return 0


if __name__ == "__main__":
    sys.exit(main())

#!/usr/bin/env python3
"""
bundle-annex-linux.py: assemble a full governed document (cover + body [+ annexes]) into one
DOCX, on Linux.

Cross-platform analog of a Word-COM bundler (win32com InsertFile, Windows-only). Here the
merge is done with docxcompose (preserves each appended part's section + headers/footers),
so the bundle runs inside the render container with no Word.

Steps:
  1. Strip the body's own duplicate cover section (the code-generated cover document is the
     authoritative one; a standalone body often carries its own copy so it renders
     self-contained). Deletes from the cover Heading-1 up to (excluding) the next Heading-1,
     at the python-docx element level so cover tables are dropped too. The heading text that
     marks the cover section is configurable via --cover-heading (prefix match,
     case-insensitive, leading numbering ignored).
  2. Compose cover + body (+ any extra DOCX annexes) into <PREFIX>_FULL.docx via docxcompose.
  3. Optionally convert the FULL to PDF via LibreOffice headless (--pdf).

NB this is the pragmatic Linux bundle: it does NOT re-run any consumer-side post-merge gates
(style harmonisation, full-document numbering, metadata scrub). The components are already
individually styled + numbered by the body/cover renderers, so the merged FULL is faithful;
run the consumer's own ship-gate tooling when that final scrub is required.

Usage:
  bundle-annex-linux.py --cover COVER.docx --body BODY.docx [--out FULL.docx] [--pdf]
                        [--cover-heading NAME] [ANNEX.docx ...]
"""
from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docxcompose.composer import Composer


def _para_is_h1(p_el) -> bool:
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return False
    st = ppr.find(qn("w:pStyle"))
    if st is None:
        return False
    val = (st.get(qn("w:val")) or "").lower()
    return val in ("heading1", "1") or val.startswith("heading1")


def _para_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def _norm_heading(text: str) -> str:
    return re.sub(r"^[\d.\s]+", "", text).strip().lower()


def strip_body_cover(doc, cover_heading: str) -> bool:
    """Remove the body's duplicate cover section: every block (paragraph/table) from the
    H1 whose text starts with cover_heading, up to (excluding) the next H1. Element-level,
    so it also drops the cover tables. Returns True if something was removed."""
    heading = cover_heading.strip().lower()
    body = doc.element.body
    children = [c for c in body if c.tag != qn("w:sectPr")]
    start = None
    for i, ch in enumerate(children):
        if ch.tag != qn("w:p") or not _para_is_h1(ch):
            continue
        norm = _norm_heading(_para_text(ch))
        if start is None and norm.startswith(heading):
            start = i
            continue
        if start is not None:
            end = i  # next H1 -> stop (exclusive)
            for victim in children[start:end]:
                victim.getparent().remove(victim)
            return True
    # the cover section was the last H1 (no following H1): drop to end
    if start is not None:
        for victim in children[start:]:
            victim.getparent().remove(victim)
        return True
    return False


def to_pdf(docx_path: Path) -> Path:
    out_dir = docx_path.parent
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
    )
    return docx_path.with_suffix(".pdf")


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Bundle cover + body (+ annexes) into one full DOCX (Linux)."
    )
    ap.add_argument("--cover", required=True)
    ap.add_argument("--body", required=True)
    ap.add_argument("--out", default=None)
    ap.add_argument("--pdf", action="store_true", help="also emit a PDF of the FULL via LibreOffice")
    ap.add_argument("--cover-heading", default="cover",
                    help="H1 text (prefix, case-insensitive) marking the body's own duplicate "
                         "cover section (default: cover)")
    ap.add_argument("--keep-body-cover", action="store_true",
                    help="do NOT strip the body's own cover section (default: strip it)")
    ap.add_argument("annexes", nargs="*", help="optional extra DOCX annexes to append after the body")
    a = ap.parse_args()

    cover = Path(a.cover); body = Path(a.body)
    if not cover.exists():
        raise SystemExit(f"cover missing: {cover}")
    if not body.exists():
        raise SystemExit(f"body missing: {body}")
    out = Path(a.out) if a.out else cover.parent / (cover.name.replace("_COVER", "_FULL"))

    body_doc = Document(str(body))
    if not a.keep_body_cover:
        removed = strip_body_cover(body_doc, a.cover_heading)
        print("  cover de-dup: " + ("removed the body's duplicate cover section (cover doc is authoritative)"
                                    if removed else "no body cover section found (nothing to remove)"))

    master = Document(str(cover))
    composer = Composer(master)
    composer.append(body_doc)
    appended = []
    for ax in a.annexes:
        p = Path(ax)
        if p.exists() and p.suffix.lower() == ".docx":
            composer.append(Document(str(p)))
            appended.append(p.name)

    out.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out))
    print(f"wrote {out}")
    print(f"  cover: {cover.name}")
    print(f"  body:  {body.name}")
    if appended:
        print(f"  annexes appended: {appended}")

    if a.pdf:
        pdf = to_pdf(out)
        print(f"  pdf:   {pdf.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
